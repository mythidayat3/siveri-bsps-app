import io
import os
import re
import time
import sqlite3
import openpyxl
import docx
import zipfile
from collections import Counter
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm, mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Body, Query, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from database import get_db_connection, DB_PATH, init_db, lookup_village_code, normalize_geo_name
init_db()

BASE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")

app = FastAPI(title="Backend BSPS DB")

# Aktifkan CORS untuk React frontend pada port 3000
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def find_header_and_data(sheet):
    header_row_idx = None
    headers = []
    rows = list(sheet.iter_rows(values_only=True))
    for r_idx, row in enumerate(rows):
        cleaned_cells = [str(c).strip().replace('\n', ' ').upper() if c is not None else '' for c in row]
        if 'NAMA' in cleaned_cells and any('KTP' in c or 'NIK' in c or 'KARTU' in c for c in cleaned_cells):
            header_row_idx = r_idx
            prev_row = rows[r_idx-1] if r_idx > 0 else []
            headers = []
            for i, c in enumerate(row):
                if c is not None and str(c).strip():
                    headers.append(str(c).strip().replace('\n', ' '))
                elif i < len(prev_row) and prev_row[i] is not None and str(prev_row[i]).strip():
                    headers.append(str(prev_row[i]).strip().replace('\n', ' '))
                else:
                    headers.append(f'COL_{i}')
            break
    if header_row_idx is not None:
        data_rows = []
        for row in rows[header_row_idx+1:]:
            if any(c is not None for c in row):
                data_rows.append(row)
        return headers, data_rows, header_row_idx
    return None, None, None

def clean_nik(val):
    if val is None:
        return ""
    s = str(val).strip()
    if s.endswith('.0'):
        s = s[:-2]
    return ''.join(c for c in s if c.isdigit())

@app.get("/")
@app.head("/")
def root_health_check():
    return {"status": "ok", "message": "SiVeri BSPS Backend API Online"}

def log_activity(username: str, action: str, entity_type: str = None, entity_name: str = None, details: str = None, ip_address: str = None, user_id: int = None, full_name: str = None):
    """Log an activity to activity_logs table safely."""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        if not full_name or not user_id:
            cursor.execute("SELECT id, full_name FROM users WHERE LOWER(username) = LOWER(?)", (username,))
            u = cursor.fetchone()
            if u:
                if not user_id: user_id = u['id']
                if not full_name: full_name = u['full_name']
                
        cursor.execute("""
            INSERT INTO activity_logs (user_id, username, full_name, action, entity_type, entity_name, details, ip_address)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (user_id, username, full_name or username, action, entity_type, entity_name, details, ip_address))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"Error logging activity: {e}")

def clean_province_for_export(prov_name: str) -> str:
    """
    Cleans a custom province name like:
    'SULAWESI SELATAN <perkotaan/perdesaan>' -> 'SULAWESI SELATAN'
    'SULAWESI SELATAN (Perkotaan/Perdesaan)' -> 'SULAWESI SELATAN'
    'SULAWESI SELATAN (Wilayah Perkotaan)' -> 'SULAWESI SELATAN'
    'PROVINSI SULAWESI SELATAN (Perdesaan)' -> 'SULAWESI SELATAN'
    """
    if not prov_name:
        return "SULAWESI SELATAN"
    
    name = str(prov_name).strip()
    name = re.sub(r'<[^>]*>', '', name)
    name = re.sub(r'\([^)]*\)', '', name)
    name = re.sub(r'\[[^\]]*\]', '', name)
    name = re.sub(r'\{[^}]*\}', '', name)
    name = re.sub(r'\s+', ' ', name).strip()
    
    if name.upper().startswith("PROVINSI "):
        name = name[9:].strip()
        
    return name.upper() if name else "SULAWESI SELATAN"

@app.post("/api/login")
def login(request: Request, body: dict = Body(...)):
    username = body.get("username", "").strip()
    password = body.get("password", "").strip()

    if not username or not password:
        raise HTTPException(status_code=400, detail="Username dan password wajib diisi")

    client_ip = request.client.host if request.client else None

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id, username, password, full_name, role FROM users WHERE LOWER(username) = LOWER(?)", (username,))
    user = cursor.fetchone()
    conn.close()

    if not user or user['password'] != password:
        log_activity(username=username, action="LOGIN_FAILED", entity_type="AUTH", entity_name="Login Gagal", details=f"Percobaan login gagal untuk username '{username}'", ip_address=client_ip)
        raise HTTPException(status_code=401, detail="Username atau password tidak valid")

    log_activity(username=user['username'], action="LOGIN", entity_type="AUTH", entity_name="Login Sukses", details=f"User '{user['username']}' ({user['full_name'] or user['username']}) berhasil masuk sebagai {user['role'].upper()}", ip_address=client_ip, user_id=user['id'], full_name=user['full_name'])

    return {
        "id": user['id'],
        "username": user['username'],
        "full_name": user['full_name'] or user['username'],
        "role": user['role'],
        "message": f"Selamat datang, {user['full_name'] or user['username']}!"
    }

@app.post("/api/change-password")
def change_password(request: Request, body: dict = Body(...)):
    target_username = body.get("username", "").strip()
    old_password = body.get("old_password", "").strip()
    new_password = body.get("new_password", "").strip()

    if not target_username or not new_password:
        raise HTTPException(status_code=400, detail="Username dan password baru wajib diisi")

    if len(new_password) < 6:
        raise HTTPException(status_code=400, detail="Password baru minimal 6 karakter")

    client_ip = request.client.host if request.client else None

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id, username, password, full_name, role FROM users WHERE LOWER(username) = LOWER(?)", (target_username,))
    user = cursor.fetchone()

    if not user:
        conn.close()
        raise HTTPException(status_code=404, detail="Pengguna tidak ditemukan")

    if old_password and user['password'] != old_password:
        conn.close()
        log_activity(username=target_username, action="CHANGE_PASSWORD_FAILED", entity_type="USER", entity_name=target_username, details="Password lama salah", ip_address=client_ip, user_id=user['id'], full_name=user['full_name'])
        raise HTTPException(status_code=400, detail="Password lama tidak sesuai")

    cursor.execute("UPDATE users SET password = ? WHERE id = ?", (new_password, user['id']))
    conn.commit()
    conn.close()

    log_activity(username=user['username'], action="CHANGE_PASSWORD", entity_type="USER", entity_name=user['username'], details=f"Password berhasil diperbarui oleh '{user['username']}'", ip_address=client_ip, user_id=user['id'], full_name=user['full_name'])

    return {"message": f"Password untuk '{user['username']}' berhasil diperbarui"}

@app.get("/api/users")
def get_users():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id, username, full_name, role, created_at FROM users ORDER BY id ASC")
    users = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return users

@app.get("/api/activity-logs")
def get_activity_logs(
    page: int = 1,
    page_size: int = 50,
    search: str = "",
    username: str = "",
    action_type: str = "",
    start_date: str = "",
    end_date: str = ""
):
    conn = get_db_connection()
    cursor = conn.cursor()

    conditions = []
    params = []

    if search.strip():
        s = f"%{search.strip()}%"
        conditions.append("(username LIKE ? OR full_name LIKE ? OR action LIKE ? OR entity_name LIKE ? OR details LIKE ?)")
        params.extend([s, s, s, s, s])

    if username.strip():
        conditions.append("LOWER(username) = LOWER(?)")
        params.append(username.strip())

    if action_type.strip():
        conditions.append("action LIKE ?")
        params.append(f"%{action_type.strip()}%")

    if start_date.strip():
        conditions.append("date(created_at) >= date(?)")
        params.append(start_date.strip())

    if end_date.strip():
        conditions.append("date(created_at) <= date(?)")
        params.append(end_date.strip())

    where_clause = " WHERE " + " AND ".join(conditions) if conditions else ""

    # Total count
    cursor.execute(f"SELECT COUNT(*) FROM activity_logs {where_clause}", params)
    total_count = cursor.fetchone()[0]

    # Stats
    cursor.execute("SELECT COUNT(*) FROM activity_logs WHERE date(created_at) = date('now')")
    today_count = cursor.fetchone()[0]

    cursor.execute("SELECT username, COUNT(*) as cnt FROM activity_logs GROUP BY username ORDER BY cnt DESC LIMIT 1")
    top_user_row = cursor.fetchone()
    top_user = top_user_row['username'] if top_user_row else "-"

    cursor.execute("SELECT action, COUNT(*) as cnt FROM activity_logs GROUP BY action ORDER BY cnt DESC LIMIT 1")
    top_action_row = cursor.fetchone()
    top_action = top_action_row['action'] if top_action_row else "-"

    # Distinct users for filter dropdown
    cursor.execute("SELECT DISTINCT username, full_name FROM users ORDER BY full_name ASC, username ASC")
    user_options = [dict(r) for r in cursor.fetchall()]

    # Fetch records with pagination
    offset = (max(1, page) - 1) * page_size
    query = f"""
        SELECT id, user_id, username, full_name, action, entity_type, entity_name, details, ip_address, created_at
        FROM activity_logs
        {where_clause}
        ORDER BY created_at DESC, id DESC
        LIMIT ? OFFSET ?
    """
    cursor.execute(query, params + [page_size, offset])
    logs = [dict(row) for row in cursor.fetchall()]
    conn.close()

    return {
        "logs": logs,
        "total": total_count,
        "page": page,
        "page_size": page_size,
        "total_pages": (total_count + page_size - 1) // page_size if total_count > 0 else 1,
        "stats": {
            "total_logs": total_count,
            "today_logs": today_count,
            "top_user": top_user,
            "top_action": top_action
        },
        "user_options": user_options
    }

@app.delete("/api/activity-logs/clear")
def clear_activity_logs():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM activity_logs")
    conn.commit()
    conn.close()
    return {"message": "Seluruh log aktifitas berhasil dibersihkan"}

@app.get("/api/provinces")
def get_provinces():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id, name FROM provinces ORDER BY id ASC")
    provinces = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return provinces

@app.post("/api/provinces")
def create_province(body: dict = Body(...)):
    raw_name = body.get("name", "").strip()
    if not raw_name:
        raise HTTPException(status_code=400, detail="Nama provinsi tidak boleh kosong")
        
    match = re.search(r'<([^>]+)>', raw_name)
    if match:
        scope = match.group(1).strip()
        base = raw_name[:match.start()].strip().upper()
        formatted_scope = "/".join([s.strip().capitalize() for s in scope.split("/")]) if "/" in scope else scope.strip().capitalize()
        name = f"{base} ({formatted_scope})"
    else:
        name = raw_name.strip()
        if not name.endswith(')'):
            name = name.upper()
        
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("INSERT INTO provinces (name) VALUES (?)", (name,))
        conn.commit()
        prov_id = cursor.lastrowid
    except sqlite3.IntegrityError:
        cursor.execute("SELECT id FROM provinces WHERE name = ?", (name,))
        prov_id = cursor.fetchone()['id']
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"Gagal menambahkan provinsi: {str(e)}")
    conn.close()
    
    log_activity(
        username="Admin",
        action="CREATE_PROVINCE",
        entity_type="PROVINCE",
        entity_name=name,
        details=f"Menambahkan wilayah provinsi baru: '{name}'"
    )
    
    return {"id": prov_id, "name": name, "message": f"Provinsi '{name}' berhasil ditambahkan"}

@app.delete("/api/provinces/{province_id}")
def delete_province(province_id: int):
    conn = get_db_connection()
    conn.execute("PRAGMA foreign_keys = ON;")
    cursor = conn.cursor()
    cursor.execute("SELECT id, name FROM provinces WHERE id = ?", (province_id,))
    prov = cursor.fetchone()
    if not prov:
        conn.close()
        raise HTTPException(status_code=404, detail="Provinsi tidak ditemukan")
    
    # Hapus data terkait secara eksplisit untuk jaminan pembersihan data
    cursor.execute("SELECT id FROM invers_stages WHERE province_id = ?", (province_id,))
    stage_ids = [r['id'] for r in cursor.fetchall()]
    if stage_ids:
        ph = ",".join("?" for _ in stage_ids)
        cursor.execute(f"DELETE FROM verified_batches WHERE stage_id IN ({ph})", stage_ids)
        cursor.execute(f"DELETE FROM invers_revisions WHERE stage_id IN ({ph})", stage_ids)
        cursor.execute(f"DELETE FROM reconciliation_overrides WHERE stage_id IN ({ph})", stage_ids)
        cursor.execute(f"DELETE FROM invers_manual_pairs WHERE stage_id IN ({ph})", stage_ids)
        cursor.execute(f"DELETE FROM invers_stages WHERE id IN ({ph})", stage_ids)
        
    cursor.execute("SELECT id FROM sk_dirjen_batches WHERE province_id = ?", (province_id,))
    sk_batch_ids = [r['id'] for r in cursor.fetchall()]
    if sk_batch_ids:
        ph_sk = ",".join("?" for _ in sk_batch_ids)
        cursor.execute(f"DELETE FROM sk_dirjen_batches WHERE id IN ({ph_sk})", sk_batch_ids)

    cursor.execute("DELETE FROM provinces WHERE id = ?", (province_id,))
    conn.commit()
    conn.close()
    return {"status": "success", "message": f"Provinsi '{prov['name']}' beserta seluruh data terkait berhasil dihapus"}

@app.get("/api/stages")
def get_stages(province_id: int = Query(None)):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    if not province_id or province_id == 1:
        filter_sql = "WHERE (s.province_id = 1 OR s.province_id IS NULL OR s.province_id = 0)"
        params = ()
    else:
        filter_sql = "WHERE s.province_id = ?"
        params = (province_id,)
        
    cursor.execute(f"""
        SELECT s.id, s.name, s.province_id, s.created_at, 
               (SELECT COUNT(*) FROM invers_records ir 
                JOIN invers_revisions irv ON ir.revision_id = irv.id 
                WHERE irv.stage_id = s.id AND irv.is_active = 1) as record_count,
               (SELECT MAX(revision_num) FROM invers_revisions WHERE stage_id = s.id) as max_revision
        FROM invers_stages s
        {filter_sql}
        ORDER BY s.created_at DESC
    """, params)
    stages = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return stages

@app.post("/api/invers/upload")
async def upload_invers(stage_name: str = Form(...), province_id: int = Form(1), file: UploadFile = File(...)):
    if not stage_name.strip():
        raise HTTPException(status_code=400, detail="Nama tahap tidak boleh kosong")
    
    file_bytes = await file.read()
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Gagal membaca file Excel: {str(e)}")
    
    sheet = wb.active
    headers, data, header_row_idx = find_header_and_data(sheet)
    if headers is None:
        raise HTTPException(status_code=400, detail="Tidak dapat menemukan baris header yang berisi 'NAMA' dan 'NO KTP/NIK'")
    
    header_map = {}
    for idx, h in enumerate(headers):
        h_upper = h.upper().replace(' ', '')
        if 'NAMA' in h_upper and 'PENGGANTI' not in h_upper:
            header_map['nama'] = idx
        elif 'KTP' in h_upper or 'NIK' in h_upper:
            header_map['no_ktp'] = idx
        elif 'KK' in h_upper or 'KELUARGA' in h_upper:
            header_map['no_kk'] = idx
        elif 'DESA' in h_upper and 'KODE' in h_upper:
            header_map['kode_desa'] = idx
        elif 'DESA' in h_upper or 'KELURAHAN' in h_upper:
            header_map['desa_kelurahan'] = idx
        elif 'KECAMATAN' in h_upper:
            header_map['kecamatan'] = idx
        elif 'KABUPATEN' in h_upper or 'KOTA' in h_upper:
            header_map['kabupaten_kota'] = idx
        elif 'ALAMAT' in h_upper:
            header_map['alamat'] = idx
        elif 'PROVINSI' in h_upper:
            header_map['provinsi'] = idx
        elif 'DELINIASI' in h_upper:
            header_map['deliniasi'] = idx
        elif 'CATATAN' in h_upper or 'KATALOG' in h_upper:
            header_map['catatan_katalog'] = idx
        elif 'PENGGUSUL' in h_upper or 'PENGUSUL' in h_upper:
            header_map['pengusul'] = idx
        elif 'TAHAP' in h_upper:
            header_map['tahap'] = idx
        elif 'URUT' in h_upper or h_upper == 'NO':
            header_map['no_urut'] = idx
            
    required_cols = ['nama', 'no_ktp', 'no_kk']
    missing_cols = [c for c in required_cols if c not in header_map]
    if missing_cols:
        raise HTTPException(status_code=400, detail=f"Kolom wajib tidak ditemukan di sheet: {', '.join(missing_cols)}")
        
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT id FROM invers_stages WHERE name = ? AND province_id = ?", (stage_name, province_id))
    stage_row = cursor.fetchone()
    if stage_row:
        stage_id = stage_row['id']
    else:
        cursor.execute("INSERT INTO invers_stages (name, province_id) VALUES (?, ?)", (stage_name, province_id))
        stage_id = cursor.lastrowid
        
    cursor.execute("SELECT MAX(revision_num) as max_rev FROM invers_revisions WHERE stage_id = ?", (stage_id,))
    rev_row = cursor.fetchone()
    next_rev = (rev_row['max_rev'] or 0) + 1
    
    cursor.execute("UPDATE invers_revisions SET is_active = 0 WHERE stage_id = ?", (stage_id,))
    
    cursor.execute("""
        INSERT INTO invers_revisions (stage_id, revision_num, filename, is_active)
        VALUES (?, ?, ?, 1)
    """, (stage_id, next_rev, file.filename))
    revision_id = cursor.lastrowid
    
    inserted_count = 0
    duplicate_warnings = []
    seen_in_file = set()
    
    for row in data:
        no_urut = row[header_map.get('no_urut')] if 'no_urut' in header_map and header_map['no_urut'] < len(row) else None
        kode_desa = str(row[header_map.get('kode_desa')]).strip() if 'kode_desa' in header_map and header_map['kode_desa'] < len(row) and row[header_map['kode_desa']] is not None else None
        nama = str(row[header_map.get('nama')]).strip() if row[header_map['nama']] is not None else ""
        jenis_kelamin = str(row[header_map.get('jenis_kelamin')]).strip() if 'jenis_kelamin' in header_map and header_map['jenis_kelamin'] < len(row) and row[header_map['jenis_kelamin']] is not None else None
        no_ktp = str(row[header_map.get('no_ktp')]).strip() if row[header_map['no_ktp']] is not None else ""
        no_kk = str(row[header_map.get('no_kk')]).strip() if row[header_map['no_kk']] is not None else ""
        alamat = str(row[header_map.get('alamat')]).strip() if 'alamat' in header_map and header_map['alamat'] < len(row) and row[header_map['alamat']] is not None else None
        desa_kelurahan = str(row[header_map.get('desa_kelurahan')]).strip() if 'desa_kelurahan' in header_map and header_map['desa_kelurahan'] < len(row) and row[header_map['desa_kelurahan']] is not None else None
        kecamatan = str(row[header_map.get('kecamatan')]).strip() if 'kecamatan' in header_map and header_map['kecamatan'] < len(row) and row[header_map['kecamatan']] is not None else None
        kabupaten_kota = str(row[header_map.get('kabupaten_kota')]).strip() if 'kabupaten_kota' in header_map and header_map['kabupaten_kota'] < len(row) and row[header_map['kabupaten_kota']] is not None else None
        provinsi = str(row[header_map.get('provinsi')]).strip() if 'provinsi' in header_map and header_map['provinsi'] < len(row) and row[header_map['provinsi']] is not None else None
        deliniasi = str(row[header_map.get('deliniasi')]).strip() if 'deliniasi' in header_map and header_map['deliniasi'] < len(row) and row[header_map['deliniasi']] is not None else None
        catatan_katalog = str(row[header_map.get('catatan_katalog')]).strip() if 'catatan_katalog' in header_map and header_map['catatan_katalog'] < len(row) and row[header_map['catatan_katalog']] is not None else None
        pengusul = str(row[header_map.get('pengusul')]).strip() if 'pengusul' in header_map and header_map['pengusul'] < len(row) and row[header_map['pengusul']] is not None else None
        tahap = str(row[header_map.get('tahap')]).strip() if 'tahap' in header_map and header_map['tahap'] < len(row) and row[header_map['tahap']] is not None else None

        if no_ktp.endswith('.0'): no_ktp = no_ktp[:-2]
        if no_kk.endswith('.0'): no_kk = no_kk[:-2]

        if not nama and not no_ktp and not no_kk:
            continue

        key_all = (nama, no_ktp, no_kk)
        if no_ktp in [k[1] for k in seen_in_file] or no_kk in [k[2] for k in seen_in_file]:
            duplicate_warnings.append(f"Duplikat ditemukan di file: {nama} (NIK: {no_ktp}, KK: {no_kk})")
        seen_in_file.add(key_all)

        cursor.execute("""
            INSERT INTO invers_records (
                revision_id, no_urut, kode_desa, nama, jenis_kelamin, no_ktp, no_kk,
                alamat, desa_kelurahan, kecamatan, kabupaten_kota, provinsi,
                deliniasi, catatan_katalog, pengusul, tahap
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (revision_id, no_urut, kode_desa, nama, jenis_kelamin, no_ktp, no_kk,
              alamat, desa_kelurahan, kecamatan, kabupaten_kota, provinsi,
              deliniasi, catatan_katalog, pengusul, tahap))
        inserted_count += 1
        
    log_activity(
        username="Admin",
        action="UPLOAD_INVERS",
        entity_type="INVERS",
        entity_name=stage_name,
        details=f"Unggah data INVERS '{stage_name}': {inserted_count} baris data (Revisi #{next_rev_num})"
    )
    
    conn.commit()
    conn.close()
    
    return {
        "stage_id": stage_id,
        "revision_num": next_rev,
        "filename": file.filename,
        "inserted_records": inserted_count,
        "warnings": duplicate_warnings[:10]
    }

@app.post("/api/verified/upload")
async def upload_verified(
    stage_id: int = Form(...),
    batch_name: str = Form(...),
    file: UploadFile = File(...)
):
    if not batch_name.strip():
        raise HTTPException(status_code=400, detail="Nama Berita Acara / Batch tidak boleh kosong")
        
    file_bytes = await file.read()
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Gagal membaca file Excel: {str(e)}")
        
    sheet_names = wb.sheetnames
    if len(sheet_names) < 2:
        raise HTTPException(status_code=400, detail="Excel harus memiliki minimal 2 sheet (Sheet 1: Lamp.IIA, Sheet 2: Lamp.IIIA)")
        
    ws_lolos = wb.worksheets[0]
    headers_lolos, data_lolos, _ = find_header_and_data(ws_lolos)
    if headers_lolos is None:
        raise HTTPException(status_code=400, detail="Tidak dapat menemukan baris header di sheet 1 (Lamp.IIA)")
        
    ws_tidak = wb.worksheets[1]
    headers_tidak, data_tidak, _ = find_header_and_data(ws_tidak)
    if headers_tidak is None:
        raise HTTPException(status_code=400, detail="Tidak dapat menemukan baris header di sheet 2 (Lamp.IIIA)")
        
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT id FROM invers_stages WHERE id = ?", (stage_id,))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=400, detail="Tahap INVERS yang dipilih tidak terdaftar")
        
    cursor.execute("SELECT COALESCE(MAX(sort_order), 0) + 1 FROM verified_batches WHERE stage_id = ?", (stage_id,))
    next_sort_order = cursor.fetchone()[0]
    cursor.execute("INSERT INTO verified_batches (stage_id, name, sort_order) VALUES (?, ?, ?)", (stage_id, batch_name, next_sort_order))
    batch_id = cursor.lastrowid
    
    h_lolos_map = {}
    for idx, h in enumerate(headers_lolos):
        h_upper = h.upper().replace(' ', '')
        if 'NAMA' in h_upper: h_lolos_map['nama'] = idx
        elif 'KTP' in h_upper or 'NIK' in h_upper: h_lolos_map['no_ktp'] = idx
        elif 'KK' in h_upper or 'KELUARGA' in h_upper: h_lolos_map['no_kk'] = idx
        elif 'URUT' in h_upper or h_upper == 'NO.': h_lolos_map['no_urut'] = idx
        elif 'DESA' in h_upper and 'KODE' in h_upper: h_lolos_map['kode_desa'] = idx
        elif 'DESA' in h_upper or 'KELURAHAN' in h_upper: h_lolos_map['desa_kelurahan'] = idx
        elif 'KECAMATAN' in h_upper: h_lolos_map['kecamatan'] = idx
        elif 'KABUPATEN' in h_upper or 'KOTA' in h_upper: h_lolos_map['kabupaten_kota'] = idx
        elif 'ALAMAT' in h_upper: h_lolos_map['alamat'] = idx
        elif 'LATITUDE' in h_upper: h_lolos_map['latitude'] = idx
        elif 'LONGITUDE' in h_upper: h_lolos_map['longitude'] = idx
        elif 'TAHAP' in h_upper: h_lolos_map['tahap'] = idx
        elif 'TANGGAL' in h_upper: h_lolos_map['tanggal'] = idx
        elif 'KETERANGAN' in h_upper: h_lolos_map['keterangan'] = idx
        elif 'KELAMIN' in h_upper or 'JENIS' in h_upper: h_lolos_map['jenis_kelamin'] = idx
        
    h_tidak_map = {}
    for idx, h in enumerate(headers_tidak):
        h_upper = h.upper().replace(' ', '').replace('\n', '')
        if h_upper == 'NO.': h_tidak_map['no_urut'] = idx
        elif idx == 1: h_tidak_map['nama'] = idx
        elif idx == 2: h_tidak_map['jenis_kelamin'] = idx
        elif idx == 3: h_tidak_map['no_ktp'] = idx
        elif idx == 4: h_tidak_map['no_kk'] = idx
        elif idx == 5: h_tidak_map['alamat'] = idx
        elif idx == 6: h_tidak_map['desa_kelurahan'] = idx
        elif idx == 7: h_tidak_map['kecamatan'] = idx
        elif idx == 8: h_tidak_map['kabupaten_kota'] = idx
        elif idx == 9: h_tidak_map['alasan_tidak_lolos'] = idx
        elif idx == 10: h_tidak_map['bnba'] = idx
        elif idx == 11: h_tidak_map['nama_pengganti'] = idx
        elif idx == 12: h_tidak_map['jenis_kelamin_pengganti'] = idx
        elif idx == 13: h_tidak_map['no_ktp_pengganti'] = idx
        elif idx == 14: h_tidak_map['no_kk_pengganti'] = idx
        elif idx == 15: h_tidak_map['alamat_pengganti'] = idx
        elif idx == 16: h_tidak_map['desa_kelurahan_pengganti'] = idx
        elif idx == 17: h_tidak_map['kecamatan_pengganti'] = idx
        elif idx == 18: h_tidak_map['kabupaten_pengganti'] = idx
        elif idx == 19: h_tidak_map['tahap'] = idx
        elif idx == 20: h_tidak_map['tanggal'] = idx
        elif idx == 21: h_tidak_map['keterangan'] = idx

    if 'nama' not in h_tidak_map:
        for idx, h in enumerate(headers_tidak):
            h_upper = h.upper()
            if 'NAMA' in h_upper and 'PENGGANTI' not in h_upper: h_tidak_map['nama'] = idx
            elif 'KTP' in h_upper and 'PENGGANTI' not in h_upper: h_tidak_map['no_ktp'] = idx
            elif 'KK' in h_upper and 'PENGGANTI' not in h_upper: h_tidak_map['no_kk'] = idx
            elif 'ALASAN' in h_upper: h_tidak_map['alasan_tidak_lolos'] = idx
            elif 'NAMA' in h_upper and 'PENGGANTI' in h_upper: h_tidak_map['nama_pengganti'] = idx
            elif 'KTP' in h_upper and 'PENGGANTI' in h_upper: h_tidak_map['no_ktp_pengganti'] = idx
            elif 'KK' in h_upper and 'PENGGANTI' in h_upper: h_tidak_map['no_kk_pengganti'] = idx
            
    cursor.execute("""
        SELECT no_ktp, no_kk, status FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        WHERE vb.stage_id = ?
    """, (stage_id,))
    # Map: key = NIK atau KK, value = set of statuses
    previously_verified_map = {}
    for row in cursor.fetchall():
        nik = row['no_ktp'].strip()
        kk = row['no_kk'].strip()
        status = row['status']
        if nik not in previously_verified_map:
            previously_verified_map[nik] = set()
        previously_verified_map[nik].add(status)
        if kk not in previously_verified_map:
            previously_verified_map[kk] = set()
        previously_verified_map[kk].add(status)
    
    def is_duplicate(nik, kk, status):
        nik_match = nik in previously_verified_map and status in previously_verified_map[nik]
        kk_match = kk in previously_verified_map and status in previously_verified_map[kk]
        return nik_match or kk_match
    
    def register_verified(nik, kk, status):
        if nik not in previously_verified_map:
            previously_verified_map[nik] = set()
        previously_verified_map[nik].add(status)
        if kk not in previously_verified_map:
            previously_verified_map[kk] = set()
        previously_verified_map[kk].add(status)
    
    stats = {
        "lolos_total": 0,
        "lolos_added": 0,
        "tidak_lolos_total": 0,
        "tidak_lolos_added": 0
    }
    
    # Lolos
    for row in data_lolos:
        nama = str(row[h_lolos_map.get('nama')]).strip() if h_lolos_map.get('nama') is not None and h_lolos_map['nama'] < len(row) and row[h_lolos_map['nama']] is not None else ""
        no_ktp = str(row[h_lolos_map.get('no_ktp')]).strip() if h_lolos_map.get('no_ktp') is not None and h_lolos_map['no_ktp'] < len(row) and row[h_lolos_map['no_ktp']] is not None else ""
        no_kk = str(row[h_lolos_map.get('no_kk')]).strip() if h_lolos_map.get('no_kk') is not None and h_lolos_map['no_kk'] < len(row) and row[h_lolos_map['no_kk']] is not None else ""
        
        if no_ktp.endswith('.0'): no_ktp = no_ktp[:-2]
        if no_kk.endswith('.0'): no_kk = no_kk[:-2]
        
        if not nama and not no_ktp and not no_kk:
            continue
            
        stats["lolos_total"] += 1
        if is_duplicate(no_ktp, no_kk, 'LOLOS'):
            is_dup = 1
        else:
            is_dup = 0
            register_verified(no_ktp, no_kk, 'LOLOS')
            stats["lolos_added"] += 1
            
        no_urut = row[h_lolos_map.get('no_urut')] if 'no_urut' in h_lolos_map and h_lolos_map['no_urut'] < len(row) else None
        kode_desa = str(row[h_lolos_map.get('kode_desa')]).strip() if 'kode_desa' in h_lolos_map and h_lolos_map['kode_desa'] < len(row) and row[h_lolos_map['kode_desa']] is not None else None
        jenis_kelamin = str(row[h_lolos_map.get('jenis_kelamin')]).strip() if 'jenis_kelamin' in h_lolos_map and h_lolos_map['jenis_kelamin'] < len(row) and row[h_lolos_map['jenis_kelamin']] is not None else None
        alamat = str(row[h_lolos_map.get('alamat')]).strip() if 'alamat' in h_lolos_map and h_lolos_map['alamat'] < len(row) and row[h_lolos_map['alamat']] is not None else None
        desa_kelurahan = str(row[h_lolos_map.get('desa_kelurahan')]).strip() if 'desa_kelurahan' in h_lolos_map and h_lolos_map['desa_kelurahan'] < len(row) and row[h_lolos_map['desa_kelurahan']] is not None else None
        kecamatan = str(row[h_lolos_map.get('kecamatan')]).strip() if 'kecamatan' in h_lolos_map and h_lolos_map['kecamatan'] < len(row) and row[h_lolos_map['kecamatan']] is not None else None
        kabupaten_kota = str(row[h_lolos_map.get('kabupaten_kota')]).strip() if 'kabupaten_kota' in h_lolos_map and h_lolos_map['kabupaten_kota'] < len(row) and row[h_lolos_map['kabupaten_kota']] is not None else None
        lat = row[h_lolos_map.get('latitude')] if 'latitude' in h_lolos_map and h_lolos_map['latitude'] < len(row) and row[h_lolos_map['latitude']] != '' else None
        lng = row[h_lolos_map.get('longitude')] if 'longitude' in h_lolos_map and h_lolos_map['longitude'] < len(row) and row[h_lolos_map['longitude']] != '' else None
        tahap = str(row[h_lolos_map.get('tahap')]).strip() if 'tahap' in h_lolos_map and h_lolos_map['tahap'] < len(row) and row[h_lolos_map['tahap']] is not None else None
        tanggal = str(row[h_lolos_map.get('tanggal')]).strip() if 'tanggal' in h_lolos_map and h_lolos_map['tanggal'] < len(row) and row[h_lolos_map['tanggal']] is not None else None
        keterangan = str(row[h_lolos_map.get('keterangan')]).strip() if 'keterangan' in h_lolos_map and h_lolos_map['keterangan'] < len(row) and row[h_lolos_map['keterangan']] is not None else None

        if not kode_desa or str(kode_desa).upper() in ("NONE", "NULL", "NAN", ""):
            kode_desa = lookup_village_code(conn, kabupaten_kota, kecamatan, desa_kelurahan)

        cursor.execute("""
            INSERT INTO verified_records (
                batch_id, no_urut, kode_desa, nama, jenis_kelamin, no_ktp, no_kk,
                alamat, desa_kelurahan, kecamatan, kabupaten_kota, status,
                latitude, longitude, tahap, tanggal, keterangan, is_duplicate_in_previous
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'LOLOS', ?, ?, ?, ?, ?, ?)
        """, (batch_id, no_urut, kode_desa, nama, jenis_kelamin, no_ktp, no_kk,
              alamat, desa_kelurahan, kecamatan, kabupaten_kota, lat, lng, tahap, tanggal, keterangan, is_dup))

    # Tidak Lolos
    for row in data_tidak:
        nama = str(row[h_tidak_map.get('nama')]).strip() if h_tidak_map.get('nama') is not None and h_tidak_map['nama'] < len(row) and row[h_tidak_map['nama']] is not None else ""
        no_ktp = str(row[h_tidak_map.get('no_ktp')]).strip() if h_tidak_map.get('no_ktp') is not None and h_tidak_map['no_ktp'] < len(row) and row[h_tidak_map['no_ktp']] is not None else ""
        no_kk = str(row[h_tidak_map.get('no_kk')]).strip() if h_tidak_map.get('no_kk') is not None and h_tidak_map['no_kk'] < len(row) and row[h_tidak_map['no_kk']] is not None else ""
        
        if no_ktp.endswith('.0'): no_ktp = no_ktp[:-2]
        if no_kk.endswith('.0'): no_kk = no_kk[:-2]
        
        if not nama and not no_ktp and not no_kk:
            continue
            
        stats["tidak_lolos_total"] += 1
        if is_duplicate(no_ktp, no_kk, 'TIDAK LOLOS'):
            is_dup = 1
        else:
            is_dup = 0
            register_verified(no_ktp, no_kk, 'TIDAK LOLOS')
            stats["tidak_lolos_added"] += 1
            
        no_urut = row[h_tidak_map.get('no_urut')] if 'no_urut' in h_tidak_map and h_tidak_map['no_urut'] < len(row) else None
        kode_desa = str(row[h_tidak_map.get('kode_desa')]).strip() if 'kode_desa' in h_tidak_map and h_tidak_map['kode_desa'] < len(row) and row[h_tidak_map['kode_desa']] is not None else None
        jenis_kelamin = str(row[h_tidak_map.get('jenis_kelamin')]).strip() if 'jenis_kelamin' in h_tidak_map and h_tidak_map['jenis_kelamin'] < len(row) and row[h_tidak_map['jenis_kelamin']] is not None else None
        alamat = str(row[h_tidak_map.get('alamat')]).strip() if 'alamat' in h_tidak_map and h_tidak_map['alamat'] < len(row) and row[h_tidak_map['alamat']] is not None else None
        desa_kelurahan = str(row[h_tidak_map.get('desa_kelurahan')]).strip() if 'desa_kelurahan' in h_tidak_map and h_tidak_map['desa_kelurahan'] < len(row) and row[h_tidak_map['desa_kelurahan']] is not None else None
        kecamatan = str(row[h_tidak_map.get('kecamatan')]).strip() if 'kecamatan' in h_tidak_map and h_tidak_map['kecamatan'] < len(row) and row[h_tidak_map['kecamatan']] is not None else None
        kabupaten_kota = str(row[h_tidak_map.get('kabupaten_kota')]).strip() if 'kabupaten_kota' in h_tidak_map and h_tidak_map['kabupaten_kota'] < len(row) and row[h_tidak_map['kabupaten_kota']] is not None else None
        alasan = str(row[h_tidak_map.get('alasan_tidak_lolos')]).strip() if 'alasan_tidak_lolos' in h_tidak_map and h_tidak_map['alasan_tidak_lolos'] < len(row) and row[h_tidak_map['alasan_tidak_lolos']] is not None else None
        tahap = str(row[h_tidak_map.get('tahap')]).strip() if 'tahap' in h_tidak_map and h_tidak_map['tahap'] < len(row) and row[h_tidak_map['tahap']] is not None else None
        tanggal = str(row[h_tidak_map.get('tanggal')]).strip() if 'tanggal' in h_tidak_map and h_tidak_map['tanggal'] < len(row) and row[h_tidak_map['tanggal']] is not None else None
        keterangan = str(row[h_tidak_map.get('keterangan')]).strip() if 'keterangan' in h_tidak_map and h_tidak_map['keterangan'] < len(row) and row[h_tidak_map['keterangan']] is not None else None

        if not kode_desa or str(kode_desa).upper() in ("NONE", "NULL", "NAN", ""):
            kode_desa = lookup_village_code(conn, kabupaten_kota, kecamatan, desa_kelurahan)

        cursor.execute("""
            INSERT INTO verified_records (
                batch_id, no_urut, kode_desa, nama, jenis_kelamin, no_ktp, no_kk,
                alamat, desa_kelurahan, kecamatan, kabupaten_kota, status,
                tahap, tanggal, alasan_tidak_lolos, keterangan, is_duplicate_in_previous
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'TIDAK LOLOS', ?, ?, ?, ?, ?)
        """, (batch_id, no_urut, kode_desa, nama, jenis_kelamin, no_ktp, no_kk,
              alamat, desa_kelurahan, kecamatan, kabupaten_kota, tahap, tanggal, alasan, keterangan, is_dup))
        record_id = cursor.lastrowid
        
        # Pengganti
        nama_pengganti = str(row[h_tidak_map.get('nama_pengganti')]).strip() if h_tidak_map.get('nama_pengganti') is not None and h_tidak_map['nama_pengganti'] < len(row) and row[h_tidak_map['nama_pengganti']] is not None else ""
        no_ktp_pengganti = str(row[h_tidak_map.get('no_ktp_pengganti')]).strip() if h_tidak_map.get('no_ktp_pengganti') is not None and h_tidak_map['no_ktp_pengganti'] < len(row) and row[h_tidak_map['no_ktp_pengganti']] is not None else ""
        no_kk_pengganti = str(row[h_tidak_map.get('no_kk_pengganti')]).strip() if h_tidak_map.get('no_kk_pengganti') is not None and h_tidak_map['no_kk_pengganti'] < len(row) and row[h_tidak_map['no_kk_pengganti']] is not None else ""

        if no_ktp_pengganti.endswith('.0'): no_ktp_pengganti = no_ktp_pengganti[:-2]
        if no_kk_pengganti.endswith('.0'): no_kk_pengganti = no_kk_pengganti[:-2]
        
        if nama_pengganti or no_ktp_pengganti or no_kk_pengganti:
            jenis_kelamin_pengganti = str(row[h_tidak_map.get('jenis_kelamin_pengganti')]).strip() if 'jenis_kelamin_pengganti' in h_tidak_map and h_tidak_map['jenis_kelamin_pengganti'] < len(row) and row[h_tidak_map['jenis_kelamin_pengganti']] is not None else None
            alamat_pengganti = str(row[h_tidak_map.get('alamat_pengganti')]).strip() if 'alamat_pengganti' in h_tidak_map and h_tidak_map['alamat_pengganti'] < len(row) and row[h_tidak_map['alamat_pengganti']] is not None else None
            desa_kelurahan_pengganti = str(row[h_tidak_map.get('desa_kelurahan_pengganti')]).strip() if 'desa_kelurahan_pengganti' in h_tidak_map and h_tidak_map['desa_kelurahan_pengganti'] < len(row) and row[h_tidak_map['desa_kelurahan_pengganti']] is not None else None
            kecamatan_pengganti = str(row[h_tidak_map.get('kecamatan_pengganti')]).strip() if 'kecamatan_pengganti' in h_tidak_map and h_tidak_map['kecamatan_pengganti'] < len(row) and row[h_tidak_map['kecamatan_pengganti']] is not None else None
            kabupaten_pengganti = str(row[h_tidak_map.get('kabupaten_pengganti')]).strip() if 'kabupaten_pengganti' in h_tidak_map and h_tidak_map['kabupaten_pengganti'] < len(row) and row[h_tidak_map['kabupaten_pengganti']] is not None else None
            
            cursor.execute("""
                INSERT INTO replacement_events (
                    disqualified_record_id, nama_pengganti, jenis_kelamin_pengganti,
                    no_ktp_pengganti, no_kk_pengganti, alamat_pengganti,
                    desa_kelurahan_pengganti, kecamatan_pengganti, kabupaten_pengganti
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (record_id, nama_pengganti, jenis_kelamin_pengganti,
                  no_ktp_pengganti, no_kk_pengganti, alamat_pengganti,
                  desa_kelurahan_pengganti, kecamatan_pengganti, kabupaten_pengganti))
            
    log_activity(
        username="Admin",
        action="UPLOAD_VERIFIKASI",
        entity_type="VERIFIKASI",
        entity_name=batch_name,
        details=f"Unggah batch verifikasi '{batch_name}': {stats['lolos_added']} Lolos, {stats['tidak_lolos_added']} Tidak Lolos"
    )
    
    conn.commit()
    conn.close()
    REKAP_CACHE.clear()
    
    return {
        "batch_id": batch_id,
        "batch_name": batch_name,
        "stats": stats
    }

@app.post("/api/verfal/upload")
async def upload_verfal(
    stage_id: int = Form(...),
    kabupaten: str = Form(...),
    batch_name: str = Form(...),
    file: UploadFile = File(...)
):
    if not kabupaten.strip():
        raise HTTPException(status_code=400, detail="Kabupaten wajib dipilih untuk Verifikasi Faktual")
    if not batch_name.strip():
        raise HTTPException(status_code=400, detail="Nama Berita Acara / Batch tidak boleh kosong")
        
    file_bytes = await file.read()
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Gagal membaca file Excel: {str(e)}")
        
    sheet_names = wb.sheetnames
    if len(sheet_names) < 2:
        raise HTTPException(status_code=400, detail="Excel Verfal harus memiliki minimal 2 sheet (Sheet 1: Lamp.IIA, Sheet 2: Lamp.IIIA)")
        
    ws_lolos = wb.worksheets[0]
    headers_lolos, data_lolos, _ = find_header_and_data(ws_lolos)
    if headers_lolos is None:
        raise HTTPException(status_code=400, detail="Tidak dapat menemukan baris header di sheet 1 (Lamp.IIA)")
        
    ws_tidak = wb.worksheets[1]
    headers_tidak, data_tidak, _ = find_header_and_data(ws_tidak)
    if headers_tidak is None:
        raise HTTPException(status_code=400, detail="Tidak dapat menemukan baris header di sheet 2 (Lamp.IIIA)")
        
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT id FROM invers_stages WHERE id = ?", (stage_id,))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=400, detail="Tahap INVERS yang dipilih tidak terdaftar")
        
    cursor.execute("SELECT COALESCE(MAX(sort_order), 0) + 1 FROM verified_batches WHERE stage_id = ? AND batch_type = 'VERFAL'", (stage_id,))
    next_sort_order = cursor.fetchone()[0]
    kab_clean = kabupaten.strip().upper()
    cursor.execute("""
        INSERT INTO verified_batches (stage_id, name, sort_order, batch_type, kabupaten) 
        VALUES (?, ?, ?, 'VERFAL', ?)
    """, (stage_id, batch_name.strip(), next_sort_order, kab_clean))
    batch_id = cursor.lastrowid
    
    h_lolos_map = {}
    for idx, h in enumerate(headers_lolos):
        h_upper = h.upper().replace(' ', '')
        if 'NAMA' in h_upper: h_lolos_map['nama'] = idx
        elif 'KTP' in h_upper or 'NIK' in h_upper: h_lolos_map['no_ktp'] = idx
        elif 'KK' in h_upper or 'KELUARGA' in h_upper: h_lolos_map['no_kk'] = idx
        elif 'URUT' in h_upper or h_upper == 'NO.': h_lolos_map['no_urut'] = idx
        elif 'DESA' in h_upper and 'KODE' in h_upper: h_lolos_map['kode_desa'] = idx
        elif 'DESA' in h_upper or 'KELURAHAN' in h_upper: h_lolos_map['desa_kelurahan'] = idx
        elif 'KECAMATAN' in h_upper: h_lolos_map['kecamatan'] = idx
        elif 'KABUPATEN' in h_upper or 'KOTA' in h_upper: h_lolos_map['kabupaten_kota'] = idx
        elif 'ALAMAT' in h_upper: h_lolos_map['alamat'] = idx
        elif 'LOLOS' in h_upper or 'STATUS' in h_upper or 'PENGGANTI' in h_upper: h_lolos_map['status'] = idx
        elif 'LATITUDE' in h_upper: h_lolos_map['latitude'] = idx
        elif 'LONGITUDE' in h_upper: h_lolos_map['longitude'] = idx
        elif 'TAHAP' in h_upper: h_lolos_map['tahap'] = idx
        elif 'TANGGAL' in h_upper: h_lolos_map['tanggal'] = idx
        elif 'KETERANGAN' in h_upper: h_lolos_map['keterangan'] = idx
        elif 'KELAMIN' in h_upper or 'JENIS' in h_upper: h_lolos_map['jenis_kelamin'] = idx
        
    h_tidak_map = {}
    bnba_col_idx = None
    for idx, h in enumerate(headers_tidak):
        h_upper = str(h or '').upper().replace('\n', ' ').strip()
        if 'BNBA' in h_upper:
            bnba_col_idx = idx
            break
    if bnba_col_idx is None:
        bnba_col_idx = 10 if len(headers_tidak) >= 22 else 9
        
    for idx, h in enumerate(headers_tidak):
        h_upper = str(h or '').upper().replace('\n', ' ').strip()
        h_clean = h_upper.replace(' ', '')
        
        if 'BNBA' in h_upper:
            h_tidak_map['bnba'] = idx
        elif 'PENGGANTI' in h_upper or idx > bnba_col_idx:
            if 'NAMA' in h_upper and 'nama_pengganti' not in h_tidak_map:
                h_tidak_map['nama_pengganti'] = idx
            elif ('KTP' in h_upper or 'NIK' in h_upper) and 'no_ktp_pengganti' not in h_tidak_map:
                h_tidak_map['no_ktp_pengganti'] = idx
            elif 'KK' in h_upper and 'no_kk_pengganti' not in h_tidak_map:
                h_tidak_map['no_kk_pengganti'] = idx
            elif ('KELAMIN' in h_upper or 'JENIS' in h_upper) and 'jenis_kelamin_pengganti' not in h_tidak_map:
                h_tidak_map['jenis_kelamin_pengganti'] = idx
            elif 'ALAMAT' in h_upper and 'alamat_pengganti' not in h_tidak_map:
                h_tidak_map['alamat_pengganti'] = idx
            elif ('DESA' in h_upper or 'KELURAHAN' in h_upper) and 'desa_kelurahan_pengganti' not in h_tidak_map:
                h_tidak_map['desa_kelurahan_pengganti'] = idx
            elif 'KECAMATAN' in h_upper and 'kecamatan_pengganti' not in h_tidak_map:
                h_tidak_map['kecamatan_pengganti'] = idx
            elif ('KABUPATEN' in h_upper or 'KOTA' in h_upper) and 'kabupaten_pengganti' not in h_tidak_map:
                h_tidak_map['kabupaten_pengganti'] = idx
            elif 'TAHAP' in h_upper:
                h_tidak_map['tahap'] = idx
            elif 'TANGGAL' in h_upper:
                h_tidak_map['tanggal'] = idx
            elif 'KETERANGAN' in h_upper:
                h_tidak_map['keterangan'] = idx
        else:
            if h_clean == 'NO.' or 'URUT' in h_upper:
                h_tidak_map['no_urut'] = idx
            elif 'NAMA' in h_upper:
                h_tidak_map['nama'] = idx
            elif 'KELAMIN' in h_upper or 'JENIS' in h_upper:
                h_tidak_map['jenis_kelamin'] = idx
            elif 'KTP' in h_upper or 'NIK' in h_upper:
                h_tidak_map['no_ktp'] = idx
            elif 'KK' in h_upper:
                h_tidak_map['no_kk'] = idx
            elif 'ALAMAT' in h_upper:
                h_tidak_map['alamat'] = idx
            elif 'DESA' in h_upper or 'KELURAHAN' in h_upper:
                h_tidak_map['desa_kelurahan'] = idx
            elif 'KECAMATAN' in h_upper:
                h_tidak_map['kecamatan'] = idx
            elif 'KABUPATEN' in h_upper or 'KOTA' in h_upper:
                h_tidak_map['kabupaten_kota'] = idx
            elif 'ALASAN' in h_upper:
                h_tidak_map['alasan_tidak_lolos'] = idx
            
    cursor.execute("""
        SELECT no_ktp, no_kk, status FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        WHERE vb.stage_id = ? AND vb.batch_type = 'VERFAL'
    """, (stage_id,))
    previously_verified_map = {}
    for row in cursor.fetchall():
        nik = row['no_ktp'].strip()
        kk = (row['no_kk'] or '').strip()
        status = row['status']
        if nik not in previously_verified_map:
            previously_verified_map[nik] = set()
        previously_verified_map[nik].add(status)
        if kk and kk not in previously_verified_map:
            previously_verified_map[kk] = set()
        if kk:
            previously_verified_map[kk].add(status)
    
    def is_duplicate(nik, kk, status):
        nik_match = nik in previously_verified_map and status in previously_verified_map[nik]
        kk_match = kk and kk in previously_verified_map and status in previously_verified_map[kk]
        return 1 if (nik_match or kk_match) else 0

    stats = {"lolos": 0, "tidak_lolos": 0, "replacements": 0, "duplicates": 0}
    
    for row in data_lolos:
        nama = str(row[h_lolos_map.get('nama')]).strip().upper() if 'nama' in h_lolos_map and h_lolos_map['nama'] < len(row) and row[h_lolos_map['nama']] is not None else None
        if not nama or nama == 'NONE' or nama == '':
            continue
            
        no_ktp_raw = str(row[h_lolos_map.get('no_ktp')]).strip() if 'no_ktp' in h_lolos_map and h_lolos_map['no_ktp'] < len(row) and row[h_lolos_map['no_ktp']] is not None else ""
        no_ktp = clean_nik(no_ktp_raw)
        
        no_kk_raw = str(row[h_lolos_map.get('no_kk')]).strip() if 'no_kk' in h_lolos_map and h_lolos_map['no_kk'] < len(row) and row[h_lolos_map['no_kk']] is not None else ""
        no_kk = clean_nik(no_kk_raw)
        
        raw_status = str(row[h_lolos_map.get('status')]).strip().upper() if 'status' in h_lolos_map and h_lolos_map['status'] < len(row) and row[h_lolos_map['status']] is not None else "LOLOS"
        status = "TIDAK LOLOS" if "TIDAK" in raw_status else "LOLOS"
        
        no_urut = row[h_lolos_map.get('no_urut')] if 'no_urut' in h_lolos_map and h_lolos_map['no_urut'] < len(row) else None
        alamat = str(row[h_lolos_map.get('alamat')]).strip() if 'alamat' in h_lolos_map and h_lolos_map['alamat'] < len(row) and row[h_lolos_map['alamat']] is not None else None
        desa = str(row[h_lolos_map.get('desa_kelurahan')]).strip().upper() if 'desa_kelurahan' in h_lolos_map and h_lolos_map['desa_kelurahan'] < len(row) and row[h_lolos_map['desa_kelurahan']] is not None else None
        kec = str(row[h_lolos_map.get('kecamatan')]).strip().upper() if 'kecamatan' in h_lolos_map and h_lolos_map['kecamatan'] < len(row) and row[h_lolos_map['kecamatan']] is not None else None
        kab_row = str(row[h_lolos_map.get('kabupaten_kota')]).strip().upper() if 'kabupaten_kota' in h_lolos_map and h_lolos_map['kabupaten_kota'] < len(row) and row[h_lolos_map['kabupaten_kota']] is not None else None
        kab = kab_row if kab_row and kab_row != 'NONE' else kab_clean
        
        tahap = str(row[h_lolos_map.get('tahap')]).strip() if 'tahap' in h_lolos_map and h_lolos_map['tahap'] < len(row) and row[h_lolos_map['tahap']] is not None else None
        tanggal = str(row[h_lolos_map.get('tanggal')]).strip() if 'tanggal' in h_lolos_map and h_lolos_map['tanggal'] < len(row) and row[h_lolos_map['tanggal']] is not None else None
        keterangan = str(row[h_lolos_map.get('keterangan')]).strip() if 'keterangan' in h_lolos_map and h_lolos_map['keterangan'] < len(row) and row[h_lolos_map['keterangan']] is not None else None
        jenis_kelamin = str(row[h_lolos_map.get('jenis_kelamin')]).strip().upper() if 'jenis_kelamin' in h_lolos_map and h_lolos_map['jenis_kelamin'] < len(row) and row[h_lolos_map['jenis_kelamin']] is not None else None
        
        dup = is_duplicate(no_ktp, no_kk, status)
        if dup: stats["duplicates"] += 1
        if status == 'LOLOS': stats["lolos"] += 1
        else: stats["tidak_lolos"] += 1
        
        cursor.execute("""
            INSERT INTO verified_records (
                batch_id, no_urut, nama, jenis_kelamin, no_ktp, no_kk,
                alamat, desa_kelurahan, kecamatan, kabupaten_kota,
                status, tahap, tanggal, keterangan, is_duplicate_in_previous
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (batch_id, no_urut, nama, jenis_kelamin, no_ktp, no_kk,
              alamat, desa, kec, kab,
              status, tahap, tanggal, keterangan, dup))

    for row in data_tidak:
        nama = str(row[h_tidak_map.get('nama')]).strip().upper() if 'nama' in h_tidak_map and h_tidak_map['nama'] < len(row) and row[h_tidak_map['nama']] is not None else None
        if not nama or nama == 'NONE' or nama == '':
            continue
            
        no_ktp_raw = str(row[h_tidak_map.get('no_ktp')]).strip() if 'no_ktp' in h_tidak_map and h_tidak_map['no_ktp'] < len(row) and row[h_tidak_map['no_ktp']] is not None else ""
        no_ktp = clean_nik(no_ktp_raw)
        
        no_kk_raw = str(row[h_tidak_map.get('no_kk')]).strip() if 'no_kk' in h_tidak_map and h_tidak_map['no_kk'] < len(row) and row[h_tidak_map['no_kk']] is not None else ""
        no_kk = clean_nik(no_kk_raw)
        
        no_urut = row[h_tidak_map.get('no_urut')] if 'no_urut' in h_tidak_map and h_tidak_map['no_urut'] < len(row) else None
        alamat = str(row[h_tidak_map.get('alamat')]).strip() if 'alamat' in h_tidak_map and h_tidak_map['alamat'] < len(row) and row[h_tidak_map['alamat']] is not None else None
        desa = str(row[h_tidak_map.get('desa_kelurahan')]).strip().upper() if 'desa_kelurahan' in h_tidak_map and h_tidak_map['desa_kelurahan'] < len(row) and row[h_tidak_map['desa_kelurahan']] is not None else None
        kec = str(row[h_tidak_map.get('kecamatan')]).strip().upper() if 'kecamatan' in h_tidak_map and h_tidak_map['kecamatan'] < len(row) and row[h_tidak_map['kecamatan']] is not None else None
        kab_row = str(row[h_tidak_map.get('kabupaten_kota')]).strip().upper() if 'kabupaten_kota' in h_tidak_map and h_tidak_map['kabupaten_kota'] < len(row) and row[h_tidak_map['kabupaten_kota']] is not None else None
        kab = kab_row if kab_row and kab_row != 'NONE' else kab_clean
        
        alasan = str(row[h_tidak_map.get('alasan_tidak_lolos')]).strip() if 'alasan_tidak_lolos' in h_tidak_map and h_tidak_map['alasan_tidak_lolos'] < len(row) and row[h_tidak_map['alasan_tidak_lolos']] is not None else None
        tahap = str(row[h_tidak_map.get('tahap')]).strip() if 'tahap' in h_tidak_map and h_tidak_map['tahap'] < len(row) and row[h_tidak_map['tahap']] is not None else None
        tanggal = str(row[h_tidak_map.get('tanggal')]).strip() if 'tanggal' in h_tidak_map and h_tidak_map['tanggal'] < len(row) and row[h_tidak_map['tanggal']] is not None else None
        keterangan = str(row[h_tidak_map.get('keterangan')]).strip() if 'keterangan' in h_tidak_map and h_tidak_map['keterangan'] < len(row) and row[h_tidak_map['keterangan']] is not None else None
        jenis_kelamin = str(row[h_tidak_map.get('jenis_kelamin')]).strip().upper() if 'jenis_kelamin' in h_tidak_map and h_tidak_map['jenis_kelamin'] < len(row) and row[h_tidak_map['jenis_kelamin']] is not None else None
        
        dup = is_duplicate(no_ktp, no_kk, 'TIDAK LOLOS')
        if dup: stats["duplicates"] += 1
        stats["tidak_lolos"] += 1
        
        cursor.execute("""
            INSERT INTO verified_records (
                batch_id, no_urut, nama, jenis_kelamin, no_ktp, no_kk,
                alamat, desa_kelurahan, kecamatan, kabupaten_kota,
                status, alasan_tidak_lolos, tahap, tanggal, keterangan, is_duplicate_in_previous
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'TIDAK LOLOS', ?, ?, ?, ?, ?)
        """, (batch_id, no_urut, nama, jenis_kelamin, no_ktp, no_kk,
              alamat, desa, kec, kab,
              alasan, tahap, tanggal, keterangan, dup))
        record_id = cursor.lastrowid
        
        nama_pengganti = str(row[h_tidak_map.get('nama_pengganti')]).strip().upper() if 'nama_pengganti' in h_tidak_map and h_tidak_map['nama_pengganti'] < len(row) and row[h_tidak_map['nama_pengganti']] is not None else None
        if nama_pengganti and nama_pengganti != 'NONE' and nama_pengganti != '':
            stats["replacements"] += 1
            no_ktp_pengganti_raw = str(row[h_tidak_map.get('no_ktp_pengganti')]).strip() if 'no_ktp_pengganti' in h_tidak_map and h_tidak_map['no_ktp_pengganti'] < len(row) and row[h_tidak_map['no_ktp_pengganti']] is not None else ""
            no_ktp_pengganti = clean_nik(no_ktp_pengganti_raw)
            
            no_kk_pengganti_raw = str(row[h_tidak_map.get('no_kk_pengganti')]).strip() if 'no_kk_pengganti' in h_tidak_map and h_tidak_map['no_kk_pengganti'] < len(row) and row[h_tidak_map['no_kk_pengganti']] is not None else ""
            no_kk_pengganti = clean_nik(no_kk_pengganti_raw)
            
            jenis_kelamin_pengganti = str(row[h_tidak_map.get('jenis_kelamin_pengganti')]).strip().upper() if 'jenis_kelamin_pengganti' in h_tidak_map and h_tidak_map['jenis_kelamin_pengganti'] < len(row) and row[h_tidak_map['jenis_kelamin_pengganti']] is not None else None
            alamat_pengganti = str(row[h_tidak_map.get('alamat_pengganti')]).strip() if 'alamat_pengganti' in h_tidak_map and h_tidak_map['alamat_pengganti'] < len(row) and row[h_tidak_map['alamat_pengganti']] is not None else None
            desa_kelurahan_pengganti = str(row[h_tidak_map.get('desa_kelurahan_pengganti')]).strip().upper() if 'desa_kelurahan_pengganti' in h_tidak_map and h_tidak_map['desa_kelurahan_pengganti'] < len(row) and row[h_tidak_map['desa_kelurahan_pengganti']] is not None else None
            kecamatan_pengganti = str(row[h_tidak_map.get('kecamatan_pengganti')]).strip().upper() if 'kecamatan_pengganti' in h_tidak_map and h_tidak_map['kecamatan_pengganti'] < len(row) and row[h_tidak_map['kecamatan_pengganti']] is not None else None
            kabupaten_pengganti = str(row[h_tidak_map.get('kabupaten_pengganti')]).strip().upper() if 'kabupaten_pengganti' in h_tidak_map and h_tidak_map['kabupaten_pengganti'] < len(row) and row[h_tidak_map['kabupaten_pengganti']] is not None else kab_clean
            
            cursor.execute("""
                INSERT INTO replacement_events (
                    disqualified_record_id, nama_pengganti, jenis_kelamin_pengganti,
                    no_ktp_pengganti, no_kk_pengganti, alamat_pengganti,
                    desa_kelurahan_pengganti, kecamatan_pengganti, kabupaten_pengganti
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (record_id, nama_pengganti, jenis_kelamin_pengganti,
                  no_ktp_pengganti, no_kk_pengganti, alamat_pengganti,
                  desa_kelurahan_pengganti, kecamatan_pengganti, kabupaten_pengganti))
            
    log_activity(
        username="Admin",
        action="UPLOAD_VERFAL",
        entity_type="VERFAL",
        entity_name=f"{kab_clean} - {batch_name}",
        details=f"Unggah Berita Acara Verfal '{batch_name}' ({kab_clean}): {stats['lolos']} Lolos, {stats['tidak_lolos']} Tidak Lolos, {stats['replacements']} Pengganti"
    )
    
    conn.commit()
    conn.close()
    REKAP_CACHE.clear()
    
    return {
        "batch_id": batch_id,
        "batch_name": batch_name,
        "kabupaten": kab_clean,
        "stats": stats
    }

@app.get("/api/stage/{stage_id}/summary")
def get_stage_summary(stage_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT id, revision_num, filename FROM invers_revisions WHERE stage_id = ? AND is_active = 1", (stage_id,))
    active_rev = cursor.fetchone()
    
    cursor.execute("SELECT id, name, uploaded_at, is_published, nomor_ba, tanggal_ba, sort_order FROM verified_batches WHERE stage_id = ? AND (batch_type = 'REGULAR' OR batch_type IS NULL) ORDER BY sort_order ASC, uploaded_at ASC, id ASC", (stage_id,))
    batches = [dict(row) for row in cursor.fetchall()]
    
    lolos_reg = 0
    tidak_lolos_reg = 0
    replacement_reg = 0
    
    for b in batches:
        cursor.execute("SELECT COUNT(*) as cnt FROM verified_records WHERE batch_id = ? AND status = 'LOLOS'", (b['id'],))
        l_cnt = cursor.fetchone()['cnt']
        lolos_reg += l_cnt
        b['lolos_count'] = l_cnt
        
        cursor.execute("SELECT COUNT(*) as cnt FROM verified_records WHERE batch_id = ? AND status = 'TIDAK LOLOS'", (b['id'],))
        tl_cnt = cursor.fetchone()['cnt']
        tidak_lolos_reg += tl_cnt
        b['tidak_lolos_count'] = tl_cnt
        
        cursor.execute("""
            SELECT COUNT(*) as cnt FROM replacement_events re
            JOIN verified_records vr ON re.disqualified_record_id = vr.id
            WHERE vr.batch_id = ?
        """, (b['id'],))
        r_cnt = cursor.fetchone()['cnt']
        replacement_reg += r_cnt
        b['replacement_count'] = r_cnt

    # Hitung data verifikasi verfal pada tahap ini
    cursor.execute("""
        SELECT vr.status, COUNT(*) as cnt
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        WHERE vb.stage_id = ? AND vb.batch_type = 'VERFAL'
        GROUP BY vr.status
    """, (stage_id,))
    verfal_status_counts = {r['status']: r['cnt'] for r in cursor.fetchall()}
    lolos_verfal = verfal_status_counts.get('LOLOS', 0)
    tidak_lolos_verfal = verfal_status_counts.get('TIDAK LOLOS', 0)
    
    cursor.execute("""
        SELECT COUNT(*) as cnt FROM replacement_events re
        JOIN verified_records vr ON re.disqualified_record_id = vr.id
        JOIN verified_batches vb ON vr.batch_id = vb.id
        WHERE vb.stage_id = ? AND vb.batch_type = 'VERFAL'
    """, (stage_id,))
    replacement_verfal = cursor.fetchone()['cnt']
    
    conn.close()
    
    lolos_combined = lolos_reg + lolos_verfal
    tidak_lolos_combined = tidak_lolos_reg + tidak_lolos_verfal
    total_verified_combined = lolos_combined + tidak_lolos_combined
    replacement_combined = replacement_reg + replacement_verfal
    
    return {
        "active_revision": dict(active_rev) if active_rev else None,
        "batches": batches,
        "totals": {
            "lolos": lolos_combined,
            "tidak_lolos": tidak_lolos_combined,
            "replacements": replacement_combined,
            "total_verified": total_verified_combined,
            "regular": {
                "lolos": lolos_reg,
                "tidak_lolos": tidak_lolos_reg,
                "total": lolos_reg + tidak_lolos_reg,
                "replacements": replacement_reg
            },
            "verfal": {
                "lolos": lolos_verfal,
                "tidak_lolos": tidak_lolos_verfal,
                "total": lolos_verfal + tidak_lolos_verfal,
                "replacements": replacement_verfal
            }
        }
    }

@app.get("/api/stage/{stage_id}/verfal-batches-grouped")
def get_verfal_batches_grouped(stage_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # 1. Daftar seluruh Kabupaten pada tahap aktif
    cursor.execute("""
        SELECT DISTINCT UPPER(TRIM(ir.kabupaten_kota)) as kab
        FROM invers_records ir
        JOIN invers_revisions irv ON ir.revision_id = irv.id
        WHERE irv.stage_id = ? AND irv.is_active = 1 AND TRIM(COALESCE(ir.kabupaten_kota, '')) != ''
        UNION
        SELECT DISTINCT UPPER(TRIM(COALESCE(vb.kabupaten, vr.kabupaten_kota))) as kab
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        WHERE vb.stage_id = ? AND TRIM(COALESCE(vb.kabupaten, vr.kabupaten_kota, '')) != ''
        ORDER BY kab ASC
    """, (stage_id, stage_id))
    all_kabupaten = [r['kab'] for r in cursor.fetchall()]
    
    # 2. Alokasi INVERS per kabupaten
    cursor.execute("""
        SELECT UPPER(TRIM(ir.kabupaten_kota)) as kab, COUNT(*) as cnt
        FROM invers_records ir
        JOIN invers_revisions irv ON ir.revision_id = irv.id
        WHERE irv.stage_id = ? AND irv.is_active = 1
        GROUP BY UPPER(TRIM(ir.kabupaten_kota))
    """, (stage_id,))
    alokasi_map = {r['kab']: r['cnt'] for r in cursor.fetchall()}
    
    # 3. Verifikasi Reguler per kabupaten
    cursor.execute("""
        SELECT UPPER(TRIM(vr.kabupaten_kota)) as kab, vr.status, COUNT(*) as cnt
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        WHERE vb.stage_id = ? AND (vb.batch_type = 'REGULAR' OR vb.batch_type IS NULL)
          AND TRIM(COALESCE(vr.kabupaten_kota, '')) != ''
        GROUP BY UPPER(TRIM(vr.kabupaten_kota)), vr.status
    """, (stage_id,))
    reg_counts = {}
    for r in cursor.fetchall():
        kab = r['kab']
        if kab not in reg_counts:
            reg_counts[kab] = {'lolos': 0, 'tidak_lolos': 0, 'total': 0, 'pengganti': 0}
        if r['status'] == 'LOLOS':
            reg_counts[kab]['lolos'] += r['cnt']
        else:
            reg_counts[kab]['tidak_lolos'] += r['cnt']
        reg_counts[kab]['total'] += r['cnt']

    cursor.execute("""
        SELECT UPPER(TRIM(vr.kabupaten_kota)) as kab, COUNT(*) as cnt
        FROM replacement_events re
        JOIN verified_records vr ON re.disqualified_record_id = vr.id
        JOIN verified_batches vb ON vr.batch_id = vb.id
        WHERE vb.stage_id = ? AND (vb.batch_type = 'REGULAR' OR vb.batch_type IS NULL)
        GROUP BY UPPER(TRIM(vr.kabupaten_kota))
    """, (stage_id,))
    for r in cursor.fetchall():
        kab = r['kab']
        if kab not in reg_counts:
            reg_counts[kab] = {'lolos': 0, 'tidak_lolos': 0, 'total': 0, 'pengganti': 0}
        reg_counts[kab]['pengganti'] = r['cnt']

    # 4. Verifikasi Verfal per kabupaten & Batch Verfal
    cursor.execute("""
        SELECT id, name, uploaded_at, is_published, nomor_ba, tanggal_ba, sort_order, kabupaten, metadata_json
        FROM verified_batches
        WHERE stage_id = ? AND batch_type = 'VERFAL'
        ORDER BY sort_order ASC, uploaded_at ASC, id ASC
    """, (stage_id,))
    all_batches = [dict(r) for r in cursor.fetchall()]
    
    verfal_batch_counts = {}
    for b in all_batches:
        b_id = b['id']
        cursor.execute("SELECT COUNT(*) as cnt FROM verified_records WHERE batch_id = ? AND status = 'LOLOS'", (b_id,))
        l_cnt = cursor.fetchone()['cnt']
        cursor.execute("SELECT COUNT(*) as cnt FROM verified_records WHERE batch_id = ? AND status = 'TIDAK LOLOS'", (b_id,))
        tl_cnt = cursor.fetchone()['cnt']
        cursor.execute("""
            SELECT COUNT(*) as cnt FROM replacement_events re
            JOIN verified_records vr ON re.disqualified_record_id = vr.id
            WHERE vr.batch_id = ?
        """, (b_id,))
        rep_cnt = cursor.fetchone()['cnt']
        
        cursor.execute("""
            SELECT alasan_tidak_lolos, COUNT(*) as cnt
            FROM verified_records
            WHERE batch_id = ? AND status = 'TIDAK LOLOS' AND alasan_tidak_lolos IS NOT NULL AND TRIM(alasan_tidak_lolos) != ''
            GROUP BY alasan_tidak_lolos
            ORDER BY cnt DESC
            LIMIT 1
        """, (b_id,))
        m_row = cursor.fetchone()
        b['alasan_tidak_lolos_terbanyak'] = m_row['alasan_tidak_lolos'] if m_row else ''
        
        b['lolos_count'] = l_cnt
        b['tidak_lolos_count'] = tl_cnt
        b['replacement_count'] = rep_cnt
        b['verifikasi_count'] = l_cnt + tl_cnt
        
        kab_key = (b.get('kabupaten') or '').strip().upper()
        if kab_key not in verfal_batch_counts:
            verfal_batch_counts[kab_key] = []
        verfal_batch_counts[kab_key].append(b)
        
    conn.close()
    
    kabupaten_groups = []
    g_verif = 0
    g_lolos = 0
    g_tidak = 0
    g_pengganti = 0
    g_alokasi = 0

    g_verif_reg = 0
    g_lolos_reg = 0
    g_tidak_reg = 0
    g_rep_reg = 0

    g_verif_verfal = 0
    g_lolos_verfal = 0
    g_tidak_verfal = 0
    g_rep_verfal = 0
    
    for kab in all_kabupaten:
        b_list = verfal_batch_counts.get(kab, [])
        kab_alokasi = alokasi_map.get(kab, 0)
        
        # Verfal numbers for this kabupaten
        v_verif = sum(b['verifikasi_count'] for b in b_list)
        v_lolos = sum(b['lolos_count'] for b in b_list)
        v_tidak = sum(b['tidak_lolos_count'] for b in b_list)
        v_pengganti = sum(b['replacement_count'] for b in b_list)

        # Reguler numbers for this kabupaten
        r_info = reg_counts.get(kab, {'lolos': 0, 'tidak_lolos': 0, 'total': 0, 'pengganti': 0})
        r_verif = r_info['total']
        r_lolos = r_info['lolos']
        r_tidak = r_info['tidak_lolos']
        r_pengganti = r_info.get('pengganti', 0)

        # Combined numbers for this kabupaten
        c_verif = v_verif + r_verif
        c_lolos = v_lolos + r_lolos
        c_tidak = v_tidak + r_tidak
        c_pengganti = v_pengganti + r_pengganti

        g_alokasi += kab_alokasi
        g_verif += c_verif
        g_lolos += c_lolos
        g_tidak += c_tidak
        g_pengganti += c_pengganti

        g_verif_reg += r_verif
        g_lolos_reg += r_lolos
        g_tidak_reg += r_tidak
        g_rep_reg += r_pengganti

        g_verif_verfal += v_verif
        g_lolos_verfal += v_lolos
        g_tidak_verfal += v_tidak
        g_rep_verfal += v_pengganti
        
        kabupaten_groups.append({
            "kabupaten": kab,
            "total_alokasi_invers": kab_alokasi,
            "batches": b_list,
            "totals": {
                "batches_count": len(b_list),
                "verifikasi": c_verif,
                "lolos": c_lolos,
                "tidak_lolos": c_tidak,
                "pengganti": c_pengganti,
                "belum_verifikasi": max(0, kab_alokasi - c_verif),
                "regular": {
                    "verifikasi": r_verif,
                    "lolos": r_lolos,
                    "tidak_lolos": r_tidak,
                    "pengganti": r_pengganti
                },
                "verfal": {
                    "verifikasi": v_verif,
                    "lolos": v_lolos,
                    "tidak_lolos": v_tidak,
                    "pengganti": v_pengganti,
                    "batches_count": len(b_list)
                }
            }
        })
        
    return {
        "kabupaten_groups": kabupaten_groups,
        "grand_totals": {
            "alokasi": g_alokasi,
            "verifikasi": g_verif,
            "lolos": g_lolos,
            "tidak_lolos": g_tidak,
            "pengganti": g_pengganti,
            "belum_verifikasi": max(0, g_alokasi - g_verif),
            "regular": {
                "verifikasi": g_verif_reg,
                "lolos": g_lolos_reg,
                "tidak_lolos": g_tidak_reg,
                "pengganti": g_rep_reg
            },
            "verfal": {
                "verifikasi": g_verif_verfal,
                "lolos": g_lolos_verfal,
                "tidak_lolos": g_tidak_verfal,
                "pengganti": g_rep_verfal
            }
        }
    }

@app.get("/api/batch/{batch_id}/kabupaten-breakdown")
def get_batch_kabupaten_breakdown(batch_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT id FROM verified_batches WHERE id = ?", (batch_id,))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="Batch tidak ditemukan")

    cursor.execute("""
        SELECT kabupaten_kota, COUNT(*) as cnt FROM verified_records
        WHERE batch_id = ? AND status = 'LOLOS' AND kabupaten_kota IS NOT NULL
        GROUP BY kabupaten_kota
    """, (batch_id,))
    lolos_map = {row['kabupaten_kota']: row['cnt'] for row in cursor.fetchall()}

    cursor.execute("""
        SELECT kabupaten_kota, COUNT(*) as cnt FROM verified_records
        WHERE batch_id = ? AND status = 'TIDAK LOLOS' AND kabupaten_kota IS NOT NULL
        GROUP BY kabupaten_kota
    """, (batch_id,))
    tidak_lolos_map = {row['kabupaten_kota']: row['cnt'] for row in cursor.fetchall()}

    cursor.execute("""
        SELECT vr.kabupaten_kota, COUNT(*) as cnt FROM replacement_events re
        JOIN verified_records vr ON re.disqualified_record_id = vr.id
        WHERE vr.batch_id = ? AND vr.kabupaten_kota IS NOT NULL
        GROUP BY vr.kabupaten_kota
    """, (batch_id,))
    replacement_map = {row['kabupaten_kota']: row['cnt'] for row in cursor.fetchall()}

    conn.close()

    all_kabs = set(lolos_map.keys()) | set(tidak_lolos_map.keys()) | set(replacement_map.keys())
    breakdown = []
    for kab in sorted(all_kabs):
        breakdown.append({
            "kabupaten": kab,
            "lolos": lolos_map.get(kab, 0),
            "tidak_lolos": tidak_lolos_map.get(kab, 0),
            "replacement": replacement_map.get(kab, 0)
        })

    return {"breakdown": breakdown}

@app.get("/api/batch/{batch_id}/export-preview")
def get_batch_export_preview(batch_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT stage_id, name FROM verified_batches WHERE id = ?", (batch_id,))
    batch_row = cursor.fetchone()
    if not batch_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Batch tidak ditemukan")
    stage_id = batch_row['stage_id']
    batch_name = batch_row['name']

    cursor.execute("SELECT name FROM invers_stages WHERE id = ?", (stage_id,))
    stage_row = cursor.fetchone()
    stage_name = stage_row['name'] if stage_row else ""

    cursor.execute("""
        SELECT vr.*, re.nama_pengganti, re.no_ktp_pengganti, re.no_kk_pengganti, re.alamat_pengganti,
               re.desa_kelurahan_pengganti as desa_pengganti, re.kecamatan_pengganti as kec_pengganti,
               re.kabupaten_pengganti as kab_pengganti, re.jenis_kelamin_pengganti,
               ro.id as override_id
        FROM verified_records vr
        LEFT JOIN replacement_events re ON re.disqualified_record_id = vr.id
        LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = ?
        WHERE vr.batch_id = ?
    """, (stage_id, batch_id))
    records = [dict(row) for row in cursor.fetchall()]
    conn.close()

    def should_include(r):
        if r['is_duplicate_in_previous'] == 0:
            return True
        return r.get('override_id') is not None

    def clean(r):
        return {k: v for k, v in r.items() if k not in ('override_id', 'is_duplicate_in_previous', 'batch_id', 'id')}

    lolos = [clean(r) for r in records if r['status'] == 'LOLOS' and should_include(r)]
    tidak_lolos = [clean(r) for r in records if r['status'] == 'TIDAK LOLOS' and should_include(r)]

    return {
        "batch_name": batch_name,
        "stage_name": stage_name,
        "lolos_records": lolos,
        "tidak_lolos_records": tidak_lolos
    }

@app.post("/api/verified/batch/{batch_id}/delete")
def delete_verified_batch(batch_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        # Ambil info batch
        cursor.execute("SELECT stage_id FROM verified_batches WHERE id = ?", (batch_id,))
        batch = cursor.fetchone()
        if not batch:
            raise HTTPException(status_code=404, detail="Batch tidak ditemukan")
        
        stage_id = batch['stage_id']
        
        # Ambil semua NIK dari batch ini SEBELUM dihapus
        cursor.execute("SELECT DISTINCT no_ktp FROM verified_records WHERE batch_id = ?", (batch_id,))
        niks_in_batch = [row['no_ktp'] for row in cursor.fetchall()]
        
        # Hapus reconciliation overrides untuk NIK-NIK tersebut
        deleted_overrides = 0
        for nik in niks_in_batch:
            cursor.execute("DELETE FROM reconciliation_overrides WHERE stage_id = ? AND original_no_ktp = ?", (stage_id, nik))
            deleted_overrides += cursor.rowcount
        
        # Hapus batch (ON DELETE CASCADE akan menghapus verified_records dan replacement_events)
        cursor.execute("DELETE FROM verified_batches WHERE id = ?", (batch_id,))
        conn.commit()
    except HTTPException:
        conn.close()
        raise
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"Gagal menghapus batch: {str(e)}")
    conn.close()
    return {"message": f"Berita Acara {batch_id} berhasil dihapus. {deleted_overrides} hasil rekonsiliasi terkait ikut terhapus."}

@app.post("/api/verified/batch/{batch_id}/toggle-published")
def toggle_batch_published(batch_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT id, is_published FROM verified_batches WHERE id = ?", (batch_id,))
        batch = cursor.fetchone()
        if not batch:
            raise HTTPException(status_code=404, detail="Batch tidak ditemukan")
        
        new_status = 0 if batch['is_published'] else 1
        cursor.execute("UPDATE verified_batches SET is_published = ? WHERE id = ?", (new_status, batch_id))
        conn.commit()
    except HTTPException:
        conn.close()
        raise
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"Gagal mengubah status: {str(e)}")
    conn.close()
    return {"batch_id": batch_id, "is_published": new_status}

@app.post("/api/verified/batch/{batch_id}/save-metadata")
def save_batch_metadata(batch_id: int, body: dict = Body(...)):
    nomor_ba = body.get("nomor_ba", "")
    tanggal_ba = body.get("tanggal_ba", "")
    
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT id FROM verified_batches WHERE id = ?", (batch_id,))
        if not cursor.fetchone():
            raise HTTPException(status_code=404, detail="Batch tidak ditemukan")
            
        cursor.execute("""
            UPDATE verified_batches 
            SET nomor_ba = ?, tanggal_ba = ? 
            WHERE id = ?
        """, (nomor_ba, tanggal_ba, batch_id))
        conn.commit()
    except HTTPException:
        conn.close()
        raise
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"Gagal menyimpan metadata: {str(e)}")
    conn.close()
    return {"status": "success", "message": "Metadata batch berhasil disimpan"}

@app.post("/api/verified/batch/{batch_id}/rename")
def rename_batch(batch_id: int, body: dict = Body(...)):
    new_name = body.get("name", "").strip()
    if not new_name:
        raise HTTPException(status_code=400, detail="Nama Berita Acara / Batch tidak boleh kosong")
        
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT id FROM verified_batches WHERE id = ?", (batch_id,))
        if not cursor.fetchone():
            raise HTTPException(status_code=404, detail="Batch tidak ditemukan")
            
        cursor.execute("UPDATE verified_batches SET name = ? WHERE id = ?", (new_name, batch_id))
        conn.commit()
    except HTTPException:
        conn.close()
        raise
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"Gagal mengubah nama batch: {str(e)}")
    conn.close()
    return {"batch_id": batch_id, "name": new_name, "message": "Nama Berita Acara / Batch berhasil diperbarui"}

@app.post("/api/verified/batches/reorder")
def reorder_verified_batches(body: dict = Body(...)):
    orders = body.get("orders", [])
    if not isinstance(orders, list):
        raise HTTPException(status_code=400, detail="Format orders tidak valid")
        
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        for idx, item in enumerate(orders):
            b_id = item.get("id")
            s_order = item.get("sort_order", idx + 1)
            if b_id:
                cursor.execute("UPDATE verified_batches SET sort_order = ? WHERE id = ?", (s_order, b_id))
        conn.commit()
    except Exception as e:
        conn.rollback()
        conn.close()
        raise HTTPException(status_code=500, detail=f"Gagal mengurutkan batch: {str(e)}")
    conn.close()
    return {"status": "success", "message": "Urutan Berita Acara / Batch berhasil diperbarui"}

@app.post("/api/verified/record/{record_id}/update-status")
def update_verified_record_status(record_id: int, body: dict = Body(...)):
    new_status = body.get("status", "").upper().strip()
    alasan_tidak_lolos = body.get("alasan_tidak_lolos", "")
    keterangan = body.get("keterangan", "")

    if new_status not in ["LOLOS", "TIDAK LOLOS"]:
        raise HTTPException(status_code=400, detail="Status harus 'LOLOS' atau 'TIDAK LOLOS'")
        
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT id, status, batch_id, no_ktp FROM verified_records WHERE id = ?", (record_id,))
        rec = cursor.fetchone()
        if not rec:
            raise HTTPException(status_code=404, detail="Data terverifikasi tidak ditemukan")
            
        if new_status == "TIDAK LOLOS":
            cursor.execute("""
                UPDATE verified_records 
                SET status = ?, alasan_tidak_lolos = ?, keterangan = ? 
                WHERE id = ?
            """, (new_status, alasan_tidak_lolos, keterangan, record_id))
        else:
            cursor.execute("""
                UPDATE verified_records 
                SET status = ? 
                WHERE id = ?
            """, (new_status, record_id))
        conn.commit()
    except HTTPException:
        conn.close()
        raise
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"Gagal mengolah pembaruan status: {str(e)}")
    conn.close()
    return {"status": "success", "record_id": record_id, "new_status": new_status, "message": f"Status berhasil diubah menjadi {new_status}"}

@app.post("/api/verified/records/bulk-update-status")
def bulk_update_verified_records_status(body: dict = Body(...)):
    record_ids = body.get("record_ids", [])
    new_status = body.get("status", "").upper().strip()
    if not record_ids or not isinstance(record_ids, list):
        raise HTTPException(status_code=400, detail="Daftar CPB tidak boleh kosong")
    if new_status not in ["LOLOS", "TIDAK LOLOS"]:
        raise HTTPException(status_code=400, detail="Status harus 'LOLOS' atau 'TIDAK LOLOS'")
        
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        placeholders = ",".join("?" for _ in record_ids)
        cursor.execute(f"UPDATE verified_records SET status = ? WHERE id IN ({placeholders})", [new_status] + record_ids)
        conn.commit()
    except Exception as e:
        conn.rollback()
        conn.close()
    return {"status": "success", "updated_count": len(record_ids), "new_status": new_status, "message": f"{len(record_ids)} CPB berhasil diubah statusnya menjadi {new_status}"}

@app.get("/api/stage/{stage_id}/records")
def get_stage_records(stage_id: int, batch_type: str = "ALL"):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT ir.* FROM invers_records ir
        JOIN invers_revisions irv ON ir.revision_id = irv.id
        WHERE irv.stage_id = ? AND irv.is_active = 1
    """, (stage_id,))
    invers_rows = [dict(row) for row in cursor.fetchall()]
    
    if batch_type.upper() == 'VERFAL':
        vb_filter = "vb.stage_id = ? AND vb.batch_type = 'VERFAL'"
        ro_filter = "stage_id = ? AND batch_type = 'VERFAL'"
    elif batch_type.upper() == 'ALL':
        vb_filter = "vb.stage_id = ?"
        ro_filter = "stage_id = ?"
    else:
        vb_filter = "vb.stage_id = ? AND (vb.batch_type = 'REGULAR' OR vb.batch_type IS NULL)"
        ro_filter = "stage_id = ? AND (batch_type = 'REGULAR' OR batch_type IS NULL)"

    cursor.execute(f"""
        SELECT vr.*, vb.name as batch_name, 
               re.nama_pengganti, re.no_ktp_pengganti, re.no_kk_pengganti, re.alamat_pengganti,
               re.desa_kelurahan_pengganti, re.kecamatan_pengganti, re.kabupaten_pengganti
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        LEFT JOIN replacement_events re ON re.disqualified_record_id = vr.id
        WHERE {vb_filter}
    """, (stage_id,))
    verified_rows = [dict(row) for row in cursor.fetchall()]
    
    cursor.execute(f"SELECT * FROM reconciliation_overrides WHERE {ro_filter}", (stage_id,))
    overrides = {row['original_no_ktp']: dict(row) for row in cursor.fetchall()}
    
    conn.close()
    
    invers_map = {}
    for ir in invers_rows:
        key = (ir['nama'].strip().upper(), ir['no_ktp'].strip(), ir['no_kk'].strip())
        invers_map[key] = ir
        
    invers_by_nik = {}
    invers_by_name_desa = {}
    invers_by_kk = {}
    invers_by_name = {}
    for ir in invers_rows:
        n_clean = ir['nama'].strip().upper()
        d_clean = (ir['desa_kelurahan'] or '').strip().upper()
        nik_clean = ir['no_ktp'].strip()
        kk_clean = ir['no_kk'].strip()
        
        invers_by_nik[nik_clean] = ir
        invers_by_name_desa[(n_clean, d_clean)] = ir
        invers_by_kk[kk_clean] = ir
        if n_clean not in invers_by_name:
            invers_by_name[n_clean] = []
        invers_by_name[n_clean].append(ir)
        
    analyzed_verified = []
    mismatch_count = 0
    
    for vr in verified_rows:
        nama = (vr.get('nama') or '').strip()
        nik = (vr.get('no_ktp') or '').strip()
        kk = (vr.get('no_kk') or '').strip()
        errors = []
        mismatch_type = ""
        is_mismatch = False
        vr['expected_invers'] = None
        
        if len(nik) != 16:
            errors.append(f"Panjang NIK adalah {len(nik)} karakter, seharusnya 16 digit")
            mismatch_type = "NIK_INVALID"
        if len(kk) != 16:
            errors.append(f"Panjang KK adalah {len(kk)} karakter, seharusnya 16 digit")
            mismatch_type = "KK_INVALID"
            
        if nik == kk and len(nik) > 0:
            errors.append("NIK dan KK bernilai sama")
            mismatch_type = "NIK_KK_IDENTICAL"
            
        if vr['is_duplicate_in_previous'] == 1:
            errors.append("Data sudah pernah terverifikasi di Berita Acara sebelumnya (Duplikat)")
            if not mismatch_type:
                mismatch_type = "DUPLICATE"
            
        override = overrides.get(nik)
        matched_invers = None
        
        if override:
            if override['override_type'] == 'ACCEPT_VERIFIED':
                pass
            elif override['override_type'] == 'MANUAL_EDIT':
                nama = override['corrected_nama']
                nik = override['corrected_no_ktp']
                kk = override['corrected_no_kk']
        
        key = (nama.upper(), nik, kk)
        
        if key in invers_map:
            matched_invers = invers_map[key]
        else:
            expected_invers = invers_by_nik.get(nik)
            
            # Fallback jika NIK tidak cocok, cari berdasarkan Nama + Desa, KK, atau Nama Unik
            if not expected_invers:
                n_clean = nama.upper()
                d_clean = (vr.get('desa_kelurahan') or '').strip().upper()
                kk_clean = kk
                
                if (n_clean, d_clean) in invers_by_name_desa:
                    expected_invers = invers_by_name_desa[(n_clean, d_clean)]
                elif kk_clean in invers_by_kk:
                    expected_invers = invers_by_kk[kk_clean]
                elif n_clean in invers_by_name and len(invers_by_name[n_clean]) == 1:
                    expected_invers = invers_by_name[n_clean][0]
                    
            if expected_invers:
                vr['expected_invers'] = expected_invers

            if override:
                # Mismatch resolved via reconciliation override
                is_mismatch = False
            else:
                if expected_invers:
                    mismatch_fields = []
                    if expected_invers['nama'].strip().upper() != nama.upper():
                        mismatch_fields.append("Nama")
                        if not mismatch_type: mismatch_type = "NAMA_MISMATCH"
                    if expected_invers['no_ktp'].strip() != nik:
                        mismatch_fields.append("NIK")
                        if not mismatch_type: mismatch_type = "NIK_MISMATCH"
                    if expected_invers['no_kk'].strip() != kk:
                        mismatch_fields.append("KK")
                        if not mismatch_type: mismatch_type = "KK_MISMATCH"
                    if mismatch_fields:
                        is_mismatch = True
                        mismatch_count += 1
                        errors.append(f"Ketidakcocokan dengan data INVERS pada kolom: {', '.join(mismatch_fields)}. Seharusnya Nama: '{expected_invers['nama']}', NIK: '{expected_invers['no_ktp']}', KK: '{expected_invers['no_kk']}'")
                else:
                    is_mismatch = True
                    mismatch_count += 1
                    errors.append("Data tidak ditemukan dalam database INVERS")
                    mismatch_type = "MISSING_IN_INVERS"
                
        vr['errors'] = errors
        vr['has_error'] = len(errors) > 0 or is_mismatch
        vr['is_mismatch'] = is_mismatch
        vr['mismatch_type'] = mismatch_type
        vr['override'] = override
        
        analyzed_verified.append(vr)
        
    return {
        "invers_records": invers_rows,
        "verified_records": analyzed_verified,
        "mismatch_count": mismatch_count
    }

@app.post("/api/reconciliation/override")
def save_reconciliation_override(
    stage_id: int = Form(...),
    original_no_ktp: str = Form(...),
    override_type: str = Form(...),
    corrected_nama: str = Form(None),
    corrected_no_ktp: str = Form(None),
    corrected_no_kk: str = Form(None)
):
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("""
            INSERT OR REPLACE INTO reconciliation_overrides (
                stage_id, original_no_ktp, override_type, corrected_nama, corrected_no_ktp, corrected_no_kk
            ) VALUES (?, ?, ?, ?, ?, ?)
        """, (stage_id, original_no_ktp, override_type, corrected_nama, corrected_no_ktp, corrected_no_kk))

        # Sinkronisasi Data INVERS jika memilih 'ACCEPT_VERIFIED' atau 'MANUAL_EDIT'
        if override_type in ('ACCEPT_VERIFIED', 'MANUAL_EDIT'):
            cursor.execute("""
                SELECT vr.id, vr.nama, vr.no_ktp, vr.no_kk, vr.desa_kelurahan
                FROM verified_records vr
                JOIN verified_batches vb ON vb.id = vr.batch_id
                WHERE vb.stage_id = ? AND vr.no_ktp = ?
            """, (stage_id, original_no_ktp))
            v_rec = cursor.fetchone()
            
            if v_rec:
                target_nik = corrected_no_ktp if (override_type == 'MANUAL_EDIT' and corrected_no_ktp) else v_rec['no_ktp']
                target_nama = corrected_nama if (override_type == 'MANUAL_EDIT' and corrected_nama) else v_rec['nama']
                target_kk = corrected_no_kk if (override_type == 'MANUAL_EDIT' and corrected_no_kk) else v_rec['no_kk']
                
                # Cari data rujukan di INVERS yang cocok berdasarkan Nama + Desa atau KK atau NIK lama
                cursor.execute("""
                    SELECT ir.id FROM invers_records ir
                    JOIN invers_revisions irv ON ir.revision_id = irv.id
                    WHERE irv.stage_id = ? AND irv.is_active = 1
                      AND (ir.no_ktp = ? OR (UPPER(ir.nama) = UPPER(?) AND UPPER(ir.desa_kelurahan) = UPPER(?)) OR ir.no_kk = ?)
                    LIMIT 1
                """, (stage_id, original_no_ktp, target_nama, (v_rec['desa_kelurahan'] or '').strip(), target_kk))
                inv_match = cursor.fetchone()
                
                if inv_match:
                    cursor.execute("""
                        UPDATE invers_records 
                        SET no_ktp = ?, nama = ?, no_kk = ?
                        WHERE id = ?
                    """, (target_nik, target_nama, target_kk, inv_match['id']))

        conn.commit()
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=400, detail=f"Gagal menyimpan perbaikan: {str(e)}")
        
    conn.close()
    return {"message": "Perbaikan rekonsiliasi berhasil disimpan"}

# --- ENDPOINT BARU UNTUK DASHBOARD PROGRESS BAR & OVERVIEW CENTER ---

@app.get("/api/stage/{stage_id}/overview-stats")
def get_overview_stats(stage_id: int):
    # Ambil semua data (gabungan Reguler dan Verfal)
    data = get_stage_records(stage_id, batch_type="ALL")
    invers = data["invers_records"]
    verified = data["verified_records"]
    
    total_invers = len(invers)
    
    # Sertakan semua data terverifikasi termasuk duplikat (duplikat masuk ke reconciliation)
    active_verified = verified
    
    # Hitung kategori untuk Multi-Color Progress Bar
    green_count = 0  # Lolos & Cocok (atau memiliki override)
    yellow_count = 0 # Ada di INVERS tapi mismatch Nama/KK (belum override)
    orange_count = 0 # Tidak ada di INVERS sama sekali (belum override)
    
    matched_invers_niks = set()
    
    for v in active_verified:
        # Tentukan status pencocokan
        if v.get("override"):
            # Jika ada override (sudah diselesaikan), hitung sebagai Selesai/Lolos (Hijau)
            green_count += 1
            # Tambahkan NIK target (yang terverifikasi/terkoreksi) ke matched list agar tidak dihitung Red
            target_nik = v["override"]["corrected_no_ktp"] if v["override"]["override_type"] == 'MANUAL_EDIT' else v["no_ktp"]
            matched_invers_niks.add(target_nik)
        elif v["is_mismatch"]:
            if v.get("expected_invers"):
                yellow_count += 1
                matched_invers_niks.add(v["no_ktp"])
            else:
                orange_count += 1
        else:
            green_count += 1
            matched_invers_niks.add(v["no_ktp"])
            
    # Red: Belum verifikasi (nik invers yang tidak terjamah sama sekali)
    # Tambahkan manual pairs ke matched set
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT invers_nik FROM invers_manual_pairs WHERE stage_id = ?", (stage_id,))
    for row in cursor.fetchall():
        matched_invers_niks.add(row['invers_nik'].strip())
    conn.close()
    
    red_count = 0
    for ir in invers:
        if ir["no_ktp"] not in matched_invers_niks:
            red_count += 1
            
    return {
        "total_invers": total_invers,
        "segments": {
            "green": green_count,
            "yellow": yellow_count,
            "orange": orange_count,
            "red": red_count
        }
    }

@app.get("/api/stage/{stage_id}/overview-tables")
def get_overview_tables(stage_id: int):
    data = get_stage_records(stage_id, batch_type="ALL")
    invers = data["invers_records"]
    verified = data["verified_records"]
    
    # Kunci verifikasi terpasang (gunakan record terakhir per NIK jika ada duplikat)
    verified_map = {}
    for v in verified:
        verified_map[v["no_ktp"]] = v
    
    # Tambahkan manual pairs - map invers_nik -> verified record
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT mp.invers_nik, vr.*
        FROM invers_manual_pairs mp
        JOIN verified_records vr ON vr.id = mp.verified_record_id
        WHERE mp.stage_id = ?
    """, (stage_id,))
    for row in cursor.fetchall():
        row_dict = dict(row)
        invers_nik = row_dict.pop('invers_nik', None) or row_dict.get('no_ktp', '')
        if invers_nik and invers_nik not in verified_map:
            verified_map[invers_nik] = row_dict

    # Ambil NIK yang terikat SK Dirjen
    cursor.execute("""
        SELECT DISTINCT vr.no_ktp
        FROM sk_dirjen_matches m
        JOIN verified_records vr ON vr.id = m.verified_record_id
        WHERE m.verified_stage_id = ?
          AND (m.match_type = 'PERFECT' OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED') OR m.match_type = 'MANUAL_PAIR')
    """, (stage_id,))
    sk_nik_set = set(row['no_ktp'].strip() for row in cursor.fetchall() if row['no_ktp'])
    conn.close()
            
    # 1. Agregasi Kabupaten
    kab_stats = {}
    # 2. Agregasi Kecamatan
    kec_stats = {}
    # 3. Agregasi Pengusul
    pengusul_stats = {}
    
    # Inisialisasi dari data INVERS
    for ir in invers:
        kab = (ir["kabupaten_kota"] or "LAINNYA").upper().strip()
        kec = (ir["kecamatan"] or "LAINNYA").upper().strip()
        peng = (ir["pengusul"] or "LAINNYA").upper().strip()
        nik = (ir["no_ktp"] or "").strip()
        
        # Check verified match
        v = verified_map.get(nik)
        is_lolos = 1 if v and v["status"] == "LOLOS" else 0
        is_tidak_lolos = 1 if v and v["status"] == "TIDAK LOLOS" else 0
        is_belum = 1 if not v else 0
        has_sk = 1 if (is_lolos and nik in sk_nik_set) else 0
        
        # Kabupaten
        if kab not in kab_stats:
            kab_stats[kab] = {"total_cpb": 0, "lolos": 0, "tidak_lolos": 0, "belum_verifikasi": 0, "sk_dirjen_sudah": 0, "sk_dirjen_belum": 0}
        kab_stats[kab]["total_cpb"] += 1
        kab_stats[kab]["lolos"] += is_lolos
        kab_stats[kab]["tidak_lolos"] += is_tidak_lolos
        kab_stats[kab]["belum_verifikasi"] += is_belum
        kab_stats[kab]["sk_dirjen_sudah"] += has_sk
        
        # Kecamatan
        if kec not in kec_stats:
            kec_stats[kec] = {"total_cpb": 0, "lolos": 0, "tidak_lolos": 0, "belum_verifikasi": 0, "sk_dirjen_sudah": 0, "sk_dirjen_belum": 0}
        kec_stats[kec]["total_cpb"] += 1
        kec_stats[kec]["lolos"] += is_lolos
        kec_stats[kec]["tidak_lolos"] += is_tidak_lolos
        kec_stats[kec]["belum_verifikasi"] += is_belum
        kec_stats[kec]["sk_dirjen_sudah"] += has_sk
        
        # Pengusul
        if peng not in pengusul_stats:
            pengusul_stats[peng] = {"total_cpb": 0, "lolos": 0, "tidak_lolos": 0, "belum_verifikasi": 0, "sk_dirjen_sudah": 0, "sk_dirjen_belum": 0}
        pengusul_stats[peng]["total_cpb"] += 1
        pengusul_stats[peng]["lolos"] += is_lolos
        pengusul_stats[peng]["tidak_lolos"] += is_tidak_lolos
        pengusul_stats[peng]["belum_verifikasi"] += is_belum
        pengusul_stats[peng]["sk_dirjen_sudah"] += has_sk
        
    for k in kab_stats.values(): k["sk_dirjen_belum"] = max(0, k["lolos"] - k["sk_dirjen_sudah"])
    for kc in kec_stats.values(): kc["sk_dirjen_belum"] = max(0, kc["lolos"] - kc["sk_dirjen_sudah"])
    for p in pengusul_stats.values(): p["sk_dirjen_belum"] = max(0, p["lolos"] - p["sk_dirjen_sudah"])

    return {
        "kabupaten": [{"name": k, **v} for k, v in kab_stats.items()],
        "kecamatan": [{"name": k, **v} for k, v in kec_stats.items()],
        "pengusul": [{"name": k, **v} for k, v in pengusul_stats.items()]
    }

@app.get("/api/stage/{stage_id}/pengusul-tree")
def get_pengusul_tree(stage_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # 1. Get INVERS records for hierarchy + CPB counts
    cursor.execute("""
        SELECT ir.nama, ir.no_ktp, ir.pengusul, 
               UPPER(TRIM(COALESCE(ir.kabupaten_kota, ''))) as kabupaten_kota,
               UPPER(TRIM(COALESCE(ir.kecamatan, ''))) as kecamatan,
               UPPER(TRIM(COALESCE(ir.desa_kelurahan, ''))) as desa_kelurahan
        FROM invers_records ir
        JOIN invers_revisions irv ON ir.revision_id = irv.id
        WHERE irv.stage_id = ? AND irv.is_active = 1
    """, (stage_id,))
    invers_rows = [dict(r) for r in cursor.fetchall()]
    
    # 2. Get verified records with same filter as rekap
    cursor.execute("""
        SELECT vr.no_ktp, vr.status
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
        WHERE vb.stage_id = ? AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
    """, (stage_id,))
    verified_rows = [dict(r) for r in cursor.fetchall()]
    
    # 3. Get manual pairs (invers ↔ verified)
    cursor.execute("""
        SELECT mp.invers_nik, vr.status
        FROM invers_manual_pairs mp
        JOIN verified_records vr ON vr.id = mp.verified_record_id
        WHERE mp.stage_id = ?
    """, (stage_id,))
    manual_pairs = [dict(r) for r in cursor.fetchall()]

    # 4. Get SK Dirjen matched NIKs for this stage
    cursor.execute("""
        SELECT DISTINCT vr.no_ktp
        FROM sk_dirjen_matches m
        JOIN verified_records vr ON vr.id = m.verified_record_id
        WHERE m.verified_stage_id = ?
          AND (m.match_type = 'PERFECT' OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED') OR m.match_type = 'MANUAL_PAIR')
    """, (stage_id,))
    sk_nik_set = set(row['no_ktp'].strip() for row in cursor.fetchall() if row['no_ktp'])
    
    conn.close()
    
    # Build NIK -> status map
    verified_map = {}
    for vr in verified_rows:
        verified_map[vr["no_ktp"].strip()] = vr["status"]
    
    # Add manual pairs to verified_map
    for mp in manual_pairs:
        nik = mp["invers_nik"].strip()
        if nik not in verified_map:
            verified_map[nik] = mp["status"]
    
    # Build the tree hierarchy from INVERS, count via NIK matching
    tree = {}
    for ir in invers_rows:
        peng = (ir["pengusul"] or "LAINNYA").strip()
        kab = ir["kabupaten_kota"] or "TIDAK DIKETAHUI"
        kec = ir["kecamatan"] or "LAINNYA"
        desa = ir["desa_kelurahan"] or "LAINNYA"
        nik = (ir["no_ktp"] or "").strip()
        
        if peng not in tree: tree[peng] = {}
        if kab not in tree[peng]: tree[peng][kab] = {}
        if kec not in tree[peng][kab]: tree[peng][kab][kec] = {}
        if desa not in tree[peng][kab][kec]:
            tree[peng][kab][kec][desa] = {"cpb": 0, "lolos": 0, "tidak_lolos": 0, "sk_dirjen_sudah": 0}
        
        node = tree[peng][kab][kec][desa]
        node["cpb"] += 1
        
        status = verified_map.get(nik)
        if status == "LOLOS":
            node["lolos"] += 1
            if nik in sk_nik_set:
                node["sk_dirjen_sudah"] += 1
        elif status == "TIDAK LOLOS":
            node["tidak_lolos"] += 1
    
    # Konversi ke format JSON tree
    tree_list = []
    for p_name, kabs in tree.items():
        p_node = {"name": p_name, "cpb": 0, "lolos": 0, "tidak_lolos": 0, "belum_verifikasi": 0, "sk_dirjen_sudah": 0, "sk_dirjen_belum": 0, "children": []}
        for kb_name, kecs in kabs.items():
            kb_node = {"name": kb_name, "cpb": 0, "lolos": 0, "tidak_lolos": 0, "belum_verifikasi": 0, "sk_dirjen_sudah": 0, "sk_dirjen_belum": 0, "children": []}
            for kc_name, desas in kecs.items():
                kc_node = {"name": kc_name, "cpb": 0, "lolos": 0, "tidak_lolos": 0, "belum_verifikasi": 0, "sk_dirjen_sudah": 0, "sk_dirjen_belum": 0, "children": []}
                for ds_name, stats in desas.items():
                    ds_belum = stats["cpb"] - stats["lolos"] - stats["tidak_lolos"]
                    ds_sk_sudah = stats["sk_dirjen_sudah"]
                    ds_sk_belum = max(0, stats["lolos"] - ds_sk_sudah)
                    ds_node = {
                        "name": ds_name, 
                        "cpb": stats["cpb"], 
                        "lolos": stats["lolos"], 
                        "tidak_lolos": stats["tidak_lolos"], 
                        "belum_verifikasi": ds_belum,
                        "sk_dirjen_sudah": ds_sk_sudah,
                        "sk_dirjen_belum": ds_sk_belum
                    }
                    kc_node["children"].append(ds_node)
                    kc_node["cpb"] += stats["cpb"]
                    kc_node["lolos"] += stats["lolos"]
                    kc_node["tidak_lolos"] += stats["tidak_lolos"]
                    kc_node["sk_dirjen_sudah"] += ds_sk_sudah
                kc_node["belum_verifikasi"] = kc_node["cpb"] - kc_node["lolos"] - kc_node["tidak_lolos"]
                kc_node["sk_dirjen_belum"] = max(0, kc_node["lolos"] - kc_node["sk_dirjen_sudah"])
                kb_node["children"].append(kc_node)
                kb_node["cpb"] += kc_node["cpb"]
                kb_node["lolos"] += kc_node["lolos"]
                kb_node["tidak_lolos"] += kc_node["tidak_lolos"]
                kb_node["sk_dirjen_sudah"] += kc_node["sk_dirjen_sudah"]
            kb_node["belum_verifikasi"] = kb_node["cpb"] - kb_node["lolos"] - kb_node["tidak_lolos"]
            kb_node["sk_dirjen_belum"] = max(0, kb_node["lolos"] - kb_node["sk_dirjen_sudah"])
            p_node["children"].append(kb_node)
            p_node["cpb"] += kb_node["cpb"]
            p_node["lolos"] += kb_node["lolos"]
            p_node["tidak_lolos"] += kb_node["tidak_lolos"]
            p_node["sk_dirjen_sudah"] += kb_node["sk_dirjen_sudah"]
        p_node["belum_verifikasi"] = p_node["cpb"] - p_node["lolos"] - p_node["tidak_lolos"]
        p_node["sk_dirjen_belum"] = max(0, p_node["lolos"] - p_node["sk_dirjen_sudah"])
        tree_list.append(p_node)
        
    return tree_list

@app.post("/api/database/clear")
def clear_database():
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        # Bersihkan tabel transaksi utama
        cursor.execute("DELETE FROM replacement_events")
        cursor.execute("DELETE FROM verified_records")
        cursor.execute("DELETE FROM verified_batches")
        cursor.execute("DELETE FROM invers_records")
        cursor.execute("DELETE FROM invers_revisions")
        cursor.execute("DELETE FROM invers_stages")
        cursor.execute("DELETE FROM reconciliation_overrides")
        conn.commit()
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"Gagal membersihkan data database: {str(e)}")
    conn.close()
    return {"message": "Database berhasil disetel ulang (bersih)"}

@app.post("/api/database/cleanup-overrides")
def cleanup_orphaned_overrides():
    """Hapus reconciliation overrides yang tidak memiliki record verifikasi terkait"""
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        # Hapus overrides yang NIK-nya tidak ada di verified_records
        cursor.execute("""
            DELETE FROM reconciliation_overrides 
            WHERE id IN (
                SELECT ro.id FROM reconciliation_overrides ro
                LEFT JOIN verified_records vr ON ro.original_no_ktp = vr.no_ktp 
                    AND vr.batch_id IN (SELECT id FROM verified_batches WHERE stage_id = ro.stage_id)
                WHERE vr.id IS NULL
            )
        """)
        deleted_count = cursor.rowcount
        conn.commit()
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"Gagal membersihkan data: {str(e)}")
    conn.close()
    return {"message": f"Berhasil membersihkan {deleted_count} rekonsiliasi yang sudah tidak terkait dengan data verifikasi"}

@app.post("/api/stage/{stage_id}/delete")
def delete_stage(stage_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        # Cek apakah stage ada
        cursor.execute("SELECT id, name FROM invers_stages WHERE id = ?", (stage_id,))
        stage = cursor.fetchone()
        if not stage:
            raise HTTPException(status_code=404, detail="Tahap tidak ditemukan")
        
        stage_name = stage['name']
        
        # Hapus stage (ON DELETE CASCADE akan menghapus semua data terkait):
        # - invers_revisions -> invers_records
        # - verified_batches -> verified_records -> replacement_events
        # - reconciliation_overrides
        # - invers_manual_pairs
        cursor.execute("DELETE FROM invers_stages WHERE id = ?", (stage_id,))
        conn.commit()
    except HTTPException:
        conn.close()
        raise
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"Gagal menghapus tahap: {str(e)}")
    conn.close()
    return {"message": f"Tahap '{stage_name}' berhasil dihapus beserta semua data terkait"}

@app.post("/api/stage/{stage_id}/rename")
def rename_stage(stage_id: int, body: dict = Body(...)):
    new_name = body.get("name", "").strip()
    if not new_name:
        raise HTTPException(status_code=400, detail="Nama Tahap INVERS tidak boleh kosong")
        
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT id, name FROM invers_stages WHERE id = ?", (stage_id,))
        stage = cursor.fetchone()
        if not stage:
            raise HTTPException(status_code=404, detail="Tahap tidak ditemukan")
            
        cursor.execute("SELECT id FROM invers_stages WHERE LOWER(name) = LOWER(?) AND id != ?", (new_name, stage_id))
        if cursor.fetchone():
            raise HTTPException(status_code=400, detail=f"Nama Tahap '{new_name}' sudah digunakan")
            
        cursor.execute("UPDATE invers_stages SET name = ? WHERE id = ?", (new_name, stage_id))
        cursor.execute("""
            UPDATE invers_records 
            SET tahap = ? 
            WHERE revision_id IN (SELECT id FROM invers_revisions WHERE stage_id = ?)
        """, (new_name, stage_id))
        conn.commit()
    except HTTPException:
        conn.close()
        raise
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"Gagal mengubah nama tahap: {str(e)}")
    conn.close()
    return {"stage_id": stage_id, "name": new_name, "message": "Nama Tahap INVERS berhasil diperbarui"}

@app.post("/api/upload/village-codes")
async def upload_village_codes(file: UploadFile = File(...)):
    if not file.filename.endswith(".xlsx"):
        raise HTTPException(status_code=400, detail="Hanya file Excel (.xlsx) yang diperbolehkan")
        
    contents = await file.read()
    try:
        wb = openpyxl.load_workbook(io.BytesIO(contents), data_only=True)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Gagal membaca file Excel: {str(e)}")
        
    sheet = wb.active
    rows = list(sheet.iter_rows(values_only=True))
    if not rows:
        raise HTTPException(status_code=400, detail="File Excel kosong")
        
    header_idx = None
    col_map = {}
    for r_idx, row in enumerate(rows[:20]):
        row_str = [str(c).strip().upper() if c is not None else '' for c in row]
        if any("KODE" in c and "DESA" in c for c in row_str) or any("KAB" in c or "KOTA" in c for c in row_str):
            header_idx = r_idx
            for c_idx, val in enumerate(row_str):
                if "KODE" in val and "DESA" in val:
                    col_map["kode_desa"] = c_idx
                elif "PROV" in val:
                    col_map["provinsi"] = c_idx
                elif "KAB" in val or "KOTA" in val:
                    col_map["kabupaten"] = c_idx
                elif "KEC" in val or "DISTRIK" in val:
                    col_map["kecamatan"] = c_idx
                elif "DESA" in val or "KELURAHAN" in val:
                    col_map["desa"] = c_idx
                elif "DELINEASI" in val or "DELINIASI" in val:
                    col_map["delineasi"] = c_idx
            break

    if "kode_desa" not in col_map: col_map["kode_desa"] = 1
    if "provinsi" not in col_map: col_map["provinsi"] = 2
    if "kabupaten" not in col_map: col_map["kabupaten"] = 3
    if "kecamatan" not in col_map: col_map["kecamatan"] = 4
    if "desa" not in col_map: col_map["desa"] = 5
    if "delineasi" not in col_map: col_map["delineasi"] = 6

    start_row = (header_idx + 1) if header_idx is not None else 1

    conn = get_db_connection()
    cursor = conn.cursor()

    records_to_insert = []
    for row in rows[start_row:]:
        if not row or all(c is None for c in row):
            continue
        
        def get_val(idx):
            if idx < len(row) and row[idx] is not None:
                return str(row[idx]).strip()
            return ""

        kode_desa = get_val(col_map.get("kode_desa", 1))
        provinsi = get_val(col_map.get("provinsi", 2))
        kabupaten = get_val(col_map.get("kabupaten", 3))
        kecamatan = get_val(col_map.get("kecamatan", 4))
        desa = get_val(col_map.get("desa", 5))
        delineasi = get_val(col_map.get("delineasi", 6))

        if not kode_desa and not desa:
            continue

        c_kab = normalize_geo_name(kabupaten)
        c_kec = normalize_geo_name(kecamatan)
        c_desa = normalize_geo_name(desa)

        records_to_insert.append((kode_desa, provinsi, kabupaten, kecamatan, desa, delineasi, c_kab, c_kec, c_desa))

    count = 0
    if records_to_insert:
        cursor.executemany("""
        INSERT OR REPLACE INTO village_codes 
        (kode_desa, provinsi, kabupaten_kota, kecamatan, desa_kelurahan, delineasi, clean_kab, clean_kec, clean_desa)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, records_to_insert)
        conn.commit()
        count = len(records_to_insert)

    conn.close()
    return {
        "status": "success",
        "message": f"Berhasil mengunggah & memperbarui {count:,} database Kode Desa/Kelurahan!",
        "count": count
    }

@app.get("/api/templates/download/{template_type}")
def download_template(template_type: str):
    if template_type == "village_codes":
        filepath = os.path.join(BASE_DIR, "DATABASE_KODE_DESA.xlsx")
        filename = "TEMPLATE_DATABASE_KODE_DESA.xlsx"
    elif template_type == "invers":
        filepath = os.path.join(BASE_DIR, "INVERS.xlsx")
        filename = "TEMPLATE_INVERS.xlsx"
    elif template_type == "verified":
        filepath = os.path.join(BASE_DIR, "TEMPLATE_TERVERIFIKASI.xlsx")
        filename = "TEMPLATE_VERIFIKASI.xlsx"
    elif template_type == "sk_dirjen":
        filepath = os.path.join(BASE_DIR, "TEMPLATE SK DIRJEN.xlsx")
        filename = "TEMPLATE_SK_DIRJEN.xlsx"
    elif template_type == "verfal":
        filepath = os.path.join(BASE_DIR, "TEMPLATE_VERFAL.xlsx")
        filename = "TEMPLATE_VERFAL.xlsx"
    else:
        raise HTTPException(status_code=400, detail="Tipe template tidak dikenal")
        
    if os.path.exists(filepath):
        return FileResponse(filepath, filename=filename)
        
    # Dynamic Openpyxl fallback if file does not exist locally
    wb = openpyxl.Workbook()
    ws = wb.active
    header_fill = PatternFill(start_color="1E3A8A", end_color="1E3A8A", fill_type="solid")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    thin_border = Border(
        left=Side(style='thin', color='D3D3D3'), right=Side(style='thin', color='D3D3D3'),
        top=Side(style='thin', color='D3D3D3'), bottom=Side(style='thin', color='D3D3D3')
    )

    if template_type == "village_codes":
        ws.title = "Database Kode Desa"
        headers = ["NO", "KODE DESA", "PROVINSI", "KAB./KOTA", "KECAMATAN/DISTRIK", "DESA/KELURAHAN", "DELINEASI"]
        sample_rows = [
            [1, "1101012001", "ACEH", "KAB. ACEH SELATAN", "BAKONGAN", "KEUDE BAKONGAN", "PESISIR"],
            [2, "1101012002", "ACEH", "KAB. ACEH SELATAN", "BAKONGAN", "UJONG MANGKI", "PESISIR"]
        ]
    elif template_type == "invers":
        ws.title = "INVERS"
        headers = ["NO", "NAMA", "NIK", "NO KK", "ALAMAT", "DESA/KELURAHAN", "KECAMATAN", "KABUPATEN/KOTA"]
        sample_rows = [[1, "JOHN DOE", "7301010101010001", "7301010101010002", "JL. SUDIRMAN NO. 1", "DESA A", "KECAMATAN B", "KABUPATEN C"]]
    elif template_type == "sk_dirjen":
        ws.title = "SK DIRJEN"
        headers = ["NO", "NAMA", "NIK", "NO KK", "ALAMAT", "DESA/KELURAHAN", "KECAMATAN", "KABUPATEN/KOTA"]
        sample_rows = [[1, "JANE DOE", "7301010101010003", "7301010101010004", "JL. MERDEKA NO. 2", "DESA X", "KECAMATAN Y", "KABUPATEN Z"]]
    else:
        ws.title = "Hasil Verifikasi"
        headers = ["NO", "KODE DESA", "NAMA", "NIK", "NO KK", "ALAMAT", "DESA/KELURAHAN", "KECAMATAN", "KABUPATEN/KOTA", "STATUS"]
        sample_rows = [[1, "1101012001", "JOHN DOE", "7301010101010001", "7301010101010002", "JL. SUDIRMAN NO. 1", "DESA A", "KECAMATAN B", "KABUPATEN C", "LOLOS"]]

    ws.append(headers)
    for col_idx in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    for row_data in sample_rows:
        ws.append(row_data)

    for row in ws.iter_rows(min_row=1, max_row=len(sample_rows)+1, min_col=1, max_col=len(headers)):
        for cell in row:
            cell.border = thin_border
            if cell.row > 1:
                cell.alignment = Alignment(vertical="center")

    for col in ws.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        col_letter = get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = max(max_len + 4, 14)

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    return StreamingResponse(
        output,
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'}
    )

@app.post("/api/export/docx")
async def export_docx_files(
    stage_id: int = Form(...),
    batch_id: int = Form(None),
    nomor_ba: str = Form(""),
    nomor_surat: str = Form(""),
    tanggal_ba: str = Form(""),
    lokasi_ba: str = Form(""),
    no_surat_dirjen: str = Form(""),
    tgl_surat_dirjen: str = Form(""),
    hal_surat_dirjen: str = Form("")
):
    base_dir = BASE_DIR
    path_ba_template = os.path.join(base_dir, "FORMAT BERITA ACARA.docx")
    path_sp_template = os.path.join(base_dir, "SURAT PENYAMPAIAN BA.docx")
    
    if not os.path.exists(path_ba_template) or not os.path.exists(path_sp_template):
        raise HTTPException(status_code=404, detail="File template Word (.docx) tidak ditemukan")
        
    # Ambil Data
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT s.name as stage_name, p.name as province_name 
        FROM invers_stages s 
        LEFT JOIN provinces p ON s.province_id = p.id 
        WHERE s.id = ?
    """, (stage_id,))
    stage_row = cursor.fetchone()
    stage_name = stage_row['stage_name'] if stage_row else "Tahap"
    prov_clean = clean_province_for_export(stage_row['province_name'] if stage_row and stage_row['province_name'] else "SULAWESI SELATAN")
    
    batch_name = None
    if batch_id:
        cursor.execute("SELECT name FROM verified_batches WHERE id = ?", (batch_id,))
        batch_row = cursor.fetchone()
        if batch_row:
            batch_name = batch_row['name']
        cursor.execute("""
            UPDATE verified_batches 
            SET nomor_ba = ?, tanggal_ba = ? 
            WHERE id = ?
        """, (nomor_ba, tanggal_ba, batch_id))
        conn.commit()
            
    # Total data INVERS di tahap aktif
    cursor.execute("""
        SELECT COUNT(*) as cnt FROM invers_records ir
        JOIN invers_revisions rev ON ir.revision_id = rev.id
        WHERE rev.stage_id = ? AND rev.is_active = 1
    """, (stage_id,))
    total_invers = cursor.fetchone()['cnt']
    
    # Total verifikasi aktif (Lolos + Tidak Lolos di Berita Acara yang aktif)
    if batch_id:
        cursor.execute("""
            SELECT COUNT(*) as cnt FROM verified_records vr
            JOIN verified_batches vb ON vr.batch_id = vb.id
            LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
            WHERE vr.batch_id = ? AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
        """, (batch_id,))
    else:
        cursor.execute("""
            SELECT COUNT(*) as cnt FROM verified_records vr
            JOIN verified_batches vb ON vr.batch_id = vb.id
            LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
            WHERE vb.stage_id = ? AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
        """, (stage_id,))
    total_verifikasi_aktif = cursor.fetchone()['cnt']
    
    # Total verifikasi seluruhnya di tahap aktif (untuk menghitung sisa belum verifikasi)
    cursor.execute("""
        SELECT COUNT(*) as cnt FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
        WHERE vb.stage_id = ? AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
    """, (stage_id,))
    total_verif_tahap_seluruhnya = cursor.fetchone()['cnt']
    sisa_belum_verif = max(0, total_invers - total_verif_tahap_seluruhnya)
    
    # Daftar rincian usulan per kabupaten untuk Berita Acara yang aktif
    if batch_id:
        cursor.execute("""
            SELECT vr.kabupaten_kota,
                   SUM(CASE WHEN vr.status = 'LOLOS' THEN 1 ELSE 0 END) as lolos,
                   SUM(CASE WHEN vr.status = 'TIDAK LOLOS' THEN 1 ELSE 0 END) as tidak_lolos,
                   COUNT(*) as total
            FROM verified_records vr
            LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = ?
            WHERE vr.batch_id = ? AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
            GROUP BY UPPER(TRIM(vr.kabupaten_kota))
            ORDER BY vr.kabupaten_kota ASC
        """, (stage_id, batch_id))
    else:
        cursor.execute("""
            SELECT vr.kabupaten_kota,
                   SUM(CASE WHEN vr.status = 'LOLOS' THEN 1 ELSE 0 END) as lolos,
                   SUM(CASE WHEN vr.status = 'TIDAK LOLOS' THEN 1 ELSE 0 END) as tidak_lolos,
                   COUNT(*) as total
            FROM verified_records vr
            JOIN verified_batches vb ON vr.batch_id = vb.id
            LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
            WHERE vb.stage_id = ? AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
            GROUP BY UPPER(TRIM(vr.kabupaten_kota))
            ORDER BY vr.kabupaten_kota ASC
        """, (stage_id,))
    kab_details = [dict(row) for row in cursor.fetchall()]
    conn.close()
    
    # Perihal surat penyampaian
    perihal_surat_penyampaian = f"Penyampaian Hasil Verifikasi Calon Penerima Bantuan (CPB) Kegiatan Bantuan Stimulan Perumahan Swadaya (BSPS) {stage_name} TA 2026 Provinsi {prov_clean.title()}"
    
    replacements = {
        "[NOMOR BA]": nomor_ba,
        "[Nomor Berita Acara]": nomor_ba,
        "[Nomor Surat]": nomor_surat,
        "[NAMA TAHAP]": stage_name.upper(),
        "[Nama Tahap]": stage_name,
        "[NAMA PROVINSI]": prov_clean,
        "[Nama Provinsi]": prov_clean.title(),
        "[Provinsi Aktif]": prov_clean,
        "[Tanggal Eksport BA]": tanggal_ba,
        "[Lokasi BA]": lokasi_ba,
        "[Nomor Surat Dirjen]": no_surat_dirjen,
        "[Tanggal Surat Dirjen]": tgl_surat_dirjen,
        "[Hal Surat Dirjen]": hal_surat_dirjen,
        "[Perihal Surat]": perihal_surat_penyampaian,
        "[Total Hasil Verifikasi di Tahap Aktif]": str(total_verifikasi_aktif),
        "[Total Data Invers Di Tahap aktif]": str(total_invers),
        "[Tanggal Batas Proses Verifikasi]": tanggal_ba
    }
    
    def process_document(template_path, reps):
        doc = docx.Document(template_path)
        for p in doc.paragraphs:
            for k, v in reps.items():
                if k in p.text:
                    for run in p.runs:
                        if k in run.text:
                            run.text = run.text.replace(k, v)
                    if k in p.text:
                        p.text = p.text.replace(k, v)
                        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for k, v in reps.items():
                            for run in p.runs:
                                if k in run.text:
                                    run.text = run.text.replace(k, v)
                            if k in p.text:
                                p.text = p.text.replace(k, v)
        return doc
        
    doc_ba = process_document(path_ba_template, replacements)
    doc_sp = process_document(path_sp_template, replacements)
    
    # FORMATTING OVERRIDES - 1. FORMAT BERITA ACARA
    for p in doc_ba.paragraphs:
        # baris Nomor BA samakan font dengan judul (Bookman Old Style, 11 pt, Bold)
        if p.text.strip().upper().startswith("NOMOR"):
            for run in p.runs:
                run.font.name = "Bookman Old Style"
                run.font.size = docx.shared.Pt(11)
                run.font.bold = True
        # Paragraf "Sehubungan dengan... dst." menggunakan font size 11
        if p.text.startswith("Sehubungan dengan"):
            for run in p.runs:
                run.font.size = docx.shared.Pt(11)
                
    # FORMATTING OVERRIDES - 2. SURAT PENYAMPAIAN BA
    for p in doc_sp.paragraphs:
        # baris yang menuliskan nomor dan tanggal surat menjadi Arial, font size 11, tambah [Lokasi BA] didepan tanggal
        if p.text.startswith("Nomor") and "[Tanggal Eksport BA]" in p.text:
            p.text = f"Nomor\t\t: {nomor_surat}\t\t\t\t\t{lokasi_ba}, {tanggal_ba}"
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = docx.shared.Pt(11)
                
    # Isian tabel di Surat Penyampaian BA
    if len(doc_sp.tables) > 0:
        table = doc_sp.tables[0]
        if len(table.rows) > 1:
            # Kolom Jumlah Usulan (Kolom 3 / indeks 3)
            cell_js = table.cell(1, 3)
            cell_js.text = str(total_verifikasi_aktif)
            for p in cell_js.paragraphs:
                p.alignment = 1  # Center
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = docx.shared.Pt(10)
            
            # Kolom Rincian Usulan per Kabupaten/Kota (Kolom 4 / indeks 4)
            cell_rinci = table.cell(1, 4)
            # Bersihkan isi lama
            p_first = cell_rinci.paragraphs[0]
            p_first.text = ""
            while len(cell_rinci.paragraphs) > 1:
                p_to_del = cell_rinci.paragraphs[-1]
                p_to_del._element.getparent().remove(p_to_del._element)
                
            is_first = True
            for kab_row in kab_details:
                kab_name = (kab_row['kabupaten_kota'] or "LAINNYA").upper().strip()
                total = kab_row['total']
                lolos = kab_row['lolos']
                tl = kab_row['tidak_lolos']
                
                if is_first:
                    p_kab = p_first
                    is_first = False
                else:
                    p_kab = cell_rinci.add_paragraph()
                    
                p_kab.paragraph_format.left_indent = docx.shared.Pt(12)
                run_kab = p_kab.add_run(f"•  {kab_name} (Total Verifikasi: {total} unit)")
                run_kab.bold = True
                run_kab.font.name = "Arial"
                run_kab.font.size = docx.shared.Pt(10)
                
                p_lolos = cell_rinci.add_paragraph()
                p_lolos.paragraph_format.left_indent = docx.shared.Pt(24)
                run_lolos = p_lolos.add_run(f"- Lolos: {lolos} unit")
                run_lolos.font.name = "Arial"
                run_lolos.font.size = docx.shared.Pt(9.5)
                
                p_tl = cell_rinci.add_paragraph()
                p_tl.paragraph_format.left_indent = docx.shared.Pt(24)
                run_tl = p_tl.add_run(f"- Tidak Lolos: {tl} unit")
                run_tl.font.name = "Arial"
                run_tl.font.size = docx.shared.Pt(9.5)
                
            # Kolom Keterangan (Kolom 5 / indeks 5)
            cell_ket = table.cell(1, 5)
            cell_ket.text = f"Sebanyak {total_verifikasi_aktif} unit dari total {total_invers} unit telah diverifikasi, sehingga yang masih proses verifikasi sebanyak {sisa_belum_verif} unit"
            for p in cell_ket.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = docx.shared.Pt(10)
                    
    # Simpan ke byte stream
    stream_ba = io.BytesIO()
    doc_ba.save(stream_ba)
    stream_ba.seek(0)
    
    stream_sp = io.BytesIO()
    doc_sp.save(stream_sp)
    stream_sp.seek(0)
    
    name_suffix = (batch_name or stage_name).replace(' ', '_')
    stage_upper = stage_name.upper()
    
    # Buat file ZIP berisi kedua dokumen
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr(f"{stage_upper}_BAHV_{batch_name or stage_name}.docx", stream_ba.getvalue())
        z.writestr(f"{stage_upper}_PENYAMPAIAN_BAHV_{batch_name or stage_name}.docx", stream_sp.getvalue())
    zip_buffer.seek(0)
    
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename=DRAFT_SURAT_BA_{name_suffix}.zip"}
    )

@app.post("/api/export/pdf")
async def export_pdf_files(
    stage_id: int = Form(...),
    batch_id: int = Form(None),
    nomor_ba: str = Form(""),
    nomor_surat: str = Form(""),
    tanggal_ba: str = Form(""),
    lokasi_ba: str = Form(""),
    no_surat_dirjen: str = Form(""),
    tgl_surat_dirjen: str = Form(""),
    hal_surat_dirjen: str = Form("")
):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT name FROM invers_stages WHERE id = ?", (stage_id,))
    stage_name = cursor.fetchone()['name']
    
    batch_name = None
    if batch_id:
        cursor.execute("SELECT name FROM verified_batches WHERE id = ?", (batch_id,))
        batch_row = cursor.fetchone()
        if batch_row:
            batch_name = batch_row['name']
        cursor.execute("""
            UPDATE verified_batches 
            SET nomor_ba = ?, tanggal_ba = ? 
            WHERE id = ?
        """, (nomor_ba, tanggal_ba, batch_id))
        conn.commit()
    
    cursor.execute("""
        SELECT COUNT(*) as cnt FROM invers_records ir
        JOIN invers_revisions irv ON ir.revision_id = irv.id
        WHERE irv.stage_id = ? AND irv.is_active = 1
    """, (stage_id,))
    total_invers = cursor.fetchone()['cnt']
    
    if batch_id:
        cursor.execute("""
            SELECT COUNT(*) as cnt FROM verified_records vr
            LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = ?
            WHERE vr.batch_id = ? AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
        """, (stage_id, batch_id))
    else:
        cursor.execute("""
            SELECT COUNT(*) as cnt FROM verified_records vr
            JOIN verified_batches vb ON vr.batch_id = vb.id
            LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
            WHERE vb.stage_id = ? AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
        """, (stage_id,))
    total_verifikasi_aktif = cursor.fetchone()['cnt']
    
    cursor.execute("""
        SELECT COUNT(*) as cnt FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
        WHERE vb.stage_id = ? AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
    """, (stage_id,))
    total_verif_tahap_seluruhnya = cursor.fetchone()['cnt']
    sisa_belum_verif = max(0, total_invers - total_verif_tahap_seluruhnya)
    
    if batch_id:
        cursor.execute("""
            SELECT vr.kabupaten_kota,
                   SUM(CASE WHEN vr.status = 'LOLOS' THEN 1 ELSE 0 END) as lolos,
                   SUM(CASE WHEN vr.status = 'TIDAK LOLOS' THEN 1 ELSE 0 END) as tidak_lolos,
                   COUNT(*) as total
            FROM verified_records vr
            LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = ?
            WHERE vr.batch_id = ? AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
            GROUP BY UPPER(TRIM(vr.kabupaten_kota))
            ORDER BY vr.kabupaten_kota ASC
        """, (stage_id, batch_id))
    else:
        cursor.execute("""
            SELECT vr.kabupaten_kota,
                   SUM(CASE WHEN vr.status = 'LOLOS' THEN 1 ELSE 0 END) as lolos,
                   SUM(CASE WHEN vr.status = 'TIDAK LOLOS' THEN 1 ELSE 0 END) as tidak_lolos,
                   COUNT(*) as total
            FROM verified_records vr
            JOIN verified_batches vb ON vr.batch_id = vb.id
            LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
            WHERE vb.stage_id = ? AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
            GROUP BY UPPER(TRIM(vr.kabupaten_kota))
            ORDER BY vr.kabupaten_kota ASC
        """, (stage_id,))
    kab_details = [dict(row) for row in cursor.fetchall()]
    conn.close()
    
    perihal_surat = f"Penyampaian Hasil Verifikasi Calon Penerima Bantuan (CPB) Kegiatan Bantuan Stimulan Perumahan Swadaya (BSPS) {stage_name} TA 2026 Provinsi Sulawesi Selatan"
    
    styles = getSampleStyleSheet()
    
    style_title = ParagraphStyle('TitleBA', parent=styles['Title'], fontName='Times-Bold', fontSize=14, alignment=TA_CENTER, spaceAfter=6)
    style_nomor = ParagraphStyle('NomorBA', parent=styles['Normal'], fontName='Times-Bold', fontSize=11, alignment=TA_CENTER, spaceAfter=4)
    style_subtitle = ParagraphStyle('SubtitleBA', parent=styles['Normal'], fontName='Times-Bold', fontSize=11, alignment=TA_CENTER, spaceAfter=2)
    style_normal = ParagraphStyle('NormalBA', parent=styles['Normal'], fontName='Times-Roman', fontSize=11, alignment=TA_JUSTIFY, spaceAfter=6, leading=14)
    style_body = ParagraphStyle('BodyBA', parent=styles['Normal'], fontName='Times-Roman', fontSize=11, alignment=TA_JUSTIFY, spaceAfter=8, leading=15, firstLineIndent=2*cm)
    style_footer = ParagraphStyle('FooterBA', parent=styles['Normal'], fontName='Times-Roman', fontSize=11, alignment=TA_LEFT, spaceAfter=4)
    style_small = ParagraphStyle('SmallBA', parent=styles['Normal'], fontName='Times-Roman', fontSize=9.5, alignment=TA_LEFT, spaceAfter=2, leading=12)
    
    elements_ba = []
    elements_ba.append(Spacer(1, 1*cm))
    elements_ba.append(Paragraph("BERITA ACARA", style_title))
    elements_ba.append(Paragraph(f"NOMOR : {nomor_ba}", style_nomor))
    elements_ba.append(Spacer(1, 0.3*cm))
    elements_ba.append(Paragraph(f"HASIL VERIFIKASI CALON PENERIMA BANTUAN (CPB) {stage_name.upper()}", style_subtitle))
    elements_ba.append(Paragraph("KEGIATAN BANTUAN STIMULAN PERUMAHAN SWADAYA (BSPS) TA.2026", style_subtitle))
    elements_ba.append(Paragraph("PROVINSI SULAWESI SELATAN", style_subtitle))
    elements_ba.append(Spacer(1, 0.8*cm))
    elements_ba.append(Paragraph(
        f"Sehubungan dengan pelaksanaan verifikasi calon penerima bantuan yang telah dilakukan di "
        f"Provinsi Sulawesi Selatan dengan {stage_name}, maka bersama ini disampaikan Berita Acara "
        f"Hasil Verifikasi Calon Penerima Bantuan (CPB) Kegiatan Bantuan Stimulan Perumahan Swadaya "
        f"(BSPS) Tahun Anggaran 2026.", style_body))
    elements_ba.append(Paragraph(
        f"Data usulan calon penerima bantuan di Provinsi Sulawesi Selatan dengan Alokasi "
        f"{total_invers} unit, telah dilakukan verifikasi lapangan terhadap calon penerima bantuan "
        f"dengan hasil sebagaimana berikut:", style_body))
    elements_ba.append(Spacer(1, 0.3*cm))
    elements_ba.append(Paragraph("Lampiran I  Hasil Verifikasi Calon Penerima Bantuan Kegiatan BSPS Tahun Anggaran 2026 di Provinsi Sulawesi Selatan", style_normal))
    elements_ba.append(Paragraph("Lampiran II  Daftar Calon Penerima Bantuan Kegiatan BSPS Tahun Anggaran 2026 Provinsi Sulawesi Selatan", style_normal))
    elements_ba.append(Paragraph("Lampiran III  Daftar Calon Penerima Bantuan Pengganti Kegiatan BSPS Tahun Anggaran 2026 Provinsi Sulawesi Selatan", style_normal))
    elements_ba.append(Spacer(1, 0.3*cm))
    elements_ba.append(Paragraph(
        "Penggantian data calon penerima bantuan bersumber dari data pengusul K/L, aplikasi e-RTLH, "
        "DTKS, sistem informasi lainnya, serta hasil verifikasi dan validasi lapangan.", style_body))
    elements_ba.append(Paragraph(
        "Demikian Berita Acara ini dibuat dengan sebenarnya dan dapat dipergunakan sebagaimana mestinya.", style_body))
    elements_ba.append(Spacer(1, 1.5*cm))
    elements_ba.append(Paragraph(f"{lokasi_ba}, {tanggal_ba}", style_footer))
    elements_ba.append(Spacer(1, 1.5*cm))
    
    elements_ba.append(Paragraph("Membuat Berita Acara,", style_footer))
    elements_ba.append(Spacer(1, 1.5*cm))
    elements_ba.append(Paragraph("_________________________", style_footer))
    elements_ba.append(Paragraph("Operator Verifikasi", style_small))
    
    stream_ba = io.BytesIO()
    doc_ba = SimpleDocTemplate(stream_ba, pagesize=A4, leftMargin=2.5*cm, rightMargin=2.5*cm, topMargin=2*cm, bottomMargin=2*cm)
    doc_ba.build(elements_ba)
    stream_ba.seek(0)
    
    elements_sp = []
    elements_sp.append(Spacer(1, 1*cm))
    elements_sp.append(Paragraph(f"Nomor\t\t: {nomor_surat}\t\t\t\t\t{lokasi_ba}, {tanggal_ba}", style_normal))
    elements_sp.append(Paragraph("Sifat\t\t: Penting", style_normal))
    elements_sp.append(Paragraph("Lampiran\t: 1 (Satu) Berkas", style_normal))
    elements_sp.append(Paragraph(f"Perihal\t\t: {perihal_surat}.", style_normal))
    elements_sp.append(Spacer(1, 0.8*cm))
    elements_sp.append(Paragraph("Yth,", style_normal))
    elements_sp.append(Paragraph("Direktur Jenderal Kawasan Permukiman", style_normal))
    elements_sp.append(Paragraph("di-", style_normal))
    elements_sp.append(Paragraph("          Jakarta", style_normal))
    elements_sp.append(Spacer(1, 0.5*cm))
    elements_sp.append(Paragraph(
        f"Menindaklanjuti Surat Direktur Jenderal Kawasan Permukiman Nomor {no_surat_dirjen} "
        f"Tanggal {tgl_surat_dirjen} {hal_surat_dirjen}.", style_body))
    elements_sp.append(Paragraph(
        f"Bersama surat ini kami sampaikan Hasil Verifikasi Calon Penerima Bantuan (CPB) "
        f"Kegiatan BSPS {stage_name} Tahun Anggaran 2026 Provinsi Sulawesi Selatan.", style_body))
    elements_sp.append(Spacer(1, 0.5*cm))
    elements_sp.append(Paragraph("Rincian usulan per Kabupaten/Kota:", style_normal))
    elements_sp.append(Spacer(1, 0.2*cm))
    
    table_data = [['No.', 'Kabupaten/Kota', 'Lolos', 'Tidak Lolos', 'Total']]
    for idx, kab in enumerate(kab_details, 1):
        kab_name = (kab['kabupaten_kota'] or "LAINNYA").upper().strip()
        table_data.append([str(idx), kab_name, str(kab['lolos']), str(kab['tidak_lolos']), str(kab['total'])])
    table_data.append(['', 'TOTAL', '', '', str(total_verifikasi_aktif)])
    
    t = Table(table_data, colWidths=[1*cm, 6*cm, 2.5*cm, 2.5*cm, 2.5*cm])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.9, 0.9, 0.9)),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('FONTNAME', (0, 0), (-1, 0), 'Times-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('FONTNAME', (0, 1), (-1, -1), 'Times-Roman'),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('ALIGN', (2, 1), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('FONTNAME', (0, -1), (-1, -1), 'Times-Bold'),
        ('BACKGROUND', (0, -1), (-1, -1), colors.Color(0.95, 0.95, 0.95)),
    ]))
    elements_sp.append(t)
    elements_sp.append(Spacer(1, 0.5*cm))
    elements_sp.append(Paragraph(
        f"Sebanyak {total_verifikasi_aktif} unit dari total {total_invers} unit telah diverifikasi, "
        f"sehingga yang masih dalam proses verifikasi sebanyak {sisa_belum_verif} unit.", style_body))
    elements_sp.append(Paragraph(
        "Demikian surat ini disampaikan, atas perhatian dan perkenaan Bapak, kami ucapkan terima kasih.", style_body))
    elements_sp.append(Spacer(1, 1.5*cm))
    elements_sp.append(Paragraph("Kepala Balai,", style_footer))
    elements_sp.append(Spacer(1, 2*cm))
    elements_sp.append(Paragraph("_________________________", style_footer))
    elements_sp.append(Paragraph("Bakhtiar", style_small))
    elements_sp.append(Paragraph("NIP. 19711009 200212 1 003", style_small))
    elements_sp.append(Spacer(1, 0.8*cm))
    elements_sp.append(Paragraph("Tembusan:", style_small))
    elements_sp.append(Paragraph("1. Direktur Jenderal Tata Kelola dan Pengendalian Resiko;", style_small))
    elements_sp.append(Paragraph("2. Kepala Pusat Data dan Informasi, Sekretariat Jenderal Kementerian PKP.", style_small))
    
    stream_sp = io.BytesIO()
    doc_sp = SimpleDocTemplate(stream_sp, pagesize=A4, leftMargin=2.5*cm, rightMargin=2.5*cm, topMargin=2*cm, bottomMargin=2*cm)
    doc_sp.build(elements_sp)
    stream_sp.seek(0)
    
    name_suffix = (batch_name or stage_name).replace(' ', '_')
    stage_upper = stage_name.upper()
    
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr(f"{stage_upper}_BAHV_{batch_name or stage_name}.pdf", stream_ba.getvalue())
        z.writestr(f"{stage_upper}_PENYAMPAIAN_BAHV_{batch_name or stage_name}.pdf", stream_sp.getvalue())
    zip_buffer.seek(0)
    
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename=DRAFT_SURAT_BA_{name_suffix}.zip"}
    )

@app.post("/api/verified/record/{record_id}/delete")
def delete_verified_record(record_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        # Check if record exists
        cursor.execute("SELECT id, no_ktp FROM verified_records WHERE id = ?", (record_id,))
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Data verifikasi tidak ditemukan")
        
        nik = row['no_ktp']
        # Delete replacement events if any
        cursor.execute("DELETE FROM replacement_events WHERE disqualified_record_id = ?", (record_id,))
        # Delete verified record
        cursor.execute("DELETE FROM verified_records WHERE id = ?", (record_id,))
        
        # Also clean up reconciliation override if any associated with this NIK
        cursor.execute("DELETE FROM reconciliation_overrides WHERE original_no_ktp = ?", (nik,))
        
        conn.commit()
    except Exception as e:
        conn.rollback()
        conn.close()
        raise HTTPException(status_code=500, detail=f"Gagal menghapus data: {str(e)}")
    conn.close()
    return {"message": "Data verifikasi berhasil dihapus"}

# --- BULK OPERATIONS ---

def replace_docx_placeholder(paragraph, search_text, replace_text):
    while search_text in paragraph.text:
        # 1. Direct replacement if in single run
        replaced = False
        for run in paragraph.runs:
            if search_text in run.text:
                run.text = run.text.replace(search_text, replace_text)
                replaced = True
                break
        if replaced:
            continue
            
        # 2. Multi-run spanning replacement (preserves shapes, drawings, and other runs)
        full_text = ''.join(r.text for r in paragraph.runs)
        start_idx = full_text.find(search_text)
        if start_idx == -1:
            break
        end_idx = start_idx + len(search_text)
        
        cur_pos = 0
        first_match = True
        for run in paragraph.runs:
            run_len = len(run.text)
            run_start = cur_pos
            run_end = cur_pos + run_len
            
            if run_end <= start_idx or run_start >= end_idx:
                pass
            elif run_start <= start_idx and run_end >= end_idx:
                p1 = run.text[:start_idx - run_start]
                p2 = run.text[end_idx - run_start:]
                run.text = p1 + replace_text + p2
            elif run_start <= start_idx and run_end < end_idx:
                p1 = run.text[:start_idx - run_start]
                run.text = p1 + (replace_text if first_match else '')
                first_match = False
            elif run_start > start_idx and run_end <= end_idx:
                run.text = ''
            elif run_start > start_idx and run_end > end_idx:
                p2 = run.text[end_idx - run_start:]
                run.text = p2

            cur_pos += run_len

@app.post("/api/export/verfal/docx")
async def export_verfal_docx(
    batch_id: int = Form(...),
    nomor_ba_verfal: str = Form(""),
    tahun_anggaran: str = Form("2026"),
    nomor_ba_versul: str = Form(""),
    tanggal_ba_verfal: str = Form(""),
    total_alokasi_versul: str = Form(""),
    total_alokasi_invers: str = Form(""),
    nama_pejabat_ketua_tim: str = Form(""),
    nama_pejabat_kepala_balai: str = Form(""),
    tanggal_terbit_ba_verfal: str = Form(""),
    alasan_tidak_lolos_terbanyak: str = Form("")
):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT vb.*, s.name as stage_name, p.name as province_name
        FROM verified_batches vb
        JOIN invers_stages s ON vb.stage_id = s.id
        LEFT JOIN provinces p ON s.province_id = p.id
        WHERE vb.id = ?
    """, (batch_id,))
    batch = cursor.fetchone()
    if not batch:
        conn.close()
        raise HTTPException(status_code=404, detail="Batch Verfal tidak ditemukan")
        
    stage_id = batch['stage_id']
    stage_name = batch['stage_name']
    prov_name_raw = batch['province_name'] or "SULAWESI SELATAN"
    prov_name = clean_province_for_export(prov_name_raw)
    kab_name = (batch['kabupaten'] or "").upper().strip()
    batch_name = batch['name']
    
    cursor.execute("""
        SELECT vr.*, re.nama_pengganti, re.jenis_kelamin_pengganti, re.no_ktp_pengganti,
               re.no_kk_pengganti, re.alamat_pengganti, re.desa_kelurahan_pengganti,
               re.kecamatan_pengganti, re.kabupaten_pengganti
        FROM verified_records vr
        LEFT JOIN replacement_events re ON re.disqualified_record_id = vr.id
        WHERE vr.batch_id = ?
        ORDER BY vr.no_urut ASC, vr.id ASC
    """, (batch_id,))
    records = [dict(r) for r in cursor.fetchall()]
    conn.close()
    
    lolos_records = [r for r in records if r['status'] == 'LOLOS']
    tidak_lolos_records = [r for r in records if r['status'] == 'TIDAK LOLOS']
    
    count_lolos = len(lolos_records)
    count_tidak_lolos = len(tidak_lolos_records)
    
    tgl_val = ""
    bln_val = ""
    thn_val = ""
    months_id = ["Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
    if "-" in tanggal_ba_verfal:
        parts = tanggal_ba_verfal.split("-")
        if len(parts) == 3:
            thn_val = parts[0]
            try:
                m_idx = int(parts[1]) - 1
                bln_val = months_id[m_idx] if 0 <= m_idx < 12 else parts[1]
            except Exception:
                bln_val = parts[1]
            tgl_val = str(int(parts[2])) if parts[2].isdigit() else parts[2]
    elif " " in tanggal_ba_verfal:
        parts = tanggal_ba_verfal.split(" ")
        if len(parts) >= 3:
            tgl_val = parts[0]
            bln_val = parts[1]
            thn_val = parts[2]
    else:
        tgl_val = tanggal_ba_verfal
        bln_val = ""
        
    if not alasan_tidak_lolos_terbanyak and tidak_lolos_records:
        reasons = [r['alasan_tidak_lolos'] for r in tidak_lolos_records if r['alasan_tidak_lolos']]
        if reasons:
            most_common = Counter(reasons).most_common(1)
            alasan_tidak_lolos_terbanyak = most_common[0][0]
    if not alasan_tidak_lolos_terbanyak:
        alasan_tidak_lolos_terbanyak = "-"
    
    template_candidates = [
        os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "FORMAT BA VERFAL.docx"),
        os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "BERITA ACARA.docx"),
        os.path.join(BASE_DIR, "FORMAT BERITA ACARA VERFAL.docx"),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates", "FORMAT BERITA ACARA VERFAL.docx"),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates", "FORMAT BA VERFAL.docx"),
        os.path.join(os.getcwd(), "FORMAT BA VERFAL.docx"),
        os.path.join(os.getcwd(), "BERITA ACARA.docx")
    ]
    template_path = None
    for candidate in template_candidates:
        if os.path.exists(candidate):
            template_path = candidate
            break
            
    if not template_path:
        raise HTTPException(status_code=404, detail="Template FORMAT BA VERFAL.docx / BERITA ACARA.docx tidak ditemukan")
        
    doc = docx.Document(template_path)
    
    replacements = {
        "[Nomor BA Verfal]": nomor_ba_verfal,
        "[Tahun Anggaran]": tahun_anggaran,
        "[Nomor BA Versul]": nomor_ba_versul,
        "[Tanggal BA Verfal]": tgl_val,
        "[Bulan BA Verfal]": bln_val,
        "[Tahun BA Verfal]": thn_val,
        "[Jumlah usulan]": str(total_alokasi_invers),
        "[Total Alokasi Invers]": str(total_alokasi_invers),
        "[Kabupaten Aktif]": kab_name,
        "[Provinsi Aktif]": prov_name,
        "[Total Alokasi Versul]": str(total_alokasi_versul or total_alokasi_invers),
        "[Jumlah lolos]": str(count_lolos),
        "[Jumlah Tidak lolos]": str(count_tidak_lolos),
        "[Alasan Tidak lolos terbanyak]": alasan_tidak_lolos_terbanyak,
        "[Nama Pejabat Kepala Balai]": nama_pejabat_kepala_balai or "( ................................................. )",
        "[Tanggal Terbit BA Verfal]": tanggal_terbit_ba_verfal or f"{kab_name.title()}, {tgl_val} {bln_val} {thn_val}",
        "[Nama Pejabat Ketua Tim]": nama_pejabat_ketua_tim or "( ................................................. )"
    }
    
    for p in doc.paragraphs:
        for k, v in replacements.items():
            replace_docx_placeholder(p, k, str(v))

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in replacements.items():
                        replace_docx_placeholder(p, k, str(v))
                        
    for s in doc.sections:
        for hp in s.header.paragraphs:
            for k, v in replacements.items():
                replace_docx_placeholder(hp, k, str(v))

    stream = io.BytesIO()
    doc.save(stream)
    docx_bytes = stream.getvalue()
    stream.seek(0)
    
    clean_kab = kab_name.replace(' ', '_')
    clean_batch = batch_name.replace(' ', '_')
    filename = f"BA_VERFAL_{clean_kab}_{clean_batch}.docx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.post("/api/export/verfal/pdf")
async def export_verfal_pdf(
    batch_id: int = Form(...),
    nomor_ba_verfal: str = Form(""),
    tahun_anggaran: str = Form("2026"),
    nomor_ba_versul: str = Form(""),
    tanggal_ba_verfal: str = Form(""),
    total_alokasi_versul: str = Form(""),
    total_alokasi_invers: str = Form(""),
    nama_pejabat_ketua_tim: str = Form(""),
    nama_pejabat_kepala_balai: str = Form(""),
    tanggal_terbit_ba_verfal: str = Form(""),
    alasan_tidak_lolos_terbanyak: str = Form("")
):
    docx_resp = await export_verfal_docx(
        batch_id, nomor_ba_verfal, tahun_anggaran, nomor_ba_versul,
        tanggal_ba_verfal, total_alokasi_versul, total_alokasi_invers,
        nama_pejabat_ketua_tim, nama_pejabat_kepala_balai,
        tanggal_terbit_ba_verfal, alasan_tidak_lolos_terbanyak
    )
    
    docx_bytes = b""
    if hasattr(docx_resp, 'body') and docx_resp.body:
        docx_bytes = docx_resp.body
    elif hasattr(docx_resp, 'body_iterator'):
        if isinstance(docx_resp.body_iterator, io.BytesIO):
            docx_bytes = docx_resp.body_iterator.getvalue()
        elif isinstance(docx_resp.body_iterator, bytes):
            docx_bytes = docx_resp.body_iterator
        else:
            chunks = []
            async for chunk in docx_resp.body_iterator:
                chunks.append(chunk)
            docx_bytes = b"".join(chunks)
            
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
        tmp_docx.write(docx_bytes)
        tmp_docx_path = tmp_docx.name
        
    tmp_pdf_path = tmp_docx_path.replace(".docx", ".pdf")
    try:
        import subprocess
        converted = False
        try:
            from docx2pdf import convert
            convert(tmp_docx_path, tmp_pdf_path)
            converted = os.path.exists(tmp_pdf_path)
        except Exception:
            pass
            
        if not converted:
            try:
                subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(tmp_docx_path), tmp_docx_path], check=True, timeout=30)
                converted = os.path.exists(tmp_pdf_path)
            except Exception:
                pass
                
        if converted and os.path.exists(tmp_pdf_path):
            with open(tmp_pdf_path, "rb") as f:
                pdf_bytes = f.read()
            os.remove(tmp_docx_path)
            os.remove(tmp_pdf_path)
            return StreamingResponse(
                io.BytesIO(pdf_bytes),
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename=BA_VERFAL_{batch_id}.pdf"}
            )
        else:
            if os.path.exists(tmp_docx_path):
                os.remove(tmp_docx_path)
            return StreamingResponse(
                io.BytesIO(docx_bytes),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename=BA_VERFAL_{batch_id}.docx"}
            )
    except Exception as e:
        if os.path.exists(tmp_docx_path):
            os.remove(tmp_docx_path)
        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=BA_VERFAL_{batch_id}.docx"}
        )

@app.get("/api/export/verfal/excel/{batch_id}")
def export_verfal_excel(batch_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT vb.*, s.name as stage_name, p.name as province_name
        FROM verified_batches vb
        JOIN invers_stages s ON vb.stage_id = s.id
        LEFT JOIN provinces p ON s.province_id = p.id
        WHERE vb.id = ?
    """, (batch_id,))
    batch = cursor.fetchone()
    if not batch:
        conn.close()
        raise HTTPException(status_code=404, detail="Batch Verfal tidak ditemukan")
        
    stage_id = batch['stage_id']
    stage_name = batch['stage_name']
    prov_name_raw = batch['province_name'] or "SULAWESI SELATAN"
    prov_name = clean_province_for_export(prov_name_raw)
    kab_name = (batch['kabupaten'] or "").upper().strip()
    batch_name = batch['name']
    
    import json
    meta = {}
    if batch['metadata_json']:
        try: meta = json.loads(batch['metadata_json'])
        except Exception: pass
    tahun_anggaran = meta.get("tahun_anggaran", "2026")
    nama_pejabat_ketua_tim = meta.get("nama_pejabat_ketua_tim", "")
    nama_pejabat_kepala_balai = meta.get("nama_pejabat_kepala_balai", "")
    tanggal_terbit_ba_verfal = meta.get("tanggal_terbit_ba_verfal", "")
    tanggal_ba = meta.get("tanggal_ba_verfal") or batch.get("tanggal_ba") or ""
    
    tgl_val = ""
    bln_val = ""
    thn_val = ""
    months_id = ["Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
    if tanggal_ba and "-" in str(tanggal_ba):
        parts = str(tanggal_ba).split("-")
        if len(parts) == 3:
            thn_val = parts[0]
            try:
                m_idx = int(parts[1]) - 1
                bln_val = months_id[m_idx] if 0 <= m_idx < 12 else parts[1]
            except Exception:
                bln_val = parts[1]
            tgl_val = str(int(parts[2])) if parts[2].isdigit() else parts[2]
            
    tanggal_terbit_display = tanggal_terbit_ba_verfal or (f"{kab_name.title()}, {tgl_val} {bln_val} {thn_val}".strip(", ") if (tgl_val or bln_val) else f"{kab_name.title()}, [Tanggal Terbit Verifikasi]")
    
    cursor.execute("""
        SELECT vr.*, re.nama_pengganti, re.jenis_kelamin_pengganti, re.no_ktp_pengganti,
               re.no_kk_pengganti, re.alamat_pengganti, re.desa_kelurahan_pengganti,
               re.kecamatan_pengganti, re.kabupaten_pengganti
        FROM verified_records vr
        LEFT JOIN replacement_events re ON re.disqualified_record_id = vr.id
        WHERE vr.batch_id = ?
        ORDER BY vr.no_urut ASC, vr.id ASC
    """, (batch_id,))
    records = [dict(r) for r in cursor.fetchall()]
    conn.close()
    
    lolos_records = [r for r in records if r['status'] == 'LOLOS']
    tidak_lolos_records = [r for r in records if r['status'] == 'TIDAK LOLOS']
    
    template_candidates = [
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates", "EXPORT_VERFAL.xlsx"),
        os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "EXPORT_VERFAL.xlsx"),
        os.path.join(BASE_DIR, "EXPORT_VERFAL.xlsx"),
        os.path.join(os.getcwd(), "backend", "templates", "EXPORT_VERFAL.xlsx"),
        os.path.join(os.getcwd(), "EXPORT_VERFAL.xlsx")
    ]
    template_path = None
    for candidate in template_candidates:
        if os.path.exists(candidate):
            template_path = candidate
            break
            
    if template_path:
        wb = openpyxl.load_workbook(template_path)
    else:
        wb = openpyxl.Workbook()
        
    font_data = Font(name='Bookman Old Style', size=10)
    font_sig = Font(name='Bookman Old Style', size=14)
    font_sig_bold = Font(name='Bookman Old Style', size=14, bold=True)
    align_sig_center = Alignment(horizontal='center', vertical='center')
    
    border_thin = Border(
        left=Side(style='thin', color='000000'),
        right=Side(style='thin', color='000000'),
        top=Side(style='thin', color='000000'),
        bottom=Side(style='thin', color='000000')
    )
    align_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    align_left = Alignment(horizontal='left', vertical='center')

    # 1. Sheet Lamp.IIA (Lolos & Pengganti Terurut per Desa)
    ws_iia = wb["Lamp.IIA"] if "Lamp.IIA" in wb.sheetnames else wb.active
    ws_iia.title = "Lamp.IIA"
    
    # Page setup: Landscape, 1 page wide, automatic height, 0.5 margins
    ws_iia.page_setup.orientation = ws_iia.ORIENTATION_LANDSCAPE
    ws_iia.page_setup.paperSize = ws_iia.PAPERSIZE_A4
    if ws_iia.sheet_properties and ws_iia.sheet_properties.pageSetUpPr:
        ws_iia.sheet_properties.pageSetUpPr.fitToPage = True
    ws_iia.page_setup.fitToWidth = 1
    ws_iia.page_setup.fitToHeight = 0
    ws_iia.page_margins.left = 0.5
    ws_iia.page_margins.right = 0.5
    ws_iia.page_margins.top = 0.5
    ws_iia.page_margins.bottom = 0.5

    ws_iia['A2'] = "DAFTAR HASIL VERIFIKASI FAKTUAL CALON PENERIMA BANTUAN"
    ws_iia['A3'] = f"BEDAH RUMAH TAHUN {tahun_anggaran}"
    ws_iia['A4'] = f"PROVINSI {prov_name}"
    
    # Clear existing data rows if any
    if ws_iia.max_row >= 8:
        ws_iia.delete_rows(8, ws_iia.max_row - 7)
        
    # Kumpulkan seluruh baris Lamp.IIA (Lolos + Pengganti dari Tidak Lolos)
    iia_items = []
    for rec in lolos_records:
        iia_items.append({
            'nama': rec['nama'],
            'jenis_kelamin': rec['jenis_kelamin'],
            'no_kk': rec['no_kk'],
            'no_ktp': rec['no_ktp'],
            'alamat': rec['alamat'],
            'desa_kelurahan': rec['desa_kelurahan'] or '',
            'kecamatan': rec['kecamatan'] or '',
            'kabupaten_kota': rec['kabupaten_kota'] or kab_name,
            'status_label': 'LOLOS',
            'tahap': rec['tahap'] or stage_name,
            'tanggal': rec['tanggal'] or batch.get('tanggal_ba') or '',
            'keterangan': rec['keterangan'] or ''
        })
        
    for rec in tidak_lolos_records:
        nama_p = (rec.get('nama_pengganti') or '').strip()
        if nama_p and nama_p != 'NONE':
            iia_items.append({
                'nama': nama_p,
                'jenis_kelamin': rec.get('jenis_kelamin_pengganti') or '',
                'no_kk': rec.get('no_kk_pengganti') or '',
                'no_ktp': rec.get('no_ktp_pengganti') or '',
                'alamat': rec.get('alamat_pengganti') or '',
                'desa_kelurahan': rec.get('desa_kelurahan_pengganti') or rec.get('desa_kelurahan') or '',
                'kecamatan': rec.get('kecamatan_pengganti') or rec.get('kecamatan') or '',
                'kabupaten_kota': rec.get('kabupaten_pengganti') or rec.get('kabupaten_kota') or kab_name,
                'status_label': 'PENGGANTI',
                'tahap': rec['tahap'] or stage_name,
                'tanggal': rec['tanggal'] or batch.get('tanggal_ba') or '',
                'keterangan': f"Pengganti dari {rec['nama']}" if rec.get('nama') else (rec.get('keterangan') or '')
            })
            
    # Urutkan berdasarkan Kecamatan, Desa/Kelurahan (agar berkumpul per desa), status ('LOLOS' lebih dulu baru 'PENGGANTI'), dan Nama
    def iia_sort_key(item):
        kec = (item['kecamatan'] or '').strip().upper()
        desa = (item['desa_kelurahan'] or '').strip().upper()
        st = 0 if item['status_label'] == 'LOLOS' else 1
        nama = (item['nama'] or '').strip().upper()
        return (kec, desa, st, nama)
        
    iia_items.sort(key=iia_sort_key)
    
    r_idx = 8
    for idx, item in enumerate(iia_items, 1):
        ws_iia.row_dimensions[r_idx].height = 28.05
        vals = [
            idx,
            item['nama'],
            item['jenis_kelamin'],
            item['no_kk'],
            item['no_ktp'],
            item['alamat'],
            item['desa_kelurahan'],
            item['kecamatan'],
            item['kabupaten_kota'],
            item['status_label'],
            item['tahap'],
            item['tanggal'],
            item['keterangan']
        ]
        for c_idx, v in enumerate(vals, 1):
            cell = ws_iia.cell(row=r_idx, column=c_idx, value=v)
            cell.font = font_data
            cell.border = border_thin
            if c_idx in (1, 3, 4, 5, 10, 11, 12):
                cell.alignment = align_center
            else:
                cell.alignment = align_left
        r_idx += 1
        
    # Tanda Tangan Lamp.IIA (Pic 1)
    row_sig_iia = r_idx + 2
    ws_iia.cell(row=row_sig_iia, column=11, value=tanggal_terbit_display).font = font_sig
    ws_iia.cell(row=row_sig_iia, column=11).alignment = align_sig_center
    
    ws_iia.cell(row=row_sig_iia + 1, column=11, value="Pelaksana/Ketua Tim Verifikasi Faktual,").font = font_sig
    ws_iia.cell(row=row_sig_iia + 1, column=11).alignment = align_sig_center
    
    ws_iia.cell(row=row_sig_iia + 5, column=11, value=nama_pejabat_ketua_tim or "( ................................................. )").font = font_sig
    ws_iia.cell(row=row_sig_iia + 5, column=11).alignment = align_sig_center

    # 2. Sheet Lamp.IIIA (Daftar Tidak Lolos & Pengganti Lengkap)
    ws_iiia = wb["Lamp.IIIA"] if "Lamp.IIIA" in wb.sheetnames else wb.create_sheet("Lamp.IIIA")
    ws_iiia.title = "Lamp.IIIA"
    
    # Page setup: Landscape, 1 page wide, automatic height, 0.5 margins
    ws_iiia.page_setup.orientation = ws_iiia.ORIENTATION_LANDSCAPE
    ws_iiia.page_setup.paperSize = ws_iiia.PAPERSIZE_A4
    if ws_iiia.sheet_properties and ws_iiia.sheet_properties.pageSetUpPr:
        ws_iiia.sheet_properties.pageSetUpPr.fitToPage = True
    ws_iiia.page_setup.fitToWidth = 1
    ws_iiia.page_setup.fitToHeight = 0
    ws_iiia.page_margins.left = 0.5
    ws_iiia.page_margins.right = 0.5
    ws_iiia.page_margins.top = 0.5
    ws_iiia.page_margins.bottom = 0.5

    ws_iiia['A2'] = "DAFTAR CALON PENGGANTI CALON PENERIMA BANTUAN"
    ws_iiia['A3'] = f"BEDAH RUMAH TAHUN {tahun_anggaran}"
    ws_iiia['A4'] = f"PROVINSI {prov_name}"

    # Determine columns layout from template header row 7
    has_22_cols = (ws_iiia.cell(row=7, column=5).value and 'KK' in str(ws_iiia.cell(row=7, column=5).value).upper())
    is_21_cols = (ws_iiia.cell(row=7, column=14).value and 'NO.KK' in str(ws_iiia.cell(row=7, column=14).value).upper())
    
    if ws_iiia.max_row >= 8:
        ws_iiia.delete_rows(8, ws_iiia.max_row - 7)
            
    r_idx = 8
    for idx, rec in enumerate(tidak_lolos_records, 1):
        ws_iiia.row_dimensions[r_idx].height = 28.05
        if has_22_cols:
            vals = [
                idx,
                rec['nama'],
                rec['jenis_kelamin'],
                rec['no_ktp'],
                rec['no_kk'] or "",
                rec['alamat'],
                rec['desa_kelurahan'],
                rec['kecamatan'],
                rec['kabupaten_kota'] or kab_name,
                rec['alasan_tidak_lolos'],
                "", # BNBA
                rec.get('nama_pengganti') or "",
                rec.get('jenis_kelamin_pengganti') or "",
                rec.get('no_ktp_pengganti') or "",
                rec.get('no_kk_pengganti') or "",
                rec.get('alamat_pengganti') or "",
                rec.get('desa_kelurahan_pengganti') or "",
                rec.get('kecamatan_pengganti') or "",
                rec.get('kabupaten_pengganti') or kab_name,
                rec['tahap'] or stage_name,
                rec['tanggal'] or batch.get('tanggal_ba') or "",
                rec['keterangan'] or ""
            ]
            center_cols = {1, 3, 4, 5, 10, 11, 13, 14, 15, 20, 21}
        elif is_21_cols:
            vals = [
                idx,
                rec['nama'],
                rec['jenis_kelamin'],
                rec['no_ktp'],
                rec['alamat'],
                rec['desa_kelurahan'],
                rec['kecamatan'],
                rec['kabupaten_kota'] or kab_name,
                rec['alasan_tidak_lolos'],
                "", # BNBA
                rec.get('nama_pengganti') or "",
                rec.get('jenis_kelamin_pengganti') or "",
                rec.get('no_ktp_pengganti') or "",
                rec.get('no_kk_pengganti') or "",
                rec.get('alamat_pengganti') or "",
                rec.get('desa_kelurahan_pengganti') or "",
                rec.get('kecamatan_pengganti') or "",
                rec.get('kabupaten_pengganti') or kab_name,
                rec['tahap'] or stage_name,
                rec['tanggal'] or batch.get('tanggal_ba') or "",
                rec['keterangan'] or ""
            ]
            center_cols = {1, 3, 4, 9, 10, 12, 13, 14, 19, 20}
        else:
            vals = [
                idx,
                rec['nama'],
                rec['jenis_kelamin'],
                rec['no_ktp'],
                rec['alamat'],
                rec['desa_kelurahan'],
                rec['kabupaten_kota'] or kab_name,
                rec['alasan_tidak_lolos'],
                "", # BNBA
                rec.get('nama_pengganti') or "",
                rec.get('jenis_kelamin_pengganti') or "",
                rec.get('no_ktp_pengganti') or "",
                rec.get('alamat_pengganti') or "",
                rec.get('desa_kelurahan_pengganti') or "",
                rec.get('kabupaten_pengganti') or kab_name,
                rec['tahap'] or stage_name,
                rec['tanggal'] or batch.get('tanggal_ba') or "",
                rec['keterangan'] or ""
            ]
            center_cols = {1, 3, 4, 8, 9, 11, 12, 16, 17}
            
        for c_idx, v in enumerate(vals, 1):
            cell = ws_iiia.cell(row=r_idx, column=c_idx, value=v)
            cell.font = font_data
            cell.border = border_thin
            if c_idx in center_cols:
                cell.alignment = align_center
            else:
                cell.alignment = align_left
        r_idx += 1
        
    # Tanda Tangan & Catatan Lamp.IIIA (Pic 2)
    row_sig_iiia = r_idx + 2
    col_left = 5
    col_right = 18 if (has_22_cols or is_21_cols) else 15
    
    # Left: Mengetahui - Kepala Balai
    ws_iiia.cell(row=row_sig_iiia, column=col_left, value="Mengetahui,").font = font_sig
    ws_iiia.cell(row=row_sig_iiia, column=col_left).alignment = align_sig_center
    
    ws_iiia.cell(row=row_sig_iiia + 1, column=col_left, value="Kepala Balai Pelaksana ").font = font_sig
    ws_iiia.cell(row=row_sig_iiia + 1, column=col_left).alignment = align_sig_center
    
    ws_iiia.cell(row=row_sig_iiia + 2, column=col_left, value="Penyediaan Perumahan dan ").font = font_sig
    ws_iiia.cell(row=row_sig_iiia + 2, column=col_left).alignment = align_sig_center
    
    ws_iiia.cell(row=row_sig_iiia + 3, column=col_left, value="Kawasan Permukiman Sulawesi III,").font = font_sig
    ws_iiia.cell(row=row_sig_iiia + 3, column=col_left).alignment = align_sig_center
    
    ws_iiia.cell(row=row_sig_iiia + 7, column=col_left, value=nama_pejabat_kepala_balai or "( ................................................. )").font = font_sig
    ws_iiia.cell(row=row_sig_iiia + 7, column=col_left).alignment = align_sig_center

    # Right: Pelaksana/Ketua Tim Verifikasi Faktual
    ws_iiia.cell(row=row_sig_iiia, column=col_right, value=tanggal_terbit_display).font = font_sig
    ws_iiia.cell(row=row_sig_iiia, column=col_right).alignment = align_sig_center
    
    ws_iiia.cell(row=row_sig_iiia + 1, column=col_right, value="Pelaksana/Ketua Tim ").font = font_sig
    ws_iiia.cell(row=row_sig_iiia + 1, column=col_right).alignment = align_sig_center
    
    ws_iiia.cell(row=row_sig_iiia + 2, column=col_right, value="Verifikasi Faktual,").font = font_sig
    ws_iiia.cell(row=row_sig_iiia + 2, column=col_right).alignment = align_sig_center
    
    ws_iiia.cell(row=row_sig_iiia + 7, column=col_right, value=nama_pejabat_ketua_tim or "( ................................................. )").font = font_sig
    ws_iiia.cell(row=row_sig_iiia + 7, column=col_right).alignment = align_sig_center

    # Notes (Catatan)
    row_notes = row_sig_iiia + 13
    notes_list = [
        ("Catatan:", True),
        ("*) Alasan Tidak Lolos, diisi dengan angka (1-8) sebagai berikut:", False),
        ("1. Belum memiliki KK sendiri;", False),
        ("2. Tanah bersengketa;", False),
        ("3. Rumah dalam kondisi layak;", False),
        ("4. Memiliki rumah lebih dari 1;", False),
        ("5. Pernah memperoleh bantuan dari APBN/APBD/CSR/anggaran lainnya;", False),
        ("6. Penghasilan lebih dari UMP;", False),
        ("7. Memilih untuk dibantu dengan sumber anggaran lain;", False),
        ("8. Menghuni kurang dari 3 tahun;", False),
        ("9. Lainnya (diisi pada kolom keterangan);", False)
    ]
    for n_idx, (ntxt, is_b) in enumerate(notes_list):
        cell_n = ws_iiia.cell(row=row_notes + n_idx, column=1, value=ntxt)
        cell_n.font = Font(name='Bookman Old Style', size=14, bold=is_b)
        
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    
    clean_kab = kab_name.replace(' ', '_')
    clean_batch = batch_name.replace(' ', '_')
    filename = f"EXPORT_VERFAL_{clean_kab}_{clean_batch}.xlsx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.get("/api/export/verfal/stage-excel/{stage_id}")
def export_stage_verfal_excel(stage_id: int):
    """
    Ekspor berkas Excel gabungan untuk seluruh kabupaten yang memiliki data verifikasi faktual pada tahap aktif.
    """
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT s.name as stage_name, p.name as province_name
        FROM invers_stages s
        LEFT JOIN provinces p ON s.province_id = p.id
        WHERE s.id = ?
    """, (stage_id,))
    stage_row = cursor.fetchone()
    if not stage_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Tahap tidak ditemukan")
        
    stage_name = stage_row['stage_name']
    prov_name_raw = stage_row['province_name'] or "SULAWESI SELATAN"
    prov_name = clean_province_for_export(prov_name_raw)
    
    cursor.execute("""
        SELECT vr.*, re.nama_pengganti, re.jenis_kelamin_pengganti, re.no_ktp_pengganti,
               re.no_kk_pengganti, re.alamat_pengganti, re.desa_kelurahan_pengganti,
               re.kecamatan_pengganti, re.kabupaten_pengganti, vb.metadata_json, vb.tanggal_ba,
               vb.name as batch_name
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        LEFT JOIN replacement_events re ON re.disqualified_record_id = vr.id
        WHERE vb.stage_id = ?
        ORDER BY vr.kabupaten_kota ASC, vr.kecamatan ASC, vr.desa_kelurahan ASC, vr.no_urut ASC, vr.id ASC
    """, (stage_id,))
    records = [dict(r) for r in cursor.fetchall()]
    conn.close()
    
    if not records:
        raise HTTPException(status_code=400, detail="Belum ada data verifikasi faktual pada tahap ini")
        
    lolos_records = [r for r in records if r['status'] == 'LOLOS']
    tidak_lolos_records = [r for r in records if r['status'] == 'TIDAK LOLOS']
    
    template_candidates = [
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates", "EXPORT_VERFAL.xlsx"),
        os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "EXPORT_VERFAL.xlsx"),
        os.path.join(BASE_DIR, "EXPORT_VERFAL.xlsx"),
        os.path.join(os.getcwd(), "backend", "templates", "EXPORT_VERFAL.xlsx"),
        os.path.join(os.getcwd(), "EXPORT_VERFAL.xlsx")
    ]
    template_path = None
    for candidate in template_candidates:
        if os.path.exists(candidate):
            template_path = candidate
            break
            
    if template_path:
        wb = openpyxl.load_workbook(template_path)
    else:
        wb = openpyxl.Workbook()
        
    font_data = Font(name='Bookman Old Style', size=10)
    font_sig = Font(name='Bookman Old Style', size=14)
    font_sig_bold = Font(name='Bookman Old Style', size=14, bold=True)
    align_sig_center = Alignment(horizontal='center', vertical='center')
    
    border_thin = Border(
        left=Side(style='thin', color='000000'),
        right=Side(style='thin', color='000000'),
        top=Side(style='thin', color='000000'),
        bottom=Side(style='thin', color='000000')
    )
    align_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    align_left = Alignment(horizontal='left', vertical='center')

    tahun_anggaran = "2026"
    
    # 1. Sheet Lamp.IIA (Lolos & Pengganti Terurut per Kabupaten, Kecamatan, Desa)
    ws_iia = wb["Lamp.IIA"] if "Lamp.IIA" in wb.sheetnames else wb.active
    ws_iia.title = "Lamp.IIA"
    
    ws_iia.page_setup.orientation = ws_iia.ORIENTATION_LANDSCAPE
    ws_iia.page_setup.paperSize = ws_iia.PAPERSIZE_A4
    if ws_iia.sheet_properties and ws_iia.sheet_properties.pageSetUpPr:
        ws_iia.sheet_properties.pageSetUpPr.fitToPage = True
    ws_iia.page_setup.fitToWidth = 1
    ws_iia.page_setup.fitToHeight = 0
    ws_iia.page_margins.left = 0.5
    ws_iia.page_margins.right = 0.5
    ws_iia.page_margins.top = 0.5
    ws_iia.page_margins.bottom = 0.5

    ws_iia['A2'] = "DAFTAR HASIL VERIFIKASI FAKTUAL CALON PENERIMA BANTUAN"
    ws_iia['A3'] = f"BEDAH RUMAH TAHUN {tahun_anggaran}"
    ws_iia['A4'] = f"PROVINSI {prov_name}"
    
    if ws_iia.max_row >= 8:
        ws_iia.delete_rows(8, ws_iia.max_row - 7)
        
    iia_items = []
    for rec in lolos_records:
        iia_items.append({
            'nama': rec['nama'],
            'jenis_kelamin': rec['jenis_kelamin'],
            'no_kk': rec['no_kk'],
            'no_ktp': rec['no_ktp'],
            'alamat': rec['alamat'],
            'desa_kelurahan': rec['desa_kelurahan'] or '',
            'kecamatan': rec['kecamatan'] or '',
            'kabupaten_kota': rec['kabupaten_kota'] or '',
            'status_label': 'LOLOS',
            'tahap': rec['tahap'] or stage_name,
            'tanggal': rec['tanggal'] or rec.get('tanggal_ba') or '',
            'keterangan': rec['keterangan'] or ''
        })
        
    for rec in tidak_lolos_records:
        nama_p = (rec.get('nama_pengganti') or '').strip()
        if nama_p and nama_p != 'NONE':
            iia_items.append({
                'nama': nama_p,
                'jenis_kelamin': rec.get('jenis_kelamin_pengganti') or '',
                'no_kk': rec.get('no_kk_pengganti') or '',
                'no_ktp': rec.get('no_ktp_pengganti') or '',
                'alamat': rec.get('alamat_pengganti') or '',
                'desa_kelurahan': rec.get('desa_kelurahan_pengganti') or rec.get('desa_kelurahan') or '',
                'kecamatan': rec.get('kecamatan_pengganti') or rec.get('kecamatan') or '',
                'kabupaten_kota': rec.get('kabupaten_pengganti') or rec.get('kabupaten_kota') or '',
                'status_label': 'PENGGANTI',
                'tahap': rec['tahap'] or stage_name,
                'tanggal': rec['tanggal'] or rec.get('tanggal_ba') or '',
                'keterangan': f"Pengganti dari {rec['nama']}" if rec.get('nama') else (rec.get('keterangan') or '')
            })
            
    def iia_sort_key(item):
        kab = (item['kabupaten_kota'] or '').strip().upper()
        kec = (item['kecamatan'] or '').strip().upper()
        desa = (item['desa_kelurahan'] or '').strip().upper()
        st = 0 if item['status_label'] == 'LOLOS' else 1
        nama = (item['nama'] or '').strip().upper()
        return (kab, kec, desa, st, nama)
        
    iia_items.sort(key=iia_sort_key)
    
    r_idx = 8
    for idx, item in enumerate(iia_items, 1):
        ws_iia.row_dimensions[r_idx].height = 28.05
        vals = [
            idx,
            item['nama'],
            item['jenis_kelamin'],
            item['no_kk'],
            item['no_ktp'],
            item['alamat'],
            item['desa_kelurahan'],
            item['kecamatan'],
            item['kabupaten_kota'],
            item['status_label'],
            item['tahap'],
            item['tanggal'],
            item['keterangan']
        ]
        for c_idx, v in enumerate(vals, 1):
            cell = ws_iia.cell(row=r_idx, column=c_idx, value=v)
            cell.font = font_data
            cell.border = border_thin
            if c_idx in (1, 3, 4, 5, 10, 11, 12):
                cell.alignment = align_center
            else:
                cell.alignment = align_left
        r_idx += 1
        
    # Tanda Tangan Lamp.IIA
    row_sig_iia = r_idx + 2
    ws_iia.cell(row=row_sig_iia, column=11, value="Makassar, ").font = font_sig
    ws_iia.cell(row=row_sig_iia, column=11).alignment = align_sig_center
    
    ws_iia.cell(row=row_sig_iia + 1, column=11, value="Pelaksana/Ketua Tim Verifikasi Faktual,").font = font_sig
    ws_iia.cell(row=row_sig_iia + 1, column=11).alignment = align_sig_center
    
    ws_iia.cell(row=row_sig_iia + 5, column=11, value="( ................................................. )").font = font_sig
    ws_iia.cell(row=row_sig_iia + 5, column=11).alignment = align_sig_center

    # 2. Sheet Lamp.IIIA (Daftar Tidak Lolos & Pengganti Lengkap)
    ws_iiia = wb["Lamp.IIIA"] if "Lamp.IIIA" in wb.sheetnames else wb.create_sheet("Lamp.IIIA")
    ws_iiia.title = "Lamp.IIIA"
    
    ws_iiia.page_setup.orientation = ws_iiia.ORIENTATION_LANDSCAPE
    ws_iiia.page_setup.paperSize = ws_iiia.PAPERSIZE_A4
    if ws_iiia.sheet_properties and ws_iiia.sheet_properties.pageSetUpPr:
        ws_iiia.sheet_properties.pageSetUpPr.fitToPage = True
    ws_iiia.page_setup.fitToWidth = 1
    ws_iiia.page_setup.fitToHeight = 0
    ws_iiia.page_margins.left = 0.5
    ws_iiia.page_margins.right = 0.5
    ws_iiia.page_margins.top = 0.5
    ws_iiia.page_margins.bottom = 0.5

    ws_iiia['A2'] = "DAFTAR CALON PENGGANTI CALON PENERIMA BANTUAN"
    ws_iiia['A3'] = f"BEDAH RUMAH TAHUN {tahun_anggaran}"
    ws_iiia['A4'] = f"PROVINSI {prov_name}"

    has_22_cols = (ws_iiia.cell(row=7, column=5).value and 'KK' in str(ws_iiia.cell(row=7, column=5).value).upper())
    is_21_cols = (ws_iiia.cell(row=7, column=14).value and 'NO.KK' in str(ws_iiia.cell(row=7, column=14).value).upper())
    
    if ws_iiia.max_row >= 8:
        ws_iiia.delete_rows(8, ws_iiia.max_row - 7)
            
    # Sort Tidak Lolos per Kabupaten, Kecamatan, Desa
    tidak_lolos_records.sort(key=lambda r: (
        (r.get('kabupaten_kota') or '').strip().upper(),
        (r.get('kecamatan') or '').strip().upper(),
        (r.get('desa_kelurahan') or '').strip().upper(),
        (r.get('nama') or '').strip().upper()
    ))

    r_idx = 8
    for idx, rec in enumerate(tidak_lolos_records, 1):
        ws_iiia.row_dimensions[r_idx].height = 28.05
        if has_22_cols:
            vals = [
                idx,
                rec['nama'],
                rec['jenis_kelamin'],
                rec['no_ktp'],
                rec['no_kk'] or "",
                rec['alamat'],
                rec['desa_kelurahan'],
                rec['kecamatan'],
                rec['kabupaten_kota'] or "",
                rec['alasan_tidak_lolos'],
                "", # BNBA
                rec.get('nama_pengganti') or "",
                rec.get('jenis_kelamin_pengganti') or "",
                rec.get('no_ktp_pengganti') or "",
                rec.get('no_kk_pengganti') or "",
                rec.get('alamat_pengganti') or "",
                rec.get('desa_kelurahan_pengganti') or "",
                rec.get('kecamatan_pengganti') or "",
                rec.get('kabupaten_pengganti') or rec.get('kabupaten_kota') or "",
                rec['tahap'] or stage_name,
                rec['tanggal'] or rec.get('tanggal_ba') or "",
                rec['keterangan'] or ""
            ]
            center_cols = {1, 3, 4, 5, 10, 11, 13, 14, 15, 20, 21}
        elif is_21_cols:
            vals = [
                idx,
                rec['nama'],
                rec['jenis_kelamin'],
                rec['no_ktp'],
                rec['alamat'],
                rec['desa_kelurahan'],
                rec['kecamatan'],
                rec['kabupaten_kota'] or "",
                rec['alasan_tidak_lolos'],
                "", # BNBA
                rec.get('nama_pengganti') or "",
                rec.get('jenis_kelamin_pengganti') or "",
                rec.get('no_ktp_pengganti') or "",
                rec.get('no_kk_pengganti') or "",
                rec.get('alamat_pengganti') or "",
                rec.get('desa_kelurahan_pengganti') or "",
                rec.get('kecamatan_pengganti') or "",
                rec.get('kabupaten_pengganti') or rec.get('kabupaten_kota') or "",
                rec['tahap'] or stage_name,
                rec['tanggal'] or rec.get('tanggal_ba') or "",
                rec['keterangan'] or ""
            ]
            center_cols = {1, 3, 4, 9, 10, 12, 13, 14, 19, 20}
        else:
            vals = [
                idx,
                rec['nama'],
                rec['jenis_kelamin'],
                rec['no_ktp'],
                rec['alamat'],
                rec['desa_kelurahan'],
                rec['kabupaten_kota'] or "",
                rec['alasan_tidak_lolos'],
                "", # BNBA
                rec.get('nama_pengganti') or "",
                rec.get('jenis_kelamin_pengganti') or "",
                rec.get('no_ktp_pengganti') or "",
                rec.get('alamat_pengganti') or "",
                rec.get('desa_kelurahan_pengganti') or "",
                rec.get('kabupaten_pengganti') or rec.get('kabupaten_kota') or "",
                rec['tahap'] or stage_name,
                rec['tanggal'] or rec.get('tanggal_ba') or "",
                rec['keterangan'] or ""
            ]
            center_cols = {1, 3, 4, 8, 9, 11, 12, 16, 17}
            
        for c_idx, v in enumerate(vals, 1):
            cell = ws_iiia.cell(row=r_idx, column=c_idx, value=v)
            cell.font = font_data
            cell.border = border_thin
            if c_idx in center_cols:
                cell.alignment = align_center
            else:
                cell.alignment = align_left
        r_idx += 1
        
    # Tanda Tangan & Catatan Lamp.IIIA
    row_sig_iiia = r_idx + 2
    col_left = 5
    col_right = 18 if (has_22_cols or is_21_cols) else 15
    
    # Left: Mengetahui - Kepala Balai
    ws_iiia.cell(row=row_sig_iiia, column=col_left, value="Mengetahui,").font = font_sig
    ws_iiia.cell(row=row_sig_iiia, column=col_left).alignment = align_sig_center
    
    ws_iiia.cell(row=row_sig_iiia + 1, column=col_left, value="Kepala Balai Pelaksana ").font = font_sig
    ws_iiia.cell(row=row_sig_iiia + 1, column=col_left).alignment = align_sig_center
    
    ws_iiia.cell(row=row_sig_iiia + 2, column=col_left, value="Penyediaan Perumahan dan ").font = font_sig
    ws_iiia.cell(row=row_sig_iiia + 2, column=col_left).alignment = align_sig_center
    
    ws_iiia.cell(row=row_sig_iiia + 3, column=col_left, value="Kawasan Permukiman Sulawesi III,").font = font_sig
    ws_iiia.cell(row=row_sig_iiia + 3, column=col_left).alignment = align_sig_center
    
    ws_iiia.cell(row=row_sig_iiia + 7, column=col_left, value="( ................................................. )").font = font_sig
    ws_iiia.cell(row=row_sig_iiia + 7, column=col_left).alignment = align_sig_center

    # Right: Pelaksana/Ketua Tim Verifikasi Faktual
    ws_iiia.cell(row=row_sig_iiia, column=col_right, value="Makassar, ").font = font_sig
    ws_iiia.cell(row=row_sig_iiia, column=col_right).alignment = align_sig_center
    
    ws_iiia.cell(row=row_sig_iiia + 1, column=col_right, value="Pelaksana/Ketua Tim ").font = font_sig
    ws_iiia.cell(row=row_sig_iiia + 1, column=col_right).alignment = align_sig_center
    
    ws_iiia.cell(row=row_sig_iiia + 2, column=col_right, value="Verifikasi Faktual,").font = font_sig
    ws_iiia.cell(row=row_sig_iiia + 2, column=col_right).alignment = align_sig_center
    
    ws_iiia.cell(row=row_sig_iiia + 7, column=col_right, value="( ................................................. )").font = font_sig
    ws_iiia.cell(row=row_sig_iiia + 7, column=col_right).alignment = align_sig_center

    # Notes (Catatan)
    row_notes = row_sig_iiia + 13
    notes_list = [
        ("Catatan:", True),
        ("*) Alasan Tidak Lolos, diisi dengan angka (1-8) sebagai berikut:", False),
        ("1. Belum memiliki KK sendiri;", False),
        ("2. Tanah bersengketa;", False),
        ("3. Rumah dalam kondisi layak;", False),
        ("4. Memiliki rumah lebih dari 1;", False),
        ("5. Pernah memperoleh bantuan dari APBN/APBD/CSR/anggaran lainnya;", False),
        ("6. Penghasilan lebih dari UMP;", False),
        ("7. Memilih untuk dibantu dengan sumber anggaran lain;", False),
        ("8. Menghuni kurang dari 3 tahun;", False),
        ("9. Lainnya (diisi pada kolom keterangan);", False)
    ]
    for n_idx, (ntxt, is_b) in enumerate(notes_list):
        cell_n = ws_iiia.cell(row=row_notes + n_idx, column=1, value=ntxt)
        cell_n.font = Font(name='Bookman Old Style', size=14, bold=is_b)
        
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    
    log_activity(
        username="Admin",
        action="EXPORT_EXCEL",
        entity_type="STAGE_VERFAL",
        entity_name=stage_name,
        details=f"Ekspor Excel Verfal gabungan seluruh kabupaten pada tahap '{stage_name}'"
    )
    
    clean_stage = stage_name.replace(' ', '_')
    clean_prov = prov_name.replace(' ', '_')
    filename = f"VERFAL_GABUNGAN_{clean_stage}_{clean_prov}.xlsx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.post("/api/verified/records/bulk-delete")
async def bulk_delete_verified_records(record_ids: str = Form(...)):
    """Bulk delete verified records by comma-separated IDs."""
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        ids = [int(rid.strip()) for rid in record_ids.split(',') if rid.strip()]
        if not ids:
            raise HTTPException(status_code=400, detail="Tidak ada ID yang diberikan")
        
        placeholders = ','.join(['?' for _ in ids])
        
        # Get NIKs before deleting
        cursor.execute(f"SELECT DISTINCT no_ktp FROM verified_records WHERE id IN ({placeholders})", ids)
        niks = [row['no_ktp'] for row in cursor.fetchall()]
        
        # Delete replacement events
        cursor.execute(f"DELETE FROM replacement_events WHERE disqualified_record_id IN ({placeholders})", ids)
        
        # Delete verified records
        cursor.execute(f"DELETE FROM verified_records WHERE id IN ({placeholders})", ids)
        deleted_count = cursor.rowcount
        
        # Clean up reconciliation overrides
        for nik in niks:
            cursor.execute("DELETE FROM reconciliation_overrides WHERE original_no_ktp = ?", (nik,))
        
        conn.commit()
    except HTTPException:
        conn.close()
        raise
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"Gagal menghapus data secara bulk: {str(e)}")
    conn.close()
    return {"message": f"{deleted_count} data verifikasi berhasil dihapus", "deleted_count": deleted_count}


@app.post("/api/verified/records/bulk-delete-by-nik")
async def bulk_delete_verified_records_by_nik(niks: str = Form(...), stage_id: int = Form(None)):
    """Bulk delete verified records by comma-separated NIKs. Optionally scoped to a stage."""
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        nik_list = [nik.strip() for nik in niks.split(',') if nik.strip()]
        if not nik_list:
            raise HTTPException(status_code=400, detail="Tidak ada NIK yang diberikan")
        
        placeholders = ','.join(['?' for _ in nik_list])
        
        if stage_id:
            cursor.execute(f"""
                SELECT vr.id FROM verified_records vr
                JOIN verified_batches vb ON vb.id = vr.batch_id
                WHERE vr.no_ktp IN ({placeholders}) AND vb.stage_id = ?
            """, nik_list + [stage_id])
        else:
            cursor.execute(f"SELECT id FROM verified_records WHERE no_ktp IN ({placeholders})", nik_list)
        
        record_ids = [str(row['id']) for row in cursor.fetchall()]
        
        if not record_ids:
            return {"message": "Tidak ada data yang cocok untuk dihapus", "deleted_count": 0}
        
        id_placeholders = ','.join(['?' for _ in record_ids])
        id_ints = [int(rid) for rid in record_ids]
        
        cursor.execute(f"DELETE FROM replacement_events WHERE disqualified_record_id IN ({id_placeholders})", id_ints)
        cursor.execute(f"DELETE FROM verified_records WHERE id IN ({id_placeholders})", id_ints)
        deleted_count = cursor.rowcount
        
        for nik in nik_list:
            cursor.execute("DELETE FROM reconciliation_overrides WHERE original_no_ktp = ?", (nik,))
        
        conn.commit()
    except HTTPException:
        conn.close()
        raise
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"Gagal menghapus data secara bulk: {str(e)}")
    conn.close()
    return {"message": f"{deleted_count} data verifikasi berhasil dihapus", "deleted_count": deleted_count}


@app.post("/api/reconciliation/bulk-override")
async def bulk_reconciliation_override(
    stage_id: int = Form(...),
    niks: str = Form(...),
    override_type: str = Form(...)
):
    """Bulk apply reconciliation override to multiple NIKs."""
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        nik_list = [nik.strip() for nik in niks.split(',') if nik.strip()]
        if not nik_list:
            raise HTTPException(status_code=400, detail="Tidak ada NIK yang diberikan")
        
        count = 0
        for nik in nik_list:
            cursor.execute("""
                INSERT OR REPLACE INTO reconciliation_overrides (
                    stage_id, original_no_ktp, override_type, corrected_nama, corrected_no_ktp, corrected_no_kk
                ) VALUES (?, ?, ?, NULL, NULL, NULL)
            """, (stage_id, nik, override_type))
            count += cursor.rowcount
        
        conn.commit()
    except HTTPException:
        conn.close()
        raise
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"Gagal melakukan rekonsiliasi bulk: {str(e)}")
    conn.close()
    return {"message": f"{count} data berhasil direkonsiliasi", "reconciled_count": count}


# --- INVERS MANUAL PAIRS (rekonsiliasi manual invers ↔ verified) ---

@app.get("/api/reconciliation/unmatched-invers/{stage_id}")
def get_unmatched_invers(stage_id: int):
    """Ambil invers records yang belum punya pasangan di verified_records."""
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute("""
        SELECT ir.nama, ir.no_ktp, ir.no_kk,
               UPPER(TRIM(COALESCE(ir.kabupaten_kota, ''))) as kabupaten_kota,
               UPPER(TRIM(COALESCE(ir.kecamatan, ''))) as kecamatan,
               UPPER(TRIM(COALESCE(ir.desa_kelurahan, ''))) as desa_kelurahan
        FROM invers_records ir
        JOIN invers_revisions irv ON ir.revision_id = irv.id
        WHERE irv.stage_id = ? AND irv.is_active = 1
    """, (stage_id,))
    invers_rows = [dict(r) for r in cursor.fetchall()]

    cursor.execute("""
        SELECT DISTINCT vr.no_ktp FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
        WHERE vb.stage_id = ? AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
    """, (stage_id,))
    verified_niks = {r['no_ktp'].strip() for r in cursor.fetchall()}

    cursor.execute("""
        SELECT invers_nik FROM invers_manual_pairs WHERE stage_id = ?
    """, (stage_id,))
    paired_niks = {r['invers_nik'].strip() for r in cursor.fetchall()}

    conn.close()

    unmatched = []
    for ir in invers_rows:
        nik = ir['no_ktp'].strip()
        if nik not in verified_niks and nik not in paired_niks:
            unmatched.append(ir)

    return {"records": unmatched, "total": len(unmatched)}


@app.get("/api/stage/{stage_id}/reconciliation/export")
def export_reconciliation_excel(stage_id: int, error_filter: str = "ALL", active_only: bool = False):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Ambil stage info
    cursor.execute("SELECT * FROM invers_stages WHERE id = ?", (stage_id,))
    stage = cursor.fetchone()
    if not stage:
        conn.close()
        raise HTTPException(status_code=404, detail="Tahap tidak ditemukan")
    stage_name = stage['name']
    conn.close()
    
    # Gunakan get_stage_records untuk mendapatkan analisis mismatch yang sama persis
    stage_data = get_stage_records(stage_id)
    verified = stage_data.get("verified_records", [])
    
    # Filter mismatch records
    mismatches = [r for r in verified if r.get('has_error') or r.get('is_mismatch')]
    
    if active_only:
        mismatches = [r for r in mismatches if not r.get('override')]

    if error_filter != "ALL":
        mismatches = [r for r in mismatches if r.get('mismatch_type') == error_filter]

    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    import io
    from fastapi.responses import StreamingResponse
    from datetime import datetime

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Penyandingan Data Mismatch"
    ws.views.sheetView[0].showGridLines = True

    # Palette Warna
    f_title = Font(name="Bookman Old Style", size=14, bold=True, color="1B4332")
    f_subtitle = Font(name="Bookman Old Style", size=10, italic=True, color="475569")
    f_grp_header = Font(name="Bookman Old Style", size=10, bold=True, color="FFFFFF")
    f_field_label = Font(name="Bookman Old Style", size=9, bold=True, color="1E293B")
    f_data = Font(name="Bookman Old Style", size=9, color="000000")
    f_data_bold = Font(name="Bookman Old Style", size=9, bold=True, color="000000")
    f_status_val = Font(name="Bookman Old Style", size=8.5, color="334155")

    fill_grp_num = PatternFill(start_color="475569", end_color="475569", fill_type="solid") # Slate
    fill_grp_label = PatternFill(start_color="64748B", end_color="64748B", fill_type="solid") # Slate Light
    fill_grp_invers = PatternFill(start_color="1B4332", end_color="1B4332", fill_type="solid") # Dark Green
    fill_grp_verified = PatternFill(start_color="7F1D1D", end_color="7F1D1D", fill_type="solid") # Dark Red
    fill_grp_status = PatternFill(start_color="1E3A8A", end_color="1E3A8A", fill_type="solid") # Dark Blue

    fill_block_header = PatternFill(start_color="F1F5F9", end_color="F1F5F9", fill_type="solid") # Light Grey
    fill_mismatch = PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid") # Soft Red
    fill_status_bg = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid") # Soft White

    border_thin = Border(
        left=Side(style='thin', color='CBD5E1'),
        right=Side(style='thin', color='CBD5E1'),
        top=Side(style='thin', color='CBD5E1'),
        bottom=Side(style='thin', color='CBD5E1')
    )
    
    border_block_bottom = Border(
        left=Side(style='thin', color='CBD5E1'),
        right=Side(style='thin', color='CBD5E1'),
        top=Side(style='thin', color='CBD5E1'),
        bottom=Side(style='medium', color='475569')
    )

    align_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    align_left = Alignment(horizontal='left', vertical='center', wrap_text=True)

    # Title section
    status_filter_label = "KASUS AKTIF SAJA (BELUM REKONSILIASI)" if active_only else "SEMUA KASUS"
    ws.cell(row=1, column=1, value=f"LAPORAN PENYANDINGAN DATA MISMATCH ({status_filter_label})").font = f_title
    ws.cell(row=2, column=1, value=f"Tahap: {stage_name.upper()} | Filter Status: {status_filter_label} | Kategori: {error_filter} | Total Kasus: {len(mismatches)} | Tanggal Ekspor: {datetime.now().strftime('%d-%m-%Y %H:%M')}").font = f_subtitle

    # Merge title headers
    ws.merge_cells("A1:E1")
    ws.merge_cells("A2:E2")

    # Header Row 4: Column Headers
    ws.cell(row=4, column=1, value="NO.").font = f_grp_header
    ws.cell(row=4, column=1).fill = fill_grp_num
    ws.cell(row=4, column=1).alignment = align_center

    ws.cell(row=4, column=2, value="ELEMEN DATA").font = f_grp_header
    ws.cell(row=4, column=2).fill = fill_grp_label
    ws.cell(row=4, column=2).alignment = align_center

    ws.cell(row=4, column=3, value="DATA RUJUKAN (FILE AWAL INVERS)").font = f_grp_header
    ws.cell(row=4, column=3).fill = fill_grp_invers
    ws.cell(row=4, column=3).alignment = align_center

    ws.cell(row=4, column=4, value="DATA LAPANGAN (FILE HASIL VERIFIKASI)").font = f_grp_header
    ws.cell(row=4, column=4).fill = fill_grp_verified
    ws.cell(row=4, column=4).alignment = align_center

    ws.cell(row=4, column=5, value="INFORMASI ERROR & STATUS REKONSILIASI").font = f_grp_header
    ws.cell(row=4, column=5).fill = fill_grp_status
    ws.cell(row=4, column=5).alignment = align_center

    ws.row_dimensions[4].height = 28

    # Data Blocks (6 Rows per CPB Mismatch)
    current_row = 5
    for idx, r in enumerate(mismatches):
        exp = r.get('expected_invers') or {}
        override = r.get('override')
        
        m_type = r.get('mismatch_type')
        m_label = (
            "Duplikat Data" if m_type == "DUPLICATE" else
            "Mismat Nama CPB" if m_type == "NAMA_MISMATCH" else
            "Mismat No. KK" if m_type == "KK_MISMATCH" else
            "Mismat NIK CPB" if m_type == "NIK_MISMATCH" else
            "NIK Tidak Ada di INVERS" if m_type == "MISSING_IN_INVERS" else
            "Format NIK Salah" if m_type == "NIK_INVALID" else
            "Format KK Salah" if m_type == "KK_INVALID" else
            "NIK & KK Identik" if m_type == "NIK_KK_IDENTICAL" else "Perbedaan Karakter"
        )
        
        status_recon = "Belum Selesai (Butuh Tindakan)"
        if override:
            status_recon = "Resolved (Terima Lapangan)" if override.get('override_type') == "ACCEPT_VERIFIED" else "Resolved (Edit Manual)"
            
        catatan = "; ".join(r.get('errors', [])) if r.get('errors') else "-"

        # Block Range: current_row to current_row + 5
        start_r = current_row
        end_r = current_row + 5

        # Merge No Column (Col A)
        ws.merge_cells(start_row=start_r, start_column=1, end_row=end_r, end_column=1)
        no_cell = ws.cell(row=start_r, column=1, value=idx + 1)
        no_cell.font = Font(name="Bookman Old Style", size=12, bold=True, color="1E293B")
        no_cell.alignment = align_center

        # Merge Status Column (Col E)
        ws.merge_cells(start_row=start_r, start_column=5, end_row=end_r, end_column=5)
        status_text = (
            f"KASUS ERROR:\n{m_label}\n\n"
            f"BERITA ACARA / BATCH:\n{r.get('batch_name') or '-'}\n\n"
            f"STATUS REKONSILIASI:\n{status_recon}\n\n"
            f"CATATAN / DIAGNOSA:\n{catatan}"
        )
        status_cell = ws.cell(row=start_r, column=5, value=status_text)
        status_cell.font = f_status_val
        status_cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)

        # Apply borders & background for merged Col A and Col E cells
        for r_idx in range(start_r, end_r + 1):
            cell_a = ws.cell(row=r_idx, column=1)
            cell_e = ws.cell(row=r_idx, column=5)
            b_style = border_block_bottom if r_idx == end_r else border_thin
            cell_a.border = b_style
            cell_e.border = b_style
            cell_e.fill = fill_status_bg

        # Rows Data setup
        field_rows = [
            ("Nama CPB", exp.get('nama') or "-", r.get('nama') or "-"),
            ("NIK (No. KTP)", f"'{exp.get('no_ktp')}" if exp.get('no_ktp') else "-", f"'{r.get('no_ktp')}" if r.get('no_ktp') else "-"),
            ("No. KK", f"'{exp.get('no_kk')}" if exp.get('no_kk') else "-", f"'{r.get('no_kk')}" if r.get('no_kk') else "-"),
            ("Kabupaten / Kota", exp.get('kabupaten_kota') or "-", r.get('kabupaten_kota') or "-"),
            ("Kecamatan", exp.get('kecamatan') or "-", r.get('kecamatan') or "-"),
            ("Desa / Kelurahan", exp.get('desa_kelurahan') or "-", r.get('desa_kelurahan') or "-")
        ]

        for i, (field_lbl, inv_val, ver_val) in enumerate(field_rows):
            r_curr = start_r + i
            b_style = border_block_bottom if r_curr == end_r else border_thin

            # Col B: Field Label
            cell_b = ws.cell(row=r_curr, column=2, value=field_lbl)
            cell_b.font = f_field_label
            cell_b.border = b_style
            cell_b.alignment = align_left
            if i == 0:
                cell_b.fill = fill_block_header

            # Col C: INVERS Value
            cell_c = ws.cell(row=r_curr, column=3, value=inv_val)
            cell_c.font = f_data_bold if i == 0 else f_data
            cell_c.border = b_style
            cell_c.alignment = align_center if i in [1, 2] else align_left
            if i in [1, 2]:
                cell_c.number_format = '@'
            if i == 0:
                cell_c.fill = fill_block_header

            # Col D: VERIFIED Value
            cell_d = ws.cell(row=r_curr, column=4, value=ver_val)
            cell_d.font = f_data_bold if i == 0 else f_data
            cell_d.border = b_style
            cell_d.alignment = align_center if i in [1, 2] else align_left
            if i in [1, 2]:
                cell_d.number_format = '@'
            if i == 0:
                cell_d.fill = fill_block_header

            # Mismatch highlight
            if exp:
                if i == 0 and exp.get('nama', '').strip().upper() != r.get('nama', '').strip().upper():
                    cell_d.fill = fill_mismatch
                    cell_d.font = f_data_bold
                elif i == 1 and exp.get('no_ktp', '').strip() != r.get('no_ktp', '').strip():
                    cell_d.fill = fill_mismatch
                    cell_d.font = f_data_bold
                elif i == 2 and exp.get('no_kk', '').strip() != r.get('no_kk', '').strip():
                    cell_d.fill = fill_mismatch
                    cell_d.font = f_data_bold

            ws.row_dimensions[r_curr].height = 22

        current_row = end_r + 1

    # Widths
    col_widths = {
        'A': 7,
        'B': 22,
        'C': 34,
        'D': 34,
        'E': 42
    }
    for col_let, width in col_widths.items():
        ws.column_dimensions[col_let].width = width

    ws.freeze_panes = "A5"

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    file_type_str = "KASUS_AKTIF" if active_only else "SEMUA_KASUS"
    filename = f"PENYANDINGAN_DATA_{file_type_str}_{stage_name.upper().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@app.get("/api/reconciliation/unmatched-verified/{stage_id}")
def get_unmatched_verified(stage_id: int):
    """Ambil verified records yang belum dipasangkan dengan invers record manapun."""
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute("""
        SELECT vr.id, vr.no_ktp, vr.no_kk, vr.nama, vr.desa_kelurahan,
               vr.kecamatan, vr.kabupaten_kota, vr.status
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
        WHERE vb.stage_id = ?
          AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
          AND vr.id NOT IN (
            SELECT mp.verified_record_id FROM invers_manual_pairs mp WHERE mp.stage_id = ?
          )
        ORDER BY vr.nama
    """, (stage_id, stage_id))
    records = [dict(r) for r in cursor.fetchall()]
    conn.close()

    return {"records": records, "total": len(records)}


@app.get("/api/reconciliation/manual-pairs/{stage_id}")
def get_manual_pairs(stage_id: int):
    """Ambil semua manual pairs untuk suatu tahap."""
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute("""
        SELECT mp.id, mp.invers_nik, mp.invers_nama, mp.invers_kabupaten,
               vr.nama as verified_nama, vr.no_ktp as verified_nik,
               vr.kabupaten_kota as verified_kabupaten, vr.status as verified_status
        FROM invers_manual_pairs mp
        LEFT JOIN verified_records vr ON vr.id = mp.verified_record_id
        WHERE mp.stage_id = ?
        ORDER BY mp.created_at DESC
    """, (stage_id,))
    pairs = [dict(r) for r in cursor.fetchall()]
    conn.close()

    return {"pairs": pairs, "total": len(pairs)}


@app.post("/api/reconciliation/pair-invers")
async def pair_invers_with_verified(
    stage_id: int = Form(...),
    invers_nik: str = Form(...),
    invers_nama: str = Form(...),
    invers_kabupaten: str = Form(''),
    verified_record_id: int = Form(...)
):
    """Pasangkan invers record dengan verified record."""
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT id, no_ktp, nama, status FROM verified_records WHERE id = ?", (verified_record_id,))
    vr = cursor.fetchone()
    if not vr:
        conn.close()
        raise HTTPException(status_code=404, detail="Verified record tidak ditemukan")

    try:
        cursor.execute("""
            INSERT OR REPLACE INTO invers_manual_pairs
                (stage_id, invers_nik, invers_nama, invers_kabupaten, verified_record_id)
            VALUES (?, ?, ?, ?, ?)
        """, (stage_id, invers_nik.strip(), invers_nama, invers_kabupaten, verified_record_id))
        conn.commit()
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=400, detail=f"Gagal menyimpan pasangan: {str(e)}")

    conn.close()
    return {
        "message": f"Berhasil memasangkan NIK {invers_nik} dengan verified record ({vr['nama']}, {vr['status']})",
        "verified_nama": vr['nama'],
        "verified_status": vr['status']
    }


@app.post("/api/reconciliation/auto-pair-nik/{stage_id}")
def auto_pair_by_nik(stage_id: int):
    """Otomatis pasangkan semua invers records yang NIK-nya cocok dengan verified_records (NIK atau KK)."""
    conn = get_db_connection()
    cursor = conn.cursor()

    # Ambil semua verified records untuk stage ini (yang aktif, bukan duplicate)
    cursor.execute("""
        SELECT vr.id, vr.no_ktp, vr.no_kk, vr.nama, vr.desa_kelurahan, vr.status
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
        WHERE vb.stage_id = ? AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
    """, (stage_id,))
    verified_map = {}
    verified_by_kk = {}
    for r in cursor.fetchall():
        d = dict(r)
        nik = r['no_ktp'].strip()
        verified_map[nik] = d
        kk = r['no_kk'].strip() if r['no_kk'] else ''
        if kk:
            verified_by_kk[kk] = d

    # Ambil NIK yang sudah dipasang
    cursor.execute("SELECT invers_nik FROM invers_manual_pairs WHERE stage_id = ?", (stage_id,))
    paired_niks = {r['invers_nik'].strip() for r in cursor.fetchall()}

    # Ambil semua invers records
    cursor.execute("""
        SELECT ir.nama, ir.no_ktp, ir.no_kk,
               UPPER(TRIM(COALESCE(ir.kabupaten_kota, ''))) as kabupaten_kota,
               UPPER(TRIM(COALESCE(ir.kecamatan, ''))) as kecamatan,
               UPPER(TRIM(COALESCE(ir.desa_kelurahan, ''))) as desa_kelurahan
        FROM invers_records ir
        JOIN invers_revisions irv ON ir.revision_id = irv.id
        WHERE irv.stage_id = ? AND irv.is_active = 1
    """, (stage_id,))
    invers_rows = [dict(r) for r in cursor.fetchall()]

    paired_records = []
    already_paired = 0
    no_match = 0

    for ir in invers_rows:
        nik = ir['no_ktp'].strip()
        if nik in paired_niks:
            already_paired += 1
            continue
        vr = verified_map.get(nik) or verified_by_kk.get(nik)
        if vr:
            try:
                cursor.execute("""
                    INSERT OR REPLACE INTO invers_manual_pairs
                        (stage_id, invers_nik, invers_nama, invers_kabupaten, verified_record_id)
                    VALUES (?, ?, ?, ?, ?)
                """, (stage_id, nik, ir['nama'], ir['kabupaten_kota'], vr['id']))
                paired_records.append({
                    "invers_nik": nik,
                    "invers_nama": ir['nama'],
                    "verified_nama": vr['nama'],
                    "verified_status": vr['status'],
                    "desa": ir['desa_kelurahan']
                })
            except Exception:
                pass
        else:
            no_match += 1

    conn.commit()
    conn.close()

    return {
        "message": f"Berhasil auto-pair {len(paired_records)} data. {no_match} data belum cocok.",
        "paired_count": len(paired_records),
        "already_paired": already_paired,
        "no_match": no_match,
        "paired_records": paired_records
    }


@app.get("/api/reconciliation/suggest-pairs/{stage_id}")
def suggest_pairs_by_name_desa(stage_id: int):
    """Sarankan pasangan berdasarkan kecocokan nama + desa untuk invers yang belum terpasang."""
    conn = get_db_connection()
    cursor = conn.cursor()

    # Ambil verified records, build map by (nama, desa)
    cursor.execute("""
        SELECT vr.id, vr.no_ktp, vr.no_kk, vr.nama, vr.desa_kelurahan, vr.status
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
        WHERE vb.stage_id = ? AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
    """, (stage_id,))
    verified_by_name_desa = {}
    for r in cursor.fetchall():
        key = (r['nama'].upper().strip(), r['desa_kelurahan'].upper().strip() if r['desa_kelurahan'] else '')
        if key not in verified_by_name_desa:
            verified_by_name_desa[key] = []
        verified_by_name_desa[key].append(dict(r))

    # Ambil NIK yang sudah dipasang
    cursor.execute("SELECT invers_nik FROM invers_manual_pairs WHERE stage_id = ?", (stage_id,))
    paired_niks = {r['invers_nik'].strip() for r in cursor.fetchall()}

    # Ambil unmatched invers
    cursor.execute("""
        SELECT ir.nama, ir.no_ktp, ir.no_kk,
               UPPER(TRIM(COALESCE(ir.kabupaten_kota, ''))) as kabupaten_kota,
               UPPER(TRIM(COALESCE(ir.kecamatan, ''))) as kecamatan,
               UPPER(TRIM(COALESCE(ir.desa_kelurahan, ''))) as desa_kelurahan
        FROM invers_records ir
        JOIN invers_revisions irv ON ir.revision_id = irv.id
        WHERE irv.stage_id = ? AND irv.is_active = 1
    """, (stage_id,))
    invers_rows = [dict(r) for r in cursor.fetchall()]

    conn.close()

    suggestions = []
    for ir in invers_rows:
        nik = ir['no_ktp'].strip()
        if nik in paired_niks:
            continue
        key = (ir['nama'].upper().strip(), ir['desa_kelurahan'].upper().strip() if ir['desa_kelurahan'] else '')
        candidates = verified_by_name_desa.get(key, [])
        if len(candidates) == 1:
            suggestions.append({
                "invers": ir,
                "verified": candidates[0]
            })

    return {"suggestions": suggestions, "total": len(suggestions)}


@app.post("/api/reconciliation/batch-pair")
def batch_pair(body: dict = Body(...)):
    """Pasangkan beberapa invers sekaligus dengan verified records."""
    stage_id = body.get("stage_id")
    pairs = body.get("pairs", [])
    if not stage_id or not pairs:
        raise HTTPException(status_code=400, detail="stage_id dan pairs harus diisi")

    conn = get_db_connection()
    cursor = conn.cursor()
    paired_count = 0
    errors = []

    for p in pairs:
        invers_nik = p.get("invers_nik", "").strip()
        invers_nama = p.get("invers_nama", "")
        invers_kabupaten = p.get("invers_kabupaten", "")
        verified_record_id = p.get("verified_record_id")
        if not invers_nik or not verified_record_id:
            errors.append(f"Data tidak lengkap: {invers_nik}")
            continue
        try:
            cursor.execute("""
                INSERT OR REPLACE INTO invers_manual_pairs
                    (stage_id, invers_nik, invers_nama, invers_kabupaten, verified_record_id)
                VALUES (?, ?, ?, ?, ?)
            """, (stage_id, invers_nik, invers_nama, invers_kabupaten, verified_record_id))
            paired_count += 1
        except Exception as e:
            errors.append(f"Gagal pair {invers_nik}: {str(e)}")

    conn.commit()
    conn.close()

    return {
        "message": f"Berhasil memasangkan {paired_count} data.",
        "paired_count": paired_count,
        "errors": errors
    }


@app.delete("/api/reconciliation/unpair-invers/{pair_id}")
async def unpair_invers(pair_id: int):
    """Hapus pasangan invers ↔ verified."""
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT id FROM invers_manual_pairs WHERE id = ?", (pair_id,))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="Pasangan tidak ditemukan")

    cursor.execute("DELETE FROM invers_manual_pairs WHERE id = ?", (pair_id,))
    conn.commit()
    conn.close()

    return {"message": "Pasangan berhasil dihapus"}


# --- PENGUSUL TREE EXPORT ---

@app.get("/api/stage/{stage_id}/pengusul-tree/export")
def export_pengusul_tree(stage_id: int):
    """Export pengusul hierarchy tree to Excel."""
    import re as re_mod
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT name FROM invers_stages WHERE id = ?", (stage_id,))
    stage_row = cursor.fetchone()
    stage_name = stage_row['name'] if stage_row else "Tahap"
    
    # Get tree data
    tree_data = get_pengusul_tree(stage_id)
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Hirarki Pengusul"
    
    header_font = Font(name='Bookman Old Style', size=11, bold=True)
    data_font = Font(name='Bookman Old Style', size=10)
    header_fill = PatternFill(start_color="0F5132", end_color="0F5132", fill_type="solid")
    header_font_white = Font(name='Bookman Old Style', size=11, bold=True, color="FFFFFF")
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    center_align = Alignment(horizontal='center', vertical='center')
    left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
    
    # Headers
    headers = ['No', 'Pengusul', 'Kabupaten/Kota', 'Kecamatan', 'Desa/Kelurahan', 'CPB', 'Lolos', 'Tidak Lolos', 'Belum', 'Sudah SK Dirjen', 'Belum SK Dirjen']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font_white
        cell.fill = header_fill
        cell.alignment = center_align
        cell.border = thin_border
    
    row_num = 2
    counter = 1
    
    for pengusul in tree_data:
        p_name = pengusul['name']
        for kab in pengusul.get('children', []):
            kb_name = kab['name']
            for kec in kab.get('children', []):
                kc_name = kec['name']
                for desa in kec.get('children', []):
                    ds_name = desa['name']
                    values = [counter, p_name, kb_name, kc_name, ds_name,
                              desa['cpb'], desa['lolos'], desa['tidak_lolos'], desa['belum_verifikasi'],
                              desa.get('sk_dirjen_sudah', 0), desa.get('sk_dirjen_belum', 0)]
                    for col, val in enumerate(values, 1):
                        cell = ws.cell(row=row_num, column=col, value=val)
                        cell.font = data_font
                        cell.border = thin_border
                        cell.alignment = center_align if col in (1, 6, 7, 8, 9, 10, 11) else left_align
                    row_num += 1
                    counter += 1
        
        # Add summary row for pengusul
        summary_values = ['', f'TOTAL {p_name}', '', '', '',
                          pengusul['cpb'], pengusul['lolos'], pengusul['tidak_lolos'], pengusul['belum_verifikasi'],
                          pengusul.get('sk_dirjen_sudah', 0), pengusul.get('sk_dirjen_belum', 0)]
        for col, val in enumerate(summary_values, 1):
            cell = ws.cell(row=row_num, column=col, value=val)
            cell.font = Font(name='Bookman Old Style', size=10, bold=True)
            cell.fill = PatternFill(start_color="D1E7DD", end_color="D1E7DD", fill_type="solid")
            cell.border = thin_border
            cell.alignment = center_align if col in (1, 6, 7, 8, 9, 10, 11) else left_align
        row_num += 1
    
    # Column widths
    col_widths = [6, 25, 25, 25, 25, 8, 8, 12, 8, 15, 15]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    
    # Auto filter
    ws.auto_filter.ref = f"A1:K{row_num - 1}"
    
    conn.close()
    
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    filename = f"Hirarki_Pengusul_{stage_name.replace(' ', '_')}.xlsx"
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@app.get("/api/stage/{stage_id}/kabupaten-pengusul-tree")
def get_kabupaten_pengusul_tree(stage_id: int):
    """Ambil data rekap per kabupaten dengan children berupa daftar pengusul."""
    tree_data = get_pengusul_tree(stage_id)

    kab_map = {}
    for pengusul in tree_data:
        p_name = pengusul['name']
        for kab in pengusul.get('children', []):
            kb_name = kab['name']
            if kb_name not in kab_map:
                kab_map[kb_name] = {
                    "name": kb_name,
                    "cpb": 0, "lolos": 0, "tidak_lolos": 0, "belum_verifikasi": 0,
                    "sk_dirjen_sudah": 0, "sk_dirjen_belum": 0,
                    "children": []
                }
            node = kab_map[kb_name]
            node["cpb"] += kab['cpb']
            node["lolos"] += kab['lolos']
            node["tidak_lolos"] += kab['tidak_lolos']
            node["belum_verifikasi"] += kab['belum_verifikasi']
            node["sk_dirjen_sudah"] += kab.get('sk_dirjen_sudah', 0)
            node["sk_dirjen_belum"] += kab.get('sk_dirjen_belum', 0)
            node["children"].append({
                "name": p_name,
                "cpb": kab['cpb'],
                "lolos": kab['lolos'],
                "tidak_lolos": kab['tidak_lolos'],
                "belum_verifikasi": kab['belum_verifikasi'],
                "sk_dirjen_sudah": kab.get('sk_dirjen_sudah', 0),
                "sk_dirjen_belum": kab.get('sk_dirjen_belum', 0)
            })

    result = sorted(kab_map.values(), key=lambda x: x['name'])
    return result


# --- END OF NEW ENDPOINTS ---

# --- REKAP KESELURUHAN (All stages, per kabupaten) ---
@app.get("/api/rekap-keseluruhan/export")
def export_rekap_keseluruhan(pengusul: str = ""):
    import re
    pengusul_list = [p.strip() for p in pengusul.split(",") if p.strip()]
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT id, name FROM invers_stages ORDER BY id ASC")
    all_stages = [dict(r) for r in cursor.fetchall()]
    def get_stage_num(s):
        match = re.search(r'\d+', s['name'])
        return int(match.group()) if match else 999
    all_stages = sorted(all_stages, key=get_stage_num)
    
    published_filter = "AND vb.is_published = 1"
    
    pengusul_clause = ""
    pengusul_params = ()
    if pengusul_list:
        ph = ",".join("?" for _ in pengusul_list)
        pengusul_clause = f"UPPER(COALESCE(ir_p.pengusul, '')) IN ({ph})"
        pengusul_params = tuple(p.upper() for p in pengusul_list)
    
    cursor.execute(f"""
        SELECT DISTINCT UPPER(TRIM(COALESCE(kabupaten_kota, ''))) as kab 
        FROM invers_records ir
        JOIN invers_revisions irv ON ir.revision_id = irv.id
        WHERE irv.is_active = 1 AND TRIM(COALESCE(kabupaten_kota, '')) != ''
        UNION
        SELECT DISTINCT UPPER(TRIM(COALESCE(vr.kabupaten_kota, ''))) as kab
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
        WHERE (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL) AND TRIM(COALESCE(vr.kabupaten_kota, '')) != ''
        {published_filter}
        ORDER BY kab ASC
    """)
    all_kabupaten = [r['kab'] for r in cursor.fetchall()]
    
    sk_where_extra = ""
    sk_params = ()
    if pengusul_list:
        ph = ",".join("?" for _ in pengusul_list)
        sk_where_extra = f"AND UPPER(COALESCE(ir_p.pengusul, '')) IN ({ph})"
        sk_params = tuple(p.upper() for p in pengusul_list)
    cursor.execute(f"""
        SELECT 
            m.verified_stage_id as stage_id,
            UPPER(TRIM(COALESCE(vr.kabupaten_kota, ''))) as kab,
            COUNT(*) as cnt
        FROM sk_dirjen_matches m
        JOIN verified_records vr ON vr.id = m.verified_record_id
        LEFT JOIN (
            SELECT DISTINCT no_ktp, pengusul FROM invers_records
            WHERE pengusul IS NOT NULL AND pengusul != ''
        ) ir_p ON ir_p.no_ktp = vr.no_ktp
        WHERE m.verified_stage_id IS NOT NULL
        AND m.verified_record_id IS NOT NULL
        AND (m.match_type = 'PERFECT' 
            OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED')
            OR m.match_type = 'MANUAL_PAIR')
        {sk_where_extra}
        GROUP BY m.verified_stage_id, kab
    """, sk_params)
    sk_by_stage_kab = {}
    for row in cursor.fetchall():
        sid = row['stage_id']
        kab = row['kab']
        if sid not in sk_by_stage_kab:
            sk_by_stage_kab[sid] = {}
        sk_by_stage_kab[sid][kab] = row['cnt']
    
    stages_data = []
    
    for stage in all_stages:
        stage_id = stage['id']
        stage_name = stage['name']
        
        if pengusul_list:
            invers_clause = f"UPPER(COALESCE(ir.pengusul, '')) IN ({','.join('?' for _ in pengusul_list)})"
            cursor.execute(f"""
                SELECT ir.no_ktp, UPPER(TRIM(COALESCE(ir.kabupaten_kota, ''))) as kab
                FROM invers_records ir
                JOIN invers_revisions irv ON ir.revision_id = irv.id
                WHERE irv.stage_id = ? AND irv.is_active = 1
                AND {invers_clause}
            """, (stage_id, *pengusul_params))
            invers_recs = [dict(r) for r in cursor.fetchall()]

            verif_query = f"""
                SELECT vr.no_ktp, vr.status, UPPER(TRIM(COALESCE(NULLIF(TRIM(vr.kabupaten_kota), ''), NULLIF(TRIM(vb.kabupaten), '')))) as kab
                FROM verified_records vr
                JOIN verified_batches vb ON vr.batch_id = vb.id
                LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
                LEFT JOIN (
                    SELECT DISTINCT no_ktp, pengusul FROM invers_records
                    WHERE pengusul IS NOT NULL AND pengusul != ''
                ) ir_p ON ir_p.no_ktp = vr.no_ktp
                WHERE vb.stage_id = ? AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
                AND {pengusul_clause}
                {published_filter}
            """
            cursor.execute(verif_query, (stage_id, *pengusul_params))
            verified_recs = [dict(r) for r in cursor.fetchall()]
        else:
            cursor.execute("""
                SELECT ir.no_ktp, UPPER(TRIM(COALESCE(ir.kabupaten_kota, ''))) as kab
                FROM invers_records ir
                JOIN invers_revisions irv ON ir.revision_id = irv.id
                WHERE irv.stage_id = ? AND irv.is_active = 1
            """, (stage_id,))
            invers_recs = [dict(r) for r in cursor.fetchall()]
            cursor.execute(f"""
                SELECT vr.no_ktp, vr.status, UPPER(TRIM(COALESCE(NULLIF(TRIM(vr.kabupaten_kota), ''), NULLIF(TRIM(vb.kabupaten), '')))) as kab
                FROM verified_records vr
                JOIN verified_batches vb ON vr.batch_id = vb.id
                LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
                WHERE vb.stage_id = ? AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL)
                {published_filter}
            """, (stage_id,))
            verified_recs = [dict(r) for r in cursor.fetchall()]
        
        alokasi_by_kab = {}
        for ir in invers_recs:
            kab = ir['kab'] if ir['kab'] else 'TIDAK DIKETAHUI'
            alokasi_by_kab[kab] = alokasi_by_kab.get(kab, 0) + 1
            
        verif_by_kab = {}
        for vr in verified_recs:
            kab = vr['kab'] if vr['kab'] else 'TIDAK DIKETAHUI'
            if kab not in verif_by_kab:
                verif_by_kab[kab] = {"lolos": 0, "tidak_lolos": 0}
            if vr['status'] == 'LOLOS':
                verif_by_kab[kab]['lolos'] += 1
            else:
                verif_by_kab[kab]['tidak_lolos'] += 1
                
        sk_data = sk_by_stage_kab.get(stage_id, {})
        
        kab_data = {}
        total_sk_sudah = 0
        total_sk_belum = 0
        for kab in all_kabupaten:
            alokasi = alokasi_by_kab.get(kab, 0)
            lolos = verif_by_kab.get(kab, {}).get('lolos', 0)
            tidak_lolos = verif_by_kab.get(kab, {}).get('tidak_lolos', 0)
            verifikasi = lolos + tidak_lolos
            belum = max(0, alokasi - verifikasi)
            
            sk_sudah = sk_data.get(kab, 0)
            sk_belum = max(0, lolos - sk_sudah)
            total_sk_sudah += sk_sudah
            total_sk_belum += sk_belum
            
            kab_data[kab] = {
                "alokasi": alokasi,
                "verifikasi": verifikasi,
                "lolos": lolos,
                "tidak_lolos": tidak_lolos,
                "belum": belum,
                "sk_dirjen_sudah": sk_sudah,
                "sk_dirjen_belum": sk_belum
            }
            
        stages_data.append({
            "stage_name": stage_name,
            "data": kab_data,
            "totals": {
                "alokasi": sum(alokasi_by_kab.values()),
                "lolos": sum(x['lolos'] for x in verif_by_kab.values()),
                "tidak_lolos": sum(x['tidak_lolos'] for x in verif_by_kab.values()),
                "sk_dirjen_sudah": total_sk_sudah,
                "sk_dirjen_belum": total_sk_belum
            }
        })
    
    conn.close()

    murni_stages = [s for s in stages_data if "pengganti" not in s['stage_name'].lower()]
    pengganti_stages = [s for s in stages_data if "pengganti" in s['stage_name'].lower()]

    wb = openpyxl.Workbook()

    font_title = Font(name='Segoe UI', size=15, bold=True, color='1F4E78')
    font_subtitle = Font(name='Segoe UI', size=10, italic=True, color='595959')
    font_header = Font(name='Segoe UI', size=9, bold=True, color='FFFFFF')
    font_sub_header = Font(name='Segoe UI', size=9, bold=True, color='2C3E50')
    font_data = Font(name='Segoe UI', size=9)
    font_total = Font(name='Segoe UI', size=9, bold=True)

    fill_title_group = PatternFill(start_color='1F4E78', end_color='1F4E78', fill_type='solid')
    fill_stage_group = PatternFill(start_color='293241', end_color='293241', fill_type='solid')
    fill_total_row = PatternFill(start_color='EAF0F6', end_color='EAF0F6', fill_type='solid')
    fill_summary_col = PatternFill(start_color='F4F8FB', end_color='F4F8FB', fill_type='solid')

    border_thin = Border(
        left=Side(style='thin', color='D3D3D3'),
        right=Side(style='thin', color='D3D3D3'),
        top=Side(style='thin', color='D3D3D3'),
        bottom=Side(style='thin', color='D3D3D3')
    )
    border_thick_right = Border(
        left=Side(style='thin', color='D3D3D3'),
        right=Side(style='medium', color='7F8C8D'),
        top=Side(style='thin', color='D3D3D3'),
        bottom=Side(style='thin', color='D3D3D3')
    )
    border_double_top = Border(
        top=Side(style='double', color='2C3E50'),
        bottom=Side(style='thin', color='D3D3D3'),
        left=Side(style='thin', color='D3D3D3'),
        right=Side(style='thin', color='D3D3D3')
    )

    fill_sk_sudah = PatternFill(start_color='D5F5E3', end_color='D5F5E3', fill_type='solid')
    fill_sk_belum = PatternFill(start_color='EAECEE', end_color='EAECEE', fill_type='solid')
    font_sk_sudah = Font(name='Segoe UI', size=9, bold=True, color='1E8449')
    font_sk_belum = Font(name='Segoe UI', size=9, bold=True, color='7F8C8D')

    align_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    align_left = Alignment(horizontal='left', vertical='center')
    sub_headers = ["ALOKASI", "VERIFIKASI", "LOLOS", "TIDAK LOLOS", "BELUM", "SUDAH", "BELUM"]

    def build_worksheet(ws, title, subtitle, group_stages):
        ws['A1'] = title
        ws['A1'].font = font_title
        ws['A2'] = subtitle
        ws['A2'].font = font_subtitle

        ws.cell(row=4, column=1, value="No").font = font_header
        ws.cell(row=4, column=1).fill = fill_stage_group
        ws.cell(row=4, column=1).alignment = align_center
        ws.merge_cells(start_row=4, start_column=1, end_row=5, end_column=1)

        ws.cell(row=4, column=2, value="Kabupaten / Kota").font = font_header
        ws.cell(row=4, column=2).fill = fill_stage_group
        ws.cell(row=4, column=2).alignment = align_center
        ws.merge_cells(start_row=4, start_column=2, end_row=5, end_column=2)

        ws.cell(row=4, column=3, value="REKAP TOTAL").font = font_header
        ws.cell(row=4, column=3).fill = fill_title_group
        ws.cell(row=4, column=3).alignment = align_center
        ws.merge_cells(start_row=4, start_column=3, end_row=4, end_column=9)

        col_idx = 10
        for stage in group_stages:
            ws.cell(row=4, column=col_idx, value=stage['stage_name'].upper()).font = font_header
            ws.cell(row=4, column=col_idx).fill = fill_stage_group
            ws.cell(row=4, column=col_idx).alignment = align_center
            ws.merge_cells(start_row=4, start_column=col_idx, end_row=4, end_column=col_idx+6)
            col_idx += 7

        for i, sh in enumerate(sub_headers):
            cell = ws.cell(row=5, column=3+i, value=sh)
            cell.font = font_sub_header
            cell.alignment = align_center
            if sh == "SUDAH":
                cell.fill = fill_sk_sudah
                cell.font = font_sk_sudah
            elif sh == "BELUM" and i == 6:
                cell.fill = fill_sk_belum
                cell.font = font_sk_belum
            else:
                cell.fill = PatternFill(start_color='D4E2EE', end_color='D4E2EE', fill_type='solid')

        col_idx = 10
        for stage in group_stages:
            for i, sh in enumerate(sub_headers):
                cell = ws.cell(row=5, column=col_idx+i, value=sh)
                cell.font = font_sub_header
                cell.alignment = align_center
                if sh == "SUDAH":
                    cell.fill = fill_sk_sudah
                    cell.font = font_sk_sudah
                elif sh == "BELUM" and i == 6:
                    cell.fill = fill_sk_belum
                    cell.font = font_sk_belum
                else:
                    cell.fill = PatternFill(start_color='E8F0F8', end_color='E8F0F8', fill_type='solid')
            col_idx += 7

        row_idx = 6
        for idx, kab in enumerate(all_kabupaten):
            ws.cell(row=row_idx, column=1, value=idx+1).alignment = align_center
            ws.cell(row=row_idx, column=1).font = font_data
            ws.cell(row=row_idx, column=2, value=kab).font = font_total

            sum_a = sum(s['data'].get(kab, {}).get('alokasi', 0) for s in group_stages)
            sum_v = sum(s['data'].get(kab, {}).get('verifikasi', 0) for s in group_stages)
            sum_l = sum(s['data'].get(kab, {}).get('lolos', 0) for s in group_stages)
            sum_tl = sum(s['data'].get(kab, {}).get('tidak_lolos', 0) for s in group_stages)
            sum_b = sum(s['data'].get(kab, {}).get('belum', 0) for s in group_stages)
            sum_sks = sum(s['data'].get(kab, {}).get('sk_dirjen_sudah', 0) for s in group_stages)
            sum_skb = sum(s['data'].get(kab, {}).get('sk_dirjen_belum', 0) for s in group_stages)

            ws.cell(row=row_idx, column=3, value=sum_a).font = font_total
            ws.cell(row=row_idx, column=4, value=sum_v).font = font_total
            ws.cell(row=row_idx, column=5, value=sum_l).font = font_total
            ws.cell(row=row_idx, column=6, value=sum_tl).font = font_total
            ws.cell(row=row_idx, column=7, value=sum_b).font = font_total
            ws.cell(row=row_idx, column=8, value=sum_sks).font = font_sk_sudah
            ws.cell(row=row_idx, column=9, value=sum_skb).font = font_sk_belum

            for c in range(3, 10):
                cell = ws.cell(row=row_idx, column=c)
                if c == 8:
                    cell.fill = fill_sk_sudah
                elif c == 9:
                    cell.fill = fill_sk_belum
                else:
                    cell.fill = fill_summary_col
                cell.alignment = align_center
                if c == 9:
                    cell.border = border_thick_right
                else:
                    cell.border = border_thin

            col_idx = 10
            for stage in group_stages:
                kd = stage["data"].get(kab, {"alokasi": 0, "verifikasi": 0, "lolos": 0, "tidak_lolos": 0, "belum": 0, "sk_dirjen_sudah": 0, "sk_dirjen_belum": 0})
                ws.cell(row=row_idx, column=col_idx, value=kd["alokasi"] or "-").alignment = align_center
                ws.cell(row=row_idx, column=col_idx+1, value=kd["verifikasi"] or "-").alignment = align_center
                ws.cell(row=row_idx, column=col_idx+2, value=kd["lolos"] or "-").alignment = align_center
                ws.cell(row=row_idx, column=col_idx+3, value=kd["tidak_lolos"] or "-").alignment = align_center
                ws.cell(row=row_idx, column=col_idx+4, value=kd["belum"] or "-").alignment = align_center
                ws.cell(row=row_idx, column=col_idx+5, value=kd.get("sk_dirjen_sudah", 0) or "-").alignment = align_center
                ws.cell(row=row_idx, column=col_idx+6, value=kd.get("sk_dirjen_belum", 0) or "-").alignment = align_center

                for offset in range(7):
                    cell = ws.cell(row=row_idx, column=col_idx+offset)
                    cell.font = font_data
                    if offset == 5:
                        cell.fill = fill_sk_sudah
                        cell.font = font_sk_sudah
                    elif offset == 6:
                        cell.fill = fill_sk_belum
                        cell.font = font_sk_belum
                    if offset == 6:
                        cell.border = border_thick_right
                    else:
                        cell.border = border_thin
                col_idx += 7

            ws.cell(row=row_idx, column=1).border = border_thin
            ws.cell(row=row_idx, column=2).border = border_thin
            row_idx += 1

        ws.cell(row=row_idx, column=2, value="TOTAL").font = font_total
        ws.cell(row=row_idx, column=2).alignment = align_left
        ws.cell(row=row_idx, column=2).fill = fill_total_row
        ws.cell(row=row_idx, column=2).border = border_double_top
        ws.cell(row=row_idx, column=1).fill = fill_total_row
        ws.cell(row=row_idx, column=1).border = border_double_top

        tot_a = sum(sum(s['data'].get(kab, {}).get('alokasi', 0) for s in group_stages) for kab in all_kabupaten)
        tot_v = sum(sum(s['data'].get(kab, {}).get('verifikasi', 0) for s in group_stages) for kab in all_kabupaten)
        tot_l = sum(sum(s['data'].get(kab, {}).get('lolos', 0) for s in group_stages) for kab in all_kabupaten)
        tot_tl = sum(sum(s['data'].get(kab, {}).get('tidak_lolos', 0) for s in group_stages) for kab in all_kabupaten)
        tot_b = sum(sum(s['data'].get(kab, {}).get('belum', 0) for s in group_stages) for kab in all_kabupaten)
        tot_sks = sum(sum(s['data'].get(kab, {}).get('sk_dirjen_sudah', 0) for s in group_stages) for kab in all_kabupaten)
        tot_skb = sum(sum(s['data'].get(kab, {}).get('sk_dirjen_belum', 0) for s in group_stages) for kab in all_kabupaten)

        total_vals = [tot_a, tot_v, tot_l, tot_tl, tot_b, tot_sks, tot_skb]
        for i, val in enumerate(total_vals):
            cell = ws.cell(row=row_idx, column=3+i, value=val)
            cell.font = font_total
            cell.alignment = align_center
            cell.fill = fill_total_row
            if i == 5:
                cell.fill = fill_sk_sudah
                cell.font = Font(name='Segoe UI', size=9, bold=True, color='1E8449')
            elif i == 6:
                cell.fill = fill_sk_belum
                cell.font = Font(name='Segoe UI', size=9, bold=True, color='7F8C8D')
            if i == 6:
                cell.border = Border(top=Side(style='double', color='2C3E50'), bottom=Side(style='thin', color='D3D3D3'), right=Side(style='medium', color='7F8C8D'))
            else:
                cell.border = border_double_top

        col_idx = 10
        for stage in group_stages:
            st_a = sum(kd["alokasi"] for kd in stage["data"].values())
            st_v = sum(kd["verifikasi"] for kd in stage["data"].values())
            st_l = sum(kd["lolos"] for kd in stage["data"].values())
            st_tl = sum(kd["tidak_lolos"] for kd in stage["data"].values())
            st_b = sum(kd["belum"] for kd in stage["data"].values())
            st_sks = sum(kd.get("sk_dirjen_sudah", 0) for kd in stage["data"].values())
            st_skb = sum(kd.get("sk_dirjen_belum", 0) for kd in stage["data"].values())

            vals = [st_a, st_v, st_l, st_tl, st_b, st_sks, st_skb]
            for i, val in enumerate(vals):
                cell = ws.cell(row=row_idx, column=col_idx+i, value=val)
                cell.font = font_total
                cell.alignment = align_center
                cell.fill = fill_total_row
                if i == 5:
                    cell.fill = fill_sk_sudah
                    cell.font = Font(name='Segoe UI', size=9, bold=True, color='1E8449')
                elif i == 6:
                    cell.fill = fill_sk_belum
                    cell.font = Font(name='Segoe UI', size=9, bold=True, color='7F8C8D')
                if i == 6:
                    cell.border = Border(top=Side(style='double', color='2C3E50'), bottom=Side(style='thin', color='D3D3D3'), right=Side(style='medium', color='7F8C8D'))
                else:
                    cell.border = border_double_top
            col_idx += 7

        for col in ws.columns:
            max_len = 0
            col_letter = get_column_letter(col[0].column)
            for cell in col:
                val = str(cell.value or '')
                if cell.row in [1, 2]:
                    continue
                if len(val) > max_len:
                    max_len = len(val)
            ws.column_dimensions[col_letter].width = max(max_len + 4, 10)
        ws.column_dimensions['B'].width = 28
        ws.freeze_panes = 'C6'

    filter_note = f" (Filter Pengusul: {', '.join(pengusul_list)})" if pengusul_list else ""
    ws_murni = wb.active
    ws_murni.title = "Rekap Invers Murni"
    build_worksheet(ws_murni, "REKAPITULASI INVERS MURNI", "Sistem Verifikasi Perumahan Swadaya — Tahap Invers Murni" + filter_note, murni_stages)

    ws_pengganti = wb.create_sheet("Rekap Invers Pengganti")
    build_worksheet(ws_pengganti, "REKAPITULASI INVERS PENGGANTI", "Sistem Verifikasi Perumahan Swadaya — Tahap Invers Pengganti" + filter_note, pengganti_stages)

    ws_keseluruhan = wb.create_sheet("Rekap Keseluruhan")
    build_worksheet(ws_keseluruhan, "REKAPITULASI KESELURUHAN INVERS", "Sistem Verifikasi Perumahan Swadaya — Semua Tahap" + filter_note, stages_data)

    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=REKAP_KESELURUHAN_INVERS.xlsx"}
    )

# --- REKAP KESELURUHAN (All stages, per kabupaten) ---
REKAP_CACHE = {}

@app.get("/api/rekap-keseluruhan")
def get_rekap_keseluruhan(published_only: int = 0, pengusul: str = "", province_id: int = 1):
    pengusul_list = [p.strip() for p in pengusul.split(",") if p.strip()]
    cache_key = f"rekap_{province_id}_{published_only}_{','.join(sorted(pengusul_list))}"
    now = time.time()
    if cache_key in REKAP_CACHE:
        ts, cached_data = REKAP_CACHE[cache_key]
        if now - ts < 10:
            return cached_data

    import re
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get stages for current province
    if not province_id or province_id == 1:
        prov_stage_sql = "WHERE (province_id = 1 OR province_id IS NULL OR province_id = 0)"
        prov_stage_params = ()
    else:
        prov_stage_sql = "WHERE province_id = ?"
        prov_stage_params = (province_id,)

    cursor.execute(f"SELECT id, name FROM invers_stages {prov_stage_sql} ORDER BY id ASC", prov_stage_params)
    all_stages = [dict(r) for r in cursor.fetchall()]
    def get_stage_num(s):
        match = re.search(r'\d+', s['name'])
        return int(match.group()) if match else 999
    all_stages = sorted(all_stages, key=get_stage_num)
    
    # Distinct list of pengusul options for the filter UI (from active invers records in current province)
    cursor.execute(f"""
        SELECT DISTINCT UPPER(TRIM(COALESCE(ir.pengusul, ''))) as pengusul
        FROM invers_records ir
        JOIN invers_revisions irv ON ir.revision_id = irv.id
        JOIN invers_stages s ON irv.stage_id = s.id
        {prov_stage_sql} AND irv.is_active = 1 AND TRIM(COALESCE(ir.pengusul, '')) != ''
        ORDER BY pengusul ASC
    """, prov_stage_params)
    pengusul_options = [r['pengusul'] for r in cursor.fetchall()]
    
    published_filter = "AND vb.is_published = 1" if published_only else ""
    
    # Get the full list of unique kabupaten across current province stages
    cursor.execute(f"""
        SELECT DISTINCT UPPER(TRIM(COALESCE(ir.kabupaten_kota, ''))) as kab 
        FROM invers_records ir
        JOIN invers_revisions irv ON ir.revision_id = irv.id
        JOIN invers_stages s ON irv.stage_id = s.id
        {prov_stage_sql} AND irv.is_active = 1 AND TRIM(COALESCE(ir.kabupaten_kota, '')) != ''
        UNION
        SELECT DISTINCT UPPER(TRIM(COALESCE(NULLIF(TRIM(vr.kabupaten_kota), ''), NULLIF(TRIM(vb.kabupaten), '')))) as kab
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        JOIN invers_stages s ON vb.stage_id = s.id
        LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
        {prov_stage_sql} AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL) AND TRIM(COALESCE(NULLIF(TRIM(vr.kabupaten_kota), ''), NULLIF(TRIM(vb.kabupaten), ''))) != ''
        {published_filter}
        ORDER BY kab ASC
    """, (*prov_stage_params, *prov_stage_params))
    all_kabupaten = [r['kab'] for r in cursor.fetchall()]
    
    # Pre-fetch SK Dirjen match counts per stage per kabupaten
    sk_where_extra = ""
    sk_params = ()
    if pengusul_list:
        ph = ",".join("?" for _ in pengusul_list)
        sk_where_extra = f"AND UPPER(COALESCE(ir_p.pengusul, '')) IN ({ph})"
        sk_params = tuple(p.upper() for p in pengusul_list)
    cursor.execute(f"""
        SELECT 
            m.verified_stage_id as stage_id,
            UPPER(TRIM(COALESCE(NULLIF(TRIM(vr.kabupaten_kota), ''), NULLIF(TRIM(vb.kabupaten), '')))) as kab,
            COUNT(*) as cnt
        FROM sk_dirjen_matches m
        JOIN verified_records vr ON vr.id = m.verified_record_id
        JOIN verified_batches vb ON vr.batch_id = vb.id
        LEFT JOIN (
            SELECT DISTINCT no_ktp, pengusul FROM invers_records
            WHERE pengusul IS NOT NULL AND pengusul != ''
        ) ir_p ON ir_p.no_ktp = vr.no_ktp
        WHERE m.verified_stage_id IS NOT NULL
        AND m.verified_record_id IS NOT NULL
        AND (m.match_type = 'PERFECT' 
            OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED')
            OR m.match_type = 'MANUAL_PAIR')
        {sk_where_extra}
        GROUP BY m.verified_stage_id, kab
    """, sk_params)
    sk_by_stage_kab = {}
    for row in cursor.fetchall():
        sid = row['stage_id']
        kab = row['kab']
        if sid not in sk_by_stage_kab:
            sk_by_stage_kab[sid] = {}
        sk_by_stage_kab[sid][kab] = row['cnt']

    # SQL Aggregation for ALL stages (Invers Alokasi)
    alokasi_where = ""
    alokasi_params = ()
    if pengusul_list:
        ph = ",".join("?" for _ in pengusul_list)
        alokasi_where = f"AND UPPER(COALESCE(ir.pengusul, '')) IN ({ph})"
        alokasi_params = tuple(p.upper() for p in pengusul_list)
    cursor.execute(f"""
        SELECT irv.stage_id, UPPER(TRIM(COALESCE(ir.kabupaten_kota, ''))) as kab, COUNT(*) as cnt
        FROM invers_records ir
        JOIN invers_revisions irv ON ir.revision_id = irv.id
        WHERE irv.is_active = 1 AND TRIM(COALESCE(ir.kabupaten_kota, '')) != ''
        {alokasi_where}
        GROUP BY irv.stage_id, kab
    """, alokasi_params)
    alokasi_map = {}
    for row in cursor.fetchall():
        sid, kab, cnt = row['stage_id'], row['kab'], row['cnt']
        if sid not in alokasi_map: alokasi_map[sid] = {}
        alokasi_map[sid][kab] = cnt

    # SQL Aggregation for ALL stages (Verified Records Lolos & Tidak Lolos)
    verif_where = ""
    verif_params = ()
    if pengusul_list:
        ph = ",".join("?" for _ in pengusul_list)
        verif_where = f"AND UPPER(COALESCE(ir_p.pengusul, '')) IN ({ph})"
        verif_params = tuple(p.upper() for p in pengusul_list)
    cursor.execute(f"""
        SELECT vb.stage_id, UPPER(TRIM(COALESCE(NULLIF(TRIM(vr.kabupaten_kota), ''), NULLIF(TRIM(vb.kabupaten), '')))) as kab, vr.status, COUNT(*) as cnt
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
        LEFT JOIN (
            SELECT DISTINCT no_ktp, pengusul FROM invers_records
            WHERE pengusul IS NOT NULL AND pengusul != ''
        ) ir_p ON ir_p.no_ktp = vr.no_ktp
        WHERE (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL) AND TRIM(COALESCE(NULLIF(TRIM(vr.kabupaten_kota), ''), NULLIF(TRIM(vb.kabupaten), ''))) != ''
        {verif_where}
        {published_filter}
        GROUP BY vb.stage_id, kab, vr.status
    """, verif_params)
    verif_map = {}
    for row in cursor.fetchall():
        sid, kab, st, cnt = row['stage_id'], row['kab'], row['status'], row['cnt']
        if sid not in verif_map: verif_map[sid] = {}
        if kab not in verif_map[sid]: verif_map[sid][kab] = {"lolos": 0, "tidak_lolos": 0}
        if st == 'LOLOS':
            verif_map[sid][kab]["lolos"] += cnt
        else:
            verif_map[sid][kab]["tidak_lolos"] += cnt

    conn.close()
    
    stages_data = []
    
    for stage in all_stages:
        stage_id = stage['id']
        stage_name = stage['name']
        
        stage_alokasi = alokasi_map.get(stage_id, {})
        stage_verif = verif_map.get(stage_id, {})
        sk_data = sk_by_stage_kab.get(stage_id, {})
        
        kab_data = []
        total_alokasi = 0
        total_verifikasi = 0
        total_lolos = 0
        total_tidak_lolos = 0
        total_belum = 0
        total_sk_sudah = 0
        total_sk_belum = 0
        
        for kab in all_kabupaten:
            alokasi = stage_alokasi.get(kab, 0)
            v_stat = stage_verif.get(kab, {"lolos": 0, "tidak_lolos": 0})
            lolos = v_stat.get('lolos', 0)
            tidak_lolos = v_stat.get('tidak_lolos', 0)
            verifikasi = lolos + tidak_lolos
            belum = max(0, alokasi - verifikasi)
            
            sk_sudah = sk_data.get(kab, 0)
            sk_belum = max(0, lolos - sk_sudah)
            
            total_alokasi += alokasi
            total_verifikasi += verifikasi
            total_lolos += lolos
            total_tidak_lolos += tidak_lolos
            total_belum += belum
            total_sk_sudah += sk_sudah
            total_sk_belum += sk_belum
            
            kab_data.append({
                "kabupaten": kab,
                "alokasi": alokasi,
                "verifikasi": verifikasi,
                "lolos": lolos,
                "tidak_lolos": tidak_lolos,
                "belum_verifikasi": belum,
                "sk_dirjen_sudah": sk_sudah,
                "sk_dirjen_belum": sk_belum
            })
        
        stage_type = "pengganti" if "pengganti" in stage_name.lower() else "murni"
        stages_data.append({
            "stage_id": stage_id,
            "stage_name": stage_name,
            "stage_type": stage_type,
            "kabupaten_data": kab_data,
            "totals": {
                "alokasi": total_alokasi,
                "verifikasi": total_verifikasi,
                "lolos": total_lolos,
                "tidak_lolos": total_tidak_lolos,
                "belum_verifikasi": total_belum,
                "sk_dirjen_sudah": total_sk_sudah,
                "sk_dirjen_belum": total_sk_belum
            }
        })
    
    result = {
        "all_kabupaten": all_kabupaten,
        "stages": stages_data,
        "pengusul_options": pengusul_options
    }
    REKAP_CACHE[cache_key] = (now, result)
    return result

# --- REKAP BATCH BERITA ACARA ---
@app.get("/api/rekap-batch-ba")
def get_rekap_batch_ba(published_only: int = 1, province_id: int = 1):
    import re
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # 1. Get stages for current province
    if not province_id or province_id == 1:
        prov_stage_sql = "WHERE (province_id = 1 OR province_id IS NULL OR province_id = 0)"
        prov_stage_params = ()
    else:
        prov_stage_sql = "WHERE province_id = ?"
        prov_stage_params = (province_id,)

    cursor.execute(f"SELECT id, name FROM invers_stages {prov_stage_sql} ORDER BY id ASC", prov_stage_params)
    all_stages = [dict(r) for r in cursor.fetchall()]
    def get_stage_num(s):
        match = re.search(r'\d+', s['name'])
        return int(match.group()) if match else 999
    all_stages = sorted(all_stages, key=get_stage_num)
    
    # 2. Get list of kabupaten/kota in current province
    published_filter = "AND vb.is_published = 1" if published_only else ""
    cursor.execute(f"""
        SELECT DISTINCT UPPER(TRIM(COALESCE(ir.kabupaten_kota, ''))) as kab 
        FROM invers_records ir
        JOIN invers_revisions irv ON ir.revision_id = irv.id
        JOIN invers_stages s ON irv.stage_id = s.id
        {prov_stage_sql} AND irv.is_active = 1 AND TRIM(COALESCE(ir.kabupaten_kota, '')) != ''
        UNION
        SELECT DISTINCT UPPER(TRIM(COALESCE(vr.kabupaten_kota, ''))) as kab
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        JOIN invers_stages s ON vb.stage_id = s.id
        LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
        {prov_stage_sql} AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL) AND TRIM(COALESCE(vr.kabupaten_kota, '')) != ''
        {published_filter}
        ORDER BY kab ASC
    """, (*prov_stage_params, *prov_stage_params))
    all_kabupaten = [r['kab'] for r in cursor.fetchall()]
    
    # 3. For each stage, get its batches
    stage_batch_filter = "AND is_published = 1" if published_only else ""
    stages_data = []
    
    for stage in all_stages:
        stage_id = stage['id']
        stage_name = stage['name']
        
        cursor.execute(f"""
            SELECT id, name, is_published, nomor_ba, tanggal_ba, sort_order 
            FROM verified_batches 
            WHERE stage_id = ? {stage_batch_filter}
            ORDER BY sort_order ASC, uploaded_at ASC, id ASC
        """, (stage_id,))
        batches = [dict(r) for r in cursor.fetchall()]
        
        if not batches:
            continue
            
        batches_data = []
        for batch in batches:
            batch_id = batch['id']
            batch_name = batch['name']
            
            cursor.execute("""
                SELECT UPPER(TRIM(COALESCE(vr.kabupaten_kota, ''))) as kab,
                       SUM(CASE WHEN vr.status = 'LOLOS' THEN 1 ELSE 0 END) as lolos,
                       SUM(CASE WHEN vr.status = 'TIDAK LOLOS' THEN 1 ELSE 0 END) as tidak_lolos,
                       SUM(CASE WHEN vr.status = 'LOLOS' AND sk.verified_record_id IS NOT NULL THEN 1 ELSE 0 END) as sk_sudah
                FROM verified_records vr
                LEFT JOIN (
                    SELECT DISTINCT verified_record_id 
                    FROM sk_dirjen_matches 
                    WHERE verified_record_id IS NOT NULL
                      AND (match_type = 'PERFECT' 
                           OR (match_type = 'NEEDS_APPROVAL' AND override_status = 'APPROVED')
                           OR match_type = 'MANUAL_PAIR')
                ) sk ON sk.verified_record_id = vr.id
                WHERE vr.batch_id = ? AND TRIM(COALESCE(vr.kabupaten_kota, '')) != ''
                GROUP BY kab
            """, (batch_id,))
            
            stats_by_kab = {}
            for row in cursor.fetchall():
                lolos = row['lolos']
                tidak_lolos = row['tidak_lolos']
                sk_sudah = row['sk_sudah']
                sk_belum = max(0, lolos - sk_sudah)
                stats_by_kab[row['kab']] = {
                    "lolos": lolos,
                    "tidak_lolos": tidak_lolos,
                    "verifikasi": lolos + tidak_lolos,
                    "sk_sudah": sk_sudah,
                    "sk_belum": sk_belum
                }
            
            kab_data = []
            total_verifikasi = 0
            total_lolos = 0
            total_tidak_lolos = 0
            total_sk_sudah = 0
            total_sk_belum = 0
            
            for kab in all_kabupaten:
                stats = stats_by_kab.get(kab, {"lolos": 0, "tidak_lolos": 0, "verifikasi": 0, "sk_sudah": 0, "sk_belum": 0})
                kab_data.append({
                    "kabupaten": kab,
                    "verifikasi": stats["verifikasi"],
                    "lolos": stats["lolos"],
                    "tidak_lolos": stats["tidak_lolos"],
                    "sk_sudah": stats["sk_sudah"],
                    "sk_belum": stats["sk_belum"]
                })
                total_verifikasi += stats["verifikasi"]
                total_lolos += stats["lolos"]
                total_tidak_lolos += stats["tidak_lolos"]
                total_sk_sudah += stats["sk_sudah"]
                total_sk_belum += stats["sk_belum"]
                
            batches_data.append({
                "batch_id": batch_id,
                "batch_name": batch_name,
                "is_published": batch["is_published"],
                "nomor_ba": batch.get("nomor_ba"),
                "tanggal_ba": batch.get("tanggal_ba"),
                "kabupaten_data": kab_data,
                "totals": {
                    "verifikasi": total_verifikasi,
                    "lolos": total_lolos,
                    "tidak_lolos": total_tidak_lolos,
                    "sk_sudah": total_sk_sudah,
                    "sk_belum": total_sk_belum
                }
            })
            
        stage_type = "pengganti" if "pengganti" in stage_name.lower() else "murni"
        stages_data.append({
            "stage_id": stage_id,
            "stage_name": stage_name,
            "stage_type": stage_type,
            "batches": batches_data
        })
        
    conn.close()
    return {
        "all_kabupaten": all_kabupaten,
        "stages": stages_data
    }

@app.get("/api/rekap-batch-ba/export")
def export_rekap_batch_ba(published_only: int = 1):
    import re
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # 1. Get all stages
    cursor.execute("SELECT id, name FROM invers_stages ORDER BY id ASC")
    all_stages = [dict(r) for r in cursor.fetchall()]
    def get_stage_num(s):
        match = re.search(r'\d+', s['name'])
        return int(match.group()) if match else 999
    all_stages = sorted(all_stages, key=get_stage_num)
    
    # 2. Get list of kabupaten/kota
    published_filter = "AND vb.is_published = 1" if published_only else ""
    cursor.execute(f"""
        SELECT DISTINCT UPPER(TRIM(COALESCE(kabupaten_kota, ''))) as kab 
        FROM invers_records ir
        JOIN invers_revisions irv ON ir.revision_id = irv.id
        WHERE irv.is_active = 1 AND TRIM(COALESCE(kabupaten_kota, '')) != ''
        UNION
        SELECT DISTINCT UPPER(TRIM(COALESCE(vr.kabupaten_kota, ''))) as kab
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
        WHERE (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL) AND TRIM(COALESCE(vr.kabupaten_kota, '')) != ''
        {published_filter}
        ORDER BY kab ASC
    """)
    all_kabupaten = [r['kab'] for r in cursor.fetchall()]
    
    stage_batch_filter = "AND is_published = 1" if published_only else ""
    stages_data = []
    
    for stage in all_stages:
        stage_id = stage['id']
        stage_name = stage['name']
        
        cursor.execute(f"""
            SELECT id, name, nomor_ba, tanggal_ba 
            FROM verified_batches 
            WHERE stage_id = ? {stage_batch_filter}
            ORDER BY sort_order ASC, uploaded_at ASC, id ASC
        """, (stage_id,))
        batches = [dict(r) for r in cursor.fetchall()]
        
        if not batches:
            continue
            
        batches_data = []
        for batch in batches:
            batch_id = batch['id']
            batch_name = batch['name']
            
            cursor.execute("""
                SELECT UPPER(TRIM(COALESCE(vr.kabupaten_kota, ''))) as kab,
                       SUM(CASE WHEN vr.status = 'LOLOS' THEN 1 ELSE 0 END) as lolos,
                       SUM(CASE WHEN vr.status = 'TIDAK LOLOS' THEN 1 ELSE 0 END) as tidak_lolos,
                       SUM(CASE WHEN vr.status = 'LOLOS' AND sk.verified_record_id IS NOT NULL THEN 1 ELSE 0 END) as sk_sudah
                FROM verified_records vr
                LEFT JOIN (
                    SELECT DISTINCT verified_record_id 
                    FROM sk_dirjen_matches 
                    WHERE verified_record_id IS NOT NULL
                      AND (match_type = 'PERFECT' 
                           OR (match_type = 'NEEDS_APPROVAL' AND override_status = 'APPROVED')
                           OR match_type = 'MANUAL_PAIR')
                ) sk ON sk.verified_record_id = vr.id
                WHERE vr.batch_id = ? AND TRIM(COALESCE(vr.kabupaten_kota, '')) != ''
                GROUP BY kab
            """, (batch_id,))
            
            stats_by_kab = {}
            for row in cursor.fetchall():
                lolos = row['lolos']
                tidak_lolos = row['tidak_lolos']
                sk_sudah = row['sk_sudah']
                sk_belum = max(0, lolos - sk_sudah)
                stats_by_kab[row['kab']] = {
                    "lolos": lolos,
                    "tidak_lolos": tidak_lolos,
                    "verifikasi": lolos + tidak_lolos,
                    "sk_sudah": sk_sudah,
                    "sk_belum": sk_belum
                }
            
            kab_data = {}
            total_verifikasi = 0
            total_lolos = 0
            total_tidak_lolos = 0
            total_sk_sudah = 0
            total_sk_belum = 0
            for kab in all_kabupaten:
                stats = stats_by_kab.get(kab, {"lolos": 0, "tidak_lolos": 0, "verifikasi": 0, "sk_sudah": 0, "sk_belum": 0})
                kab_data[kab] = {
                    "verifikasi": stats["verifikasi"],
                    "lolos": stats["lolos"],
                    "tidak_lolos": stats["tidak_lolos"],
                    "sk_sudah": stats["sk_sudah"],
                    "sk_belum": stats["sk_belum"]
                }
                total_verifikasi += stats["verifikasi"]
                total_lolos += stats["lolos"]
                total_tidak_lolos += stats["tidak_lolos"]
                total_sk_sudah += stats["sk_sudah"]
                total_sk_belum += stats["sk_belum"]
                
            batches_data.append({
                "batch_id": batch_id,
                "batch_name": batch_name,
                "nomor_ba": batch.get("nomor_ba"),
                "tanggal_ba": batch.get("tanggal_ba"),
                "data": kab_data,
                "totals": {
                    "verifikasi": total_verifikasi,
                    "lolos": total_lolos,
                    "tidak_lolos": total_tidak_lolos,
                    "sk_sudah": total_sk_sudah,
                    "sk_belum": total_sk_belum
                }
            })
            
        stage_type = "pengganti" if "pengganti" in stage_name.lower() else "murni"
        stages_data.append({
            "stage_name": stage_name,
            "stage_type": stage_type,
            "batches": batches_data
        })
        
    conn.close()
    
    murni_stages = [s for s in stages_data if s['stage_type'] == 'murni']
    pengganti_stages = [s for s in stages_data if s['stage_type'] == 'pengganti']
    
    wb = openpyxl.Workbook()
    
    font_title = Font(name='Segoe UI', size=15, bold=True, color='1F4E78')
    font_subtitle = Font(name='Segoe UI', size=10, italic=True, color='595959')
    font_header_corner = Font(name='Segoe UI', size=9, bold=True, color='2C3E50')
    font_header_l1 = Font(name='Segoe UI', size=9, bold=True, color='FFFFFF')
    font_header_l2 = Font(name='Segoe UI', size=8, bold=True, color='2C3E50')
    font_header_l3 = Font(name='Segoe UI', size=8, bold=True, color='595959')
    
    font_data = Font(name='Segoe UI', size=9)
    font_total = Font(name='Segoe UI', size=9, bold=True)
    
    fill_corner = PatternFill(start_color='EAF0F6', end_color='EAF0F6', fill_type='solid')
    fill_l1 = PatternFill(start_color='1F4E78', end_color='1F4E78', fill_type='solid')
    fill_l2 = PatternFill(start_color='FCE4D6', end_color='FCE4D6', fill_type='solid')
    fill_l3 = PatternFill(start_color='F2F4F4', end_color='F2F4F4', fill_type='solid')
    fill_total_row = PatternFill(start_color='EAF0F6', end_color='EAF0F6', fill_type='solid')
    
    fill_lolos_cell = PatternFill(start_color='E8F5E9', end_color='E8F5E9', fill_type='solid')
    fill_tidak_lolos_cell = PatternFill(start_color='FDEDEC', end_color='FDEDEC', fill_type='solid')
    fill_sk_sudah_cell = PatternFill(start_color='E0F2FE', end_color='E0F2FE', fill_type='solid')
    fill_sk_belum_cell = PatternFill(start_color='FEF3C7', end_color='FEF3C7', fill_type='solid')

    font_lolos_cell = Font(name='Segoe UI', size=9, color='2E7D32')
    font_tidak_lolos_cell = Font(name='Segoe UI', size=9, color='C0392B')
    font_sk_sudah_cell = Font(name='Segoe UI', size=9, color='0284C7', bold=True)
    font_sk_belum_cell = Font(name='Segoe UI', size=9, color='D97706')
    
    border_thin = Border(
        left=Side(style='thin', color='D3D3D3'),
        right=Side(style='thin', color='D3D3D3'),
        top=Side(style='thin', color='D3D3D3'),
        bottom=Side(style='thin', color='D3D3D3')
    )
    border_double_top = Border(
        top=Side(style='double', color='2C3E50'),
        bottom=Side(style='thin', color='D3D3D3'),
        left=Side(style='thin', color='D3D3D3'),
        right=Side(style='thin', color='D3D3D3')
    )
    
    align_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    align_left = Alignment(horizontal='left', vertical='center')
    
    def build_worksheet(ws, title, subtitle, group_stages):
        ws['A1'] = title
        ws['A1'].font = font_title
        ws['A2'] = subtitle
        ws['A2'].font = font_subtitle
        
        cell_no = ws.cell(row=4, column=1, value="No")
        cell_no.font = font_header_corner
        cell_no.fill = fill_corner
        cell_no.alignment = align_center
        ws.merge_cells(start_row=4, start_column=1, end_row=6, end_column=1)
        
        cell_kab = ws.cell(row=4, column=2, value="Kabupaten / Kota")
        cell_kab.font = font_header_corner
        cell_kab.fill = fill_corner
        cell_kab.alignment = align_center
        ws.merge_cells(start_row=4, start_column=2, end_row=6, end_column=2)
        
        col_idx = 3
        for stage in group_stages:
            num_batches = len(stage['batches'])
            if num_batches == 0:
                continue
                
            stage_width = num_batches * 5
            cell_stage = ws.cell(row=4, column=col_idx, value=stage['stage_name'].upper())
            cell_stage.font = font_header_l1
            cell_stage.fill = fill_l1
            cell_stage.alignment = align_center
            ws.merge_cells(start_row=4, start_column=col_idx, end_row=4, end_column=col_idx + stage_width - 1)
            
            b_col_idx = col_idx
            for batch in stage['batches']:
                val_str = f"{batch['batch_name'].upper()}\nNo: {batch.get('nomor_ba') or '—'}\nTgl: {batch.get('tanggal_ba') or '—'}"
                cell_batch = ws.cell(row=5, column=b_col_idx, value=val_str)
                cell_batch.font = font_header_l2
                cell_batch.fill = fill_l2
                cell_batch.alignment = align_center
                ws.merge_cells(start_row=5, start_column=b_col_idx, end_row=5, end_column=b_col_idx + 4)
                ws.row_dimensions[5].height = 42
                
                metrics = ["VERIFIKASI", "LOLOS", "TIDAK LOLOS", "SUDAH SK", "BELUM SK"]
                for i, m in enumerate(metrics):
                    cell_m = ws.cell(row=6, column=b_col_idx + i, value=m)
                    cell_m.font = font_header_l3
                    cell_m.fill = fill_l3
                    cell_m.alignment = align_center
                    
                b_col_idx += 5
            col_idx += stage_width
            
        max_col = 2
        for stage in group_stages:
            max_col += len(stage['batches']) * 5
            
        for r in range(4, 7):
            for c in range(1, max_col + 1):
                cell = ws.cell(row=r, column=c)
                cell.border = border_thin
                
        row_idx = 7
        for idx, kab in enumerate(all_kabupaten):
            ws.cell(row=row_idx, column=1, value=idx + 1).alignment = align_center
            ws.cell(row=row_idx, column=1).font = font_data
            ws.cell(row=row_idx, column=1).border = border_thin
            
            ws.cell(row=row_idx, column=2, value=kab).font = font_total
            ws.cell(row=row_idx, column=2).border = border_thin
            
            c_idx = 3
            for stage in group_stages:
                for batch in stage['batches']:
                    kd = batch['data'].get(kab, {"verifikasi": 0, "lolos": 0, "tidak_lolos": 0, "sk_sudah": 0, "sk_belum": 0})
                    val_v = kd['verifikasi']
                    val_l = kd['lolos']
                    val_tl = kd['tidak_lolos']
                    val_sks = kd['sk_sudah']
                    val_skb = kd['sk_belum']
                    
                    cell_v = ws.cell(row=row_idx, column=c_idx, value=val_v or "-")
                    cell_l = ws.cell(row=row_idx, column=c_idx + 1, value=val_l or "-")
                    cell_tl = ws.cell(row=row_idx, column=c_idx + 2, value=val_tl or "-")
                    cell_sks = ws.cell(row=row_idx, column=c_idx + 3, value=val_sks or "-")
                    cell_skb = ws.cell(row=row_idx, column=c_idx + 4, value=val_skb or "-")
                    
                    cell_v.alignment = align_center
                    cell_l.alignment = align_center
                    cell_tl.alignment = align_center
                    cell_sks.alignment = align_center
                    cell_skb.alignment = align_center
                    
                    cell_v.font = font_data
                    cell_l.font = font_lolos_cell if val_l > 0 else font_data
                    cell_tl.font = font_tidak_lolos_cell if val_tl > 0 else font_data
                    cell_sks.font = font_sk_sudah_cell if val_sks > 0 else font_data
                    cell_skb.font = font_sk_belum_cell if val_skb > 0 else font_data
                    
                    if val_l > 0:
                        cell_l.fill = fill_lolos_cell
                    if val_tl > 0:
                        cell_tl.fill = fill_tidak_lolos_cell
                    if val_sks > 0:
                        cell_sks.fill = fill_sk_sudah_cell
                    if val_skb > 0:
                        cell_skb.fill = fill_sk_belum_cell
                        
                    cell_v.border = border_thin
                    cell_l.border = border_thin
                    cell_tl.border = border_thin
                    cell_sks.border = border_thin
                    cell_skb.border = border_thin
                    c_idx += 5
            row_idx += 1
            
        cell_tot_label = ws.cell(row=row_idx, column=2, value="TOTAL")
        cell_tot_label.font = font_total
        cell_tot_label.alignment = align_left
        cell_tot_label.fill = fill_total_row
        cell_tot_label.border = border_double_top
        
        cell_tot_no = ws.cell(row=row_idx, column=1)
        cell_tot_no.fill = fill_total_row
        cell_tot_no.border = border_double_top
        
        c_idx = 3
        for stage in group_stages:
            for batch in stage['batches']:
                t = batch['totals']
                cell_tot_v = ws.cell(row=row_idx, column=c_idx, value=t['verifikasi'])
                cell_tot_l = ws.cell(row=row_idx, column=c_idx + 1, value=t['lolos'])
                cell_tot_tl = ws.cell(row=row_idx, column=c_idx + 2, value=t['tidak_lolos'])
                cell_tot_sks = ws.cell(row=row_idx, column=c_idx + 3, value=t['sk_sudah'])
                cell_tot_skb = ws.cell(row=row_idx, column=c_idx + 4, value=t['sk_belum'])
                
                cell_tot_v.font = font_total
                cell_tot_l.font = font_total
                cell_tot_tl.font = font_total
                cell_tot_sks.font = font_total
                cell_tot_skb.font = font_total
                
                cell_tot_v.alignment = align_center
                cell_tot_l.alignment = align_center
                cell_tot_tl.alignment = align_center
                cell_tot_sks.alignment = align_center
                cell_tot_skb.alignment = align_center
                
                cell_tot_v.fill = fill_total_row
                cell_tot_l.fill = fill_total_row
                cell_tot_tl.fill = fill_total_row
                cell_tot_sks.fill = fill_total_row
                cell_tot_skb.fill = fill_total_row
                
                cell_tot_v.border = border_double_top
                cell_tot_l.border = border_double_top
                cell_tot_tl.border = border_double_top
                cell_tot_sks.border = border_double_top
                cell_tot_skb.border = border_double_top
                c_idx += 5
                
        for col in ws.columns:
            col_idx = col[0].column
            col_letter = get_column_letter(col_idx)
            if col_idx == 1:
                ws.column_dimensions[col_letter].width = 6
            elif col_idx == 2:
                ws.column_dimensions[col_letter].width = 28
            else:
                ws.column_dimensions[col_letter].width = 14
        ws.freeze_panes = 'C7'
        
    ws_murni = wb.active
    ws_murni.title = "Rekap Batch Murni"
    build_worksheet(ws_murni, "REKAPITULASI BATCH INVERS MURNI", "Sistem Verifikasi Perumahan Swadaya — Laporan per Batch Murni", murni_stages)
    
    if pengganti_stages:
        ws_pengganti = wb.create_sheet("Rekap Batch Pengganti")
        build_worksheet(ws_pengganti, "REKAPITULASI BATCH INVERS PENGGANTI", "Sistem Verifikasi Perumahan Swadaya — Laporan per Batch Pengganti", pengganti_stages)
        
    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    
    filename = "REKAP_BATCH_BERITA_ACARA.xlsx"
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.get("/api/rekap-batch-verfal")
def get_rekap_batch_verfal(published_only: int = 1, province_id: int = 1):
    import re
    conn = get_db_connection()
    cursor = conn.cursor()
    
    if not province_id or province_id == 1:
        prov_stage_sql = "WHERE (s.province_id = 1 OR s.province_id IS NULL OR s.province_id = 0)"
        prov_stage_params = ()
    else:
        prov_stage_sql = "WHERE s.province_id = ?"
        prov_stage_params = (province_id,)

    cursor.execute(f"SELECT s.id, s.name FROM invers_stages s {prov_stage_sql} ORDER BY s.id ASC", prov_stage_params)
    all_stages = [dict(r) for r in cursor.fetchall()]
    def get_stage_num(s):
        match = re.search(r'\d+', s['name'])
        return int(match.group()) if match else 999
    all_stages = sorted(all_stages, key=get_stage_num)
    
    published_filter = "AND vb.is_published = 1" if published_only else ""
    cursor.execute(f"""
        SELECT DISTINCT UPPER(TRIM(COALESCE(ir.kabupaten_kota, ''))) as kab 
        FROM invers_records ir
        JOIN invers_revisions irv ON ir.revision_id = irv.id
        JOIN invers_stages s ON irv.stage_id = s.id
        {prov_stage_sql} AND irv.is_active = 1 AND TRIM(COALESCE(ir.kabupaten_kota, '')) != ''
        UNION
        SELECT DISTINCT UPPER(TRIM(COALESCE(vr.kabupaten_kota, ''))) as kab
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        JOIN invers_stages s ON vb.stage_id = s.id
        LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
        {prov_stage_sql} AND vb.batch_type = 'VERFAL' AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL) AND TRIM(COALESCE(vr.kabupaten_kota, '')) != ''
        {published_filter}
        ORDER BY kab ASC
    """, (*prov_stage_params, *prov_stage_params))
    all_kabupaten = [r['kab'] for r in cursor.fetchall()]
    
    stage_batch_filter = "AND is_published = 1" if published_only else ""
    stages_data = []
    
    for stage in all_stages:
        stage_id = stage['id']
        stage_name = stage['name']
        
        cursor.execute(f"""
            SELECT id, name, is_published, nomor_ba, tanggal_ba, sort_order, kabupaten 
            FROM verified_batches 
            WHERE stage_id = ? AND batch_type = 'VERFAL' {stage_batch_filter}
            ORDER BY sort_order ASC, uploaded_at ASC, id ASC
        """, (stage_id,))
        batches = [dict(r) for r in cursor.fetchall()]
        
        if not batches:
            continue
            
        batches_data = []
        for batch in batches:
            batch_id = batch['id']
            batch_name = batch['name']
            
            cursor.execute("""
                SELECT UPPER(TRIM(COALESCE(vr.kabupaten_kota, ''))) as kab,
                       SUM(CASE WHEN vr.status = 'LOLOS' THEN 1 ELSE 0 END) as lolos,
                       SUM(CASE WHEN vr.status = 'TIDAK LOLOS' THEN 1 ELSE 0 END) as tidak_lolos,
                       SUM(CASE WHEN vr.status = 'LOLOS' AND sk.verified_record_id IS NOT NULL THEN 1 ELSE 0 END) as sk_sudah
                FROM verified_records vr
                LEFT JOIN (
                    SELECT DISTINCT verified_record_id 
                    FROM sk_dirjen_matches 
                    WHERE verified_record_id IS NOT NULL
                      AND (match_type = 'PERFECT' 
                           OR (match_type = 'NEEDS_APPROVAL' AND override_status = 'APPROVED')
                           OR match_type = 'MANUAL_PAIR')
                ) sk ON sk.verified_record_id = vr.id
                WHERE vr.batch_id = ? AND TRIM(COALESCE(vr.kabupaten_kota, '')) != ''
                GROUP BY kab
            """, (batch_id,))
            
            stats_by_kab = {}
            for row in cursor.fetchall():
                lolos = row['lolos']
                tidak_lolos = row['tidak_lolos']
                sk_sudah = row['sk_sudah']
                sk_belum = max(0, lolos - sk_sudah)
                stats_by_kab[row['kab']] = {
                    "lolos": lolos,
                    "tidak_lolos": tidak_lolos,
                    "verifikasi": lolos + tidak_lolos,
                    "sk_sudah": sk_sudah,
                    "sk_belum": sk_belum
                }
            
            kab_data = []
            total_verifikasi = 0
            total_lolos = 0
            total_tidak_lolos = 0
            total_sk_sudah = 0
            total_sk_belum = 0
            
            for kab in all_kabupaten:
                stats = stats_by_kab.get(kab, {"lolos": 0, "tidak_lolos": 0, "verifikasi": 0, "sk_sudah": 0, "sk_belum": 0})
                kab_data.append({
                    "kabupaten": kab,
                    "verifikasi": stats["verifikasi"],
                    "lolos": stats["lolos"],
                    "tidak_lolos": stats["tidak_lolos"],
                    "sk_sudah": stats["sk_sudah"],
                    "sk_belum": stats["sk_belum"]
                })
                total_verifikasi += stats["verifikasi"]
                total_lolos += stats["lolos"]
                total_tidak_lolos += stats["tidak_lolos"]
                total_sk_sudah += stats["sk_sudah"]
                total_sk_belum += stats["sk_belum"]
                
            batches_data.append({
                "batch_id": batch_id,
                "batch_name": batch_name,
                "kabupaten": batch.get("kabupaten"),
                "is_published": batch["is_published"],
                "nomor_ba": batch.get("nomor_ba"),
                "tanggal_ba": batch.get("tanggal_ba"),
                "kabupaten_data": kab_data,
                "totals": {
                    "verifikasi": total_verifikasi,
                    "lolos": total_lolos,
                    "tidak_lolos": total_tidak_lolos,
                    "sk_sudah": total_sk_sudah,
                    "sk_belum": total_sk_belum
                }
            })
            
        stage_type = "pengganti" if "pengganti" in stage_name.lower() else "murni"
        stages_data.append({
            "stage_id": stage_id,
            "stage_name": stage_name,
            "stage_type": stage_type,
            "batches": batches_data
        })
        
    conn.close()
    return {
        "all_kabupaten": all_kabupaten,
        "stages": stages_data
    }

@app.get("/api/rekap-batch-verfal/export")
def export_rekap_batch_verfal(published_only: int = 1):
    import re
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT id, name FROM invers_stages ORDER BY id ASC")
    all_stages = [dict(r) for r in cursor.fetchall()]
    def get_stage_num(s):
        match = re.search(r'\d+', s['name'])
        return int(match.group()) if match else 999
    all_stages = sorted(all_stages, key=get_stage_num)
    
    published_filter = "AND vb.is_published = 1" if published_only else ""
    cursor.execute(f"""
        SELECT DISTINCT UPPER(TRIM(COALESCE(kabupaten_kota, ''))) as kab 
        FROM invers_records ir
        JOIN invers_revisions irv ON ir.revision_id = irv.id
        WHERE irv.is_active = 1 AND TRIM(COALESCE(kabupaten_kota, '')) != ''
        UNION
        SELECT DISTINCT UPPER(TRIM(COALESCE(vr.kabupaten_kota, ''))) as kab
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
        WHERE vb.batch_type = 'VERFAL' AND (vr.is_duplicate_in_previous = 0 OR ro.id IS NOT NULL) AND TRIM(COALESCE(vr.kabupaten_kota, '')) != ''
        {published_filter}
        ORDER BY kab ASC
    """)
    all_kabupaten = [r['kab'] for r in cursor.fetchall()]
    
    stage_batch_filter = "AND is_published = 1" if published_only else ""
    stages_data = []
    
    for stage in all_stages:
        stage_id = stage['id']
        stage_name = stage['name']
        
        cursor.execute(f"""
            SELECT id, name, nomor_ba, tanggal_ba, kabupaten 
            FROM verified_batches 
            WHERE stage_id = ? AND batch_type = 'VERFAL' {stage_batch_filter}
            ORDER BY sort_order ASC, uploaded_at ASC, id ASC
        """, (stage_id,))
        batches = [dict(r) for r in cursor.fetchall()]
        
        if not batches:
            continue
            
        batches_data = []
        for batch in batches:
            batch_id = batch['id']
            batch_name = batch['name']
            
            cursor.execute("""
                SELECT UPPER(TRIM(COALESCE(vr.kabupaten_kota, ''))) as kab,
                       SUM(CASE WHEN vr.status = 'LOLOS' THEN 1 ELSE 0 END) as lolos,
                       SUM(CASE WHEN vr.status = 'TIDAK LOLOS' THEN 1 ELSE 0 END) as tidak_lolos,
                       SUM(CASE WHEN vr.status = 'LOLOS' AND sk.verified_record_id IS NOT NULL THEN 1 ELSE 0 END) as sk_sudah
                FROM verified_records vr
                LEFT JOIN (
                    SELECT DISTINCT verified_record_id 
                    FROM sk_dirjen_matches 
                    WHERE verified_record_id IS NOT NULL
                      AND (match_type = 'PERFECT' 
                           OR (match_type = 'NEEDS_APPROVAL' AND override_status = 'APPROVED')
                           OR match_type = 'MANUAL_PAIR')
                ) sk ON sk.verified_record_id = vr.id
                WHERE vr.batch_id = ? AND TRIM(COALESCE(vr.kabupaten_kota, '')) != ''
                GROUP BY kab
            """, (batch_id,))
            
            stats_by_kab = {}
            for row in cursor.fetchall():
                lolos = row['lolos']
                tidak_lolos = row['tidak_lolos']
                sk_sudah = row['sk_sudah']
                sk_belum = max(0, lolos - sk_sudah)
                stats_by_kab[row['kab']] = {
                    "lolos": lolos,
                    "tidak_lolos": tidak_lolos,
                    "verifikasi": lolos + tidak_lolos,
                    "sk_sudah": sk_sudah,
                    "sk_belum": sk_belum
                }
            
            kab_data = {}
            total_verifikasi = 0
            total_lolos = 0
            total_tidak_lolos = 0
            total_sk_sudah = 0
            total_sk_belum = 0
            
            for kab in all_kabupaten:
                stats = stats_by_kab.get(kab, {"lolos": 0, "tidak_lolos": 0, "verifikasi": 0, "sk_sudah": 0, "sk_belum": 0})
                kab_data[kab] = stats
                total_verifikasi += stats["verifikasi"]
                total_lolos += stats["lolos"]
                total_tidak_lolos += stats["tidak_lolos"]
                total_sk_sudah += stats["sk_sudah"]
                total_sk_belum += stats["sk_belum"]
                
            batches_data.append({
                "batch_id": batch_id,
                "batch_name": batch_name,
                "kabupaten": batch.get("kabupaten"),
                "nomor_ba": batch.get("nomor_ba"),
                "tanggal_ba": batch.get("tanggal_ba"),
                "kabupaten_data": kab_data,
                "totals": {
                    "verifikasi": total_verifikasi,
                    "lolos": total_lolos,
                    "tidak_lolos": total_tidak_lolos,
                    "sk_sudah": total_sk_sudah,
                    "sk_belum": total_sk_belum
                }
            })
            
        stage_type = "pengganti" if "pengganti" in stage_name.lower() else "murni"
        stages_data.append({
            "stage_id": stage_id,
            "stage_name": stage_name,
            "stage_type": stage_type,
            "batches": batches_data
        })
        
    conn.close()
    
    wb = openpyxl.Workbook()
    
    font_main_title = Font(name='Arial', size=14, bold=True, color='1E293B')
    font_sub_title = Font(name='Arial', size=10, italic=True, color='64748B')
    font_stage_header = Font(name='Arial', size=11, bold=True, color='FFFFFF')
    font_batch_header = Font(name='Arial', size=10, bold=True, color='1E293B')
    font_sub_header = Font(name='Arial', size=9, bold=True, color='475569')
    font_sub_header_sk_sudah = Font(name='Arial', size=9, bold=True, color='047857')
    font_sub_header_sk_belum = Font(name='Arial', size=9, bold=True, color='B45309')
    font_data = Font(name='Arial', size=9, color='000000')
    font_total = Font(name='Arial', size=9, bold=True, color='000000')
    font_lolos_cell = Font(name='Arial', size=9, color='15803D', bold=True)
    font_tidak_lolos_cell = Font(name='Arial', size=9, color='DC2626', bold=True)
    font_sk_sudah_cell = Font(name='Arial', size=9, color='047857', bold=True)
    font_sk_belum_cell = Font(name='Arial', size=9, color='B45309', bold=True)
    
    fill_stage = PatternFill(start_color="1E3A8A", end_color="1E3A8A", fill_type="solid")
    fill_batch = PatternFill(start_color="E2E8F0", end_color="E2E8F0", fill_type="solid")
    fill_sub_v = PatternFill(start_color="F1F5F9", end_color="F1F5F9", fill_type="solid")
    fill_sub_l = PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid")
    fill_sub_tl = PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid")
    fill_sub_sk_sudah = PatternFill(start_color="D1FAE5", end_color="D1FAE5", fill_type="solid")
    fill_sub_sk_belum = PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid")
    
    fill_row_even = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
    fill_total_row = PatternFill(start_color="E2E8F0", end_color="E2E8F0", fill_type="solid")
    fill_lolos_cell = PatternFill(start_color="F0FDF4", end_color="F0FDF4", fill_type="solid")
    fill_tidak_lolos_cell = PatternFill(start_color="FEF2F2", end_color="FEF2F2", fill_type="solid")
    fill_sk_sudah_cell = PatternFill(start_color="ECFDF5", end_color="ECFDF5", fill_type="solid")
    fill_sk_belum_cell = PatternFill(start_color="FFFBEB", end_color="FFFBEB", fill_type="solid")
    
    border_thin = Border(
        left=Side(style='thin', color='CBD5E1'),
        right=Side(style='thin', color='CBD5E1'),
        top=Side(style='thin', color='CBD5E1'),
        bottom=Side(style='thin', color='CBD5E1')
    )
    border_double_top = Border(
        left=Side(style='thin', color='CBD5E1'),
        right=Side(style='thin', color='CBD5E1'),
        top=Side(style='double', color='000000'),
        bottom=Side(style='double', color='000000')
    )
    
    align_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    align_left = Alignment(horizontal='left', vertical='center')
    align_title = Alignment(horizontal='left', vertical='center')
    
    ws = wb.active
    ws.title = "Rekap BA Verfal"
    
    ws['A1'] = "REKAPITULASI BATCH BERITA ACARA VERIFIKASI FAKTUAL (VERFAL)"
    ws['A1'].font = font_main_title
    ws['A1'].alignment = align_title
    
    ws['A2'] = "Sistem Verifikasi Perumahan Swadaya — Laporan per Batch Verfal"
    ws['A2'].font = font_sub_title
    ws['A2'].alignment = align_title
    
    ws.cell(row=4, column=1, value="NO").font = font_stage_header
    ws.cell(row=4, column=1).alignment = align_center
    ws.cell(row=4, column=1).fill = fill_stage
    ws.cell(row=4, column=1).border = border_thin
    ws.merge_cells('A4:A6')
    
    ws.cell(row=4, column=2, value="KABUPATEN / KOTA").font = font_stage_header
    ws.cell(row=4, column=2).alignment = align_center
    ws.cell(row=4, column=2).fill = fill_stage
    ws.cell(row=4, column=2).border = border_thin
    ws.merge_cells('B4:B6')
    
    col_idx = 3
    for stage in stages_data:
        stage_batches = stage['batches']
        if not stage_batches:
            continue
        num_cols = len(stage_batches) * 5
        start_col = col_idx
        end_col = col_idx + num_cols - 1
        
        cell_stage = ws.cell(row=4, column=start_col, value=stage['stage_name'].upper())
        cell_stage.font = font_stage_header
        cell_stage.alignment = align_center
        cell_stage.fill = fill_stage
        
        for c in range(start_col, end_col + 1):
            ws.cell(row=4, column=c).border = border_thin
            ws.cell(row=4, column=c).fill = fill_stage
        if num_cols > 1:
            ws.merge_cells(start_row=4, start_column=start_col, end_row=4, end_column=end_col)
            
        for batch in stage_batches:
            b_start = col_idx
            b_end = col_idx + 4
            
            b_label = batch['batch_name']
            if batch.get('kabupaten'):
                b_label = f"{b_label} ({batch['kabupaten']})"
            if batch.get('nomor_ba'):
                b_label += f"\nNo: {batch['nomor_ba']}"
            if batch.get('tanggal_ba'):
                b_label += f"\nTgl: {batch['tanggal_ba']}"
                
            cell_batch = ws.cell(row=5, column=b_start, value=b_label)
            cell_batch.font = font_batch_header
            cell_batch.alignment = align_center
            cell_batch.fill = fill_batch
            
            for c in range(b_start, b_end + 1):
                ws.cell(row=5, column=c).border = border_thin
                ws.cell(row=5, column=c).fill = fill_batch
            ws.merge_cells(start_row=5, start_column=b_start, end_row=5, end_column=b_end)
            
            c_v = ws.cell(row=6, column=b_start, value="VERIFIKASI")
            c_l = ws.cell(row=6, column=b_start + 1, value="LOLOS")
            c_tl = ws.cell(row=6, column=b_start + 2, value="TIDAK LOLOS")
            c_sks = ws.cell(row=6, column=b_start + 3, value="SUDAH SK")
            c_skb = ws.cell(row=6, column=b_start + 4, value="BELUM SK")
            
            for c_sub, f_sub, fill_s in [
                (c_v, font_sub_header, fill_sub_v), 
                (c_l, font_sub_header, fill_sub_l), 
                (c_tl, font_sub_header, fill_sub_tl),
                (c_sks, font_sub_header_sk_sudah, fill_sub_sk_sudah),
                (c_skb, font_sub_header_sk_belum, fill_sub_sk_belum)
            ]:
                c_sub.font = f_sub
                c_sub.alignment = align_center
                c_sub.fill = fill_s
                c_sub.border = border_thin
                
            col_idx += 5
            
    row_idx = 7
    for k_idx, kab in enumerate(all_kabupaten, 1):
        is_even = (k_idx % 2 == 0)
        row_fill = fill_row_even if is_even else PatternFill(fill_type=None)
        
        cell_no = ws.cell(row=row_idx, column=1, value=k_idx)
        cell_no.font = font_data
        cell_no.alignment = align_center
        cell_no.border = border_thin
        if is_even: cell_no.fill = row_fill
        
        cell_kab = ws.cell(row=row_idx, column=2, value=kab)
        cell_kab.font = font_data
        cell_kab.alignment = align_left
        cell_kab.border = border_thin
        if is_even: cell_kab.fill = row_fill
        
        c_idx = 3
        for stage in stages_data:
            for batch in stage['batches']:
                stats = batch['kabupaten_data'].get(kab, {"verifikasi": 0, "lolos": 0, "tidak_lolos": 0, "sk_sudah": 0, "sk_belum": 0})
                val_v = stats['verifikasi']
                val_l = stats['lolos']
                val_tl = stats['tidak_lolos']
                val_sks = stats['sk_sudah']
                val_skb = stats['sk_belum']
                
                cell_v = ws.cell(row=row_idx, column=c_idx, value=val_v if val_v > 0 else "-")
                cell_l = ws.cell(row=row_idx, column=c_idx + 1, value=val_l if val_l > 0 else "-")
                cell_tl = ws.cell(row=row_idx, column=c_idx + 2, value=val_tl if val_tl > 0 else "-")
                cell_sks = ws.cell(row=row_idx, column=c_idx + 3, value=val_sks if val_sks > 0 else "-")
                cell_skb = ws.cell(row=row_idx, column=c_idx + 4, value=val_skb if val_skb > 0 else "-")
                
                for cell in (cell_v, cell_l, cell_tl, cell_sks, cell_skb):
                    cell.alignment = align_center
                    if is_even: cell.fill = row_fill
                    
                cell_v.font = font_data
                cell_l.font = font_lolos_cell if val_l > 0 else font_data
                cell_tl.font = font_tidak_lolos_cell if val_tl > 0 else font_data
                cell_sks.font = font_sk_sudah_cell if val_sks > 0 else font_data
                cell_skb.font = font_sk_belum_cell if val_skb > 0 else font_data
                
                if val_l > 0: cell_l.fill = fill_lolos_cell
                if val_tl > 0: cell_tl.fill = fill_tidak_lolos_cell
                if val_sks > 0: cell_sks.fill = fill_sk_sudah_cell
                if val_skb > 0: cell_skb.fill = fill_sk_belum_cell
                    
                cell_v.border = border_thin
                cell_l.border = border_thin
                cell_tl.border = border_thin
                cell_sks.border = border_thin
                cell_skb.border = border_thin
                c_idx += 5
        row_idx += 1
        
    cell_tot_label = ws.cell(row=row_idx, column=2, value="TOTAL")
    cell_tot_label.font = font_total
    cell_tot_label.alignment = align_left
    cell_tot_label.fill = fill_total_row
    cell_tot_label.border = border_double_top
    
    cell_tot_no = ws.cell(row=row_idx, column=1)
    cell_tot_no.fill = fill_total_row
    cell_tot_no.border = border_double_top
    
    c_idx = 3
    for stage in stages_data:
        for batch in stage['batches']:
            t = batch['totals']
            cell_tot_v = ws.cell(row=row_idx, column=c_idx, value=t['verifikasi'])
            cell_tot_l = ws.cell(row=row_idx, column=c_idx + 1, value=t['lolos'])
            cell_tot_tl = ws.cell(row=row_idx, column=c_idx + 2, value=t['tidak_lolos'])
            cell_tot_sks = ws.cell(row=row_idx, column=c_idx + 3, value=t['sk_sudah'])
            cell_tot_skb = ws.cell(row=row_idx, column=c_idx + 4, value=t['sk_belum'])
            
            cell_tot_v.font = font_total
            cell_tot_l.font = font_total
            cell_tot_tl.font = font_total
            cell_tot_sks.font = font_total
            cell_tot_skb.font = font_total
            
            cell_tot_v.alignment = align_center
            cell_tot_l.alignment = align_center
            cell_tot_tl.alignment = align_center
            cell_tot_sks.alignment = align_center
            cell_tot_skb.alignment = align_center
            
            cell_tot_v.fill = fill_total_row
            cell_tot_l.fill = fill_total_row
            cell_tot_tl.fill = fill_total_row
            cell_tot_sks.fill = fill_total_row
            cell_tot_skb.fill = fill_total_row
            
            cell_tot_v.border = border_double_top
            cell_tot_l.border = border_double_top
            cell_tot_tl.border = border_double_top
            cell_tot_sks.border = border_double_top
            cell_tot_skb.border = border_double_top
            c_idx += 5
            
    for col in ws.columns:
        col_idx = col[0].column
        col_letter = get_column_letter(col_idx)
        if col_idx == 1:
            ws.column_dimensions[col_letter].width = 6
        elif col_idx == 2:
            ws.column_dimensions[col_letter].width = 28
        else:
            ws.column_dimensions[col_letter].width = 14
    ws.freeze_panes = 'C7'
    
    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    
    filename = "REKAP_BATCH_VERFAL.xlsx"
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# --- GLOBAL SEARCH (Cross-stage search for invers + verified records) ---
@app.get("/api/global-search")
def global_search(
    q: str = "",
    kabupaten: str = "",
    kecamatan: str = "",
    desa: str = "",
    status: str = "",
    tahap: str = "",
    sk_dirjen: str = "ALL",
    pengusul: str = "",
    record_type: str = "all",
    published_only: int = 1,
    province_id: int = 1,
    page: int = 1,
    limit: int = 30,
    export_all: bool = False
):
    conn = get_db_connection()
    cursor = conn.cursor()

    conditions = []
    params = []

    if not province_id or province_id == 1:
        conditions.append("(ist.province_id = 1 OR ist.province_id IS NULL OR ist.province_id = 0)")
    else:
        conditions.append("ist.province_id = ?")
        params.append(province_id)

    # --- Build WHERE clauses ---
    if q and len(q.strip()) >= 1:
        term = f"%{q.strip().upper()}%"
        conditions.append("(UPPER(vr.nama) LIKE ? OR vr.no_ktp LIKE ? OR vr.no_kk LIKE ?)")
        params.extend([term, term, term])

    if kabupaten:
        conditions.append("UPPER(vr.kabupaten_kota) = UPPER(?)")
        params.append(kabupaten)
    if kecamatan:
        conditions.append("UPPER(vr.kecamatan) = UPPER(?)")
        params.append(kecamatan)
    if desa:
        conditions.append("UPPER(vr.desa_kelurahan) = UPPER(?)")
        params.append(desa)
    if tahap:
        conditions.append("ist.id = ?")
        params.append(int(tahap))
    if pengusul:
        conditions.append("UPPER(COALESCE(ir_p.pengusul, '')) = UPPER(?)")
        params.append(pengusul)

    where_clause = " AND ".join(conditions) if conditions else "1=1"

    # Status filter for verified records
    verified_status_clause = ""
    if status == "LOLOS":
        verified_status_clause = "AND vr.status = 'LOLOS'"
    elif status == "TIDAK_LOLOS":
        verified_status_clause = "AND vr.status = 'TIDAK LOLOS'"

    sk_dirjen_clause = ""
    if sk_dirjen == "SUDAH":
        sk_dirjen_clause = "AND sk.has_sk = 1"
    elif sk_dirjen == "BELUM":
        sk_dirjen_clause = "AND (sk.has_sk IS NULL OR sk.has_sk = 0)"

    published_clause = "AND vb.is_published = 1" if published_only == 1 else ""

    all_records = []
    summary = {"total_alokasi": 0, "total_verifikasi": 0, "total_lolos": 0, "total_tidak_lolos": 0, "total_belum": 0}

    # --- Verified records ---
    if record_type in ("all", "verified"):
        query = f"""
            SELECT vr.id, vr.nama, vr.no_ktp, vr.no_kk, vr.kabupaten_kota, vr.kecamatan,
                   vr.desa_kelurahan, vr.status, ist.id as tahap_id, ist.name as tahap_name,
                   vb.name as batch_name, 'verified' as record_type,
                   CASE WHEN sk.has_sk IS NOT NULL THEN 'SUDAH' ELSE 'BELUM' END as sk_dirjen_status,
                   sk.nomor_sk,
                   COALESCE(ir_p.pengusul, 'LAINNYA') as pengusul
            FROM verified_records vr
            JOIN verified_batches vb ON vb.id = vr.batch_id
            JOIN invers_stages ist ON ist.id = vb.stage_id
            LEFT JOIN (
                SELECT DISTINCT no_ktp, pengusul
                FROM invers_records
                WHERE pengusul IS NOT NULL AND pengusul != ''
            ) ir_p ON ir_p.no_ktp = vr.no_ktp
            LEFT JOIN (
                SELECT DISTINCT m.verified_record_id, 1 as has_sk, b.stage_name as nomor_sk
                FROM sk_dirjen_matches m
                JOIN sk_dirjen_records r ON r.id = m.sk_record_id
                JOIN sk_dirjen_batches b ON b.id = r.batch_id
                WHERE m.verified_record_id IS NOT NULL
                  AND (m.match_type = 'PERFECT' OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED') OR m.match_type = 'MANUAL_PAIR')
            ) sk ON sk.verified_record_id = vr.id
            WHERE {where_clause}
            {verified_status_clause}
            {sk_dirjen_clause}
            {published_clause}
            ORDER BY vr.nama ASC
        """
        cursor.execute(query, params)
        for row in cursor.fetchall():
            rec = dict(row)
            rec['status'] = rec['status'] if rec['status'] else 'SELESAI'
            all_records.append(rec)

    # --- Invers records (belum diverifikasi) ---
    if record_type in ("all", "invers") and sk_dirjen != "SUDAH":
        invers_conditions = []
        invers_params = []
        
        if not province_id or province_id == 1:
            invers_conditions.append("(ist.province_id = 1 OR ist.province_id IS NULL OR ist.province_id = 0)")
        else:
            invers_conditions.append("ist.province_id = ?")
            invers_params.append(province_id)

        if q and len(q.strip()) >= 1:
            term = f"%{q.strip().upper()}%"
            invers_conditions.append("(UPPER(ir.nama) LIKE ? OR ir.no_ktp LIKE ? OR ir.no_kk LIKE ?)")
            invers_params.extend([term, term, term])
        if kabupaten:
            invers_conditions.append("UPPER(ir.kabupaten_kota) = UPPER(?)")
            invers_params.append(kabupaten)
        if kecamatan:
            invers_conditions.append("UPPER(ir.kecamatan) = UPPER(?)")
            invers_params.append(kecamatan)
        if desa:
            invers_conditions.append("UPPER(ir.desa_kelurahan) = UPPER(?)")
            invers_params.append(desa)
        if tahap:
            invers_conditions.append("irv.stage_id = ?")
            invers_params.append(int(tahap))
        if pengusul:
            invers_conditions.append("UPPER(COALESCE(ir.pengusul, '')) = UPPER(?)")
            invers_params.append(pengusul)

        invers_where = " AND ".join(invers_conditions) if invers_conditions else "1=1"
        published_invers_clause = "AND vb2.is_published = 1" if published_only == 1 else ""

        invers_query = f"""
            SELECT ir.id, ir.nama, ir.no_ktp, ir.no_kk, ir.kabupaten_kota, ir.kecamatan,
                   ir.desa_kelurahan, irv.stage_id as tahap_id, ist.name as tahap_name,
                   'Belum Diverifikasi' as batch_name, 'invers' as record_type,
                   COALESCE(ir.pengusul, 'LAINNYA') as pengusul
            FROM invers_records ir
            JOIN invers_revisions irv ON ir.revision_id = irv.id
            JOIN invers_stages ist ON ist.id = irv.stage_id
            WHERE irv.is_active = 1
              AND {invers_where}
              AND ir.no_ktp NOT IN (
                  SELECT vr2.no_ktp FROM verified_records vr2
                  JOIN verified_batches vb2 ON vb2.id = vr2.batch_id
                  WHERE vb2.stage_id = irv.stage_id {published_invers_clause}
              )
            ORDER BY ir.nama ASC
        """
        cursor.execute(invers_query, invers_params)
        for row in cursor.fetchall():
            rec = dict(row)
            rec['status'] = 'BELUM'
            rec['sk_dirjen_status'] = 'BELUM'
            rec['nomor_sk'] = None
            all_records.append(rec)

    conn.close()

    # Sort: verified first, then by name
    all_records.sort(key=lambda r: (0 if r['record_type'] == 'verified' else 1, r['nama'] or ''))

    # Summary from ALL records (before pagination)
    summary['total_alokasi'] = len(all_records)
    summary['total_verifikasi'] = len([r for r in all_records if r['record_type'] == 'verified'])
    summary['total_lolos'] = len([r for r in all_records if r.get('status') == 'LOLOS'])
    summary['total_tidak_lolos'] = len([r for r in all_records if r.get('status') == 'TIDAK LOLOS'])
    summary['total_belum'] = len([r for r in all_records if r['record_type'] == 'invers'])

    total = len(all_records)
    total_pages = max(1, (total + limit - 1) // limit)

    if export_all:
        return {"records": all_records, "total": total, "page": 1, "total_pages": 1, "summary": summary}

    start = (page - 1) * limit
    paginated = all_records[start:start + limit]

    # Fetch filter options efficiently
    all_kabs = sorted(list(set(r['kabupaten_kota'] for r in all_records if r.get('kabupaten_kota'))))
    all_kecs = sorted(list(set(r['kecamatan'] for r in all_records if r.get('kecamatan'))))
    all_desas = sorted(list(set(r['desa_kelurahan'] for r in all_records if r.get('desa_kelurahan'))))
    all_pengusuls = sorted(list(set(r['pengusul'] for r in all_records if r.get('pengusul'))))
    
    all_tahaps_raw = {}
    for r in all_records:
        tid = r.get('tahap_id')
        tname = r.get('tahap_name')
        if tid and tid not in all_tahaps_raw:
            all_tahaps_raw[tid] = tname
    all_tahaps = [{"id": k, "name": v} for k, v in sorted(all_tahaps_raw.items())]

    return {
        "records": paginated,
        "total": total,
        "page": page,
        "total_pages": total_pages,
        "summary": summary,
        "filters": {
            "kabupatens": all_kabs,
            "kecamatans": all_kecs,
            "desas": all_desas,
            "tahaps": all_tahaps,
            "pengusuls": all_pengusuls
        }
    }

@app.get("/api/global-search/export")
def global_search_export(
    q: str = "",
    kabupaten: str = "",
    kecamatan: str = "",
    desa: str = "",
    status: str = "",
    tahap: str = "",
    sk_dirjen: str = "ALL",
    pengusul: str = "",
    record_type: str = "all"
):
    result = global_search(q=q, kabupaten=kabupaten, kecamatan=kecamatan, desa=desa,
                           status=status, tahap=tahap, sk_dirjen=sk_dirjen, pengusul=pengusul, record_type=record_type,
                           page=1, limit=99999, export_all=True)
    records = result["records"]

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Pencarian Global"

    headers = ["No", "Tahap", "Nama", "NIK", "No KK", "Kabupaten/Kota", "Kecamatan", "Desa/Kelurahan", "Status", "SK Dirjen", "Nomor SK", "Asal", "Pengusul"]
    header_font = openpyxl.styles.Font(bold=True, color="FFFFFF")
    header_fill = openpyxl.styles.PatternFill(start_color="1A3C40", end_color="1A3C40", fill_type="solid")
    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = openpyxl.styles.Alignment(horizontal="center")

    for row_idx, rec in enumerate(records, 2):
        ws.cell(row=row_idx, column=1, value=row_idx - 1)
        ws.cell(row=row_idx, column=2, value=rec.get("tahap_name", ""))
        ws.cell(row=row_idx, column=3, value=rec.get("nama", ""))
        ws.cell(row=row_idx, column=4, value=rec.get("no_ktp", ""))
        ws.cell(row=row_idx, column=5, value=rec.get("no_kk", ""))
        ws.cell(row=row_idx, column=6, value=rec.get("kabupaten_kota", ""))
        ws.cell(row=row_idx, column=7, value=rec.get("kecamatan", ""))
        ws.cell(row=row_idx, column=8, value=rec.get("desa_kelurahan", ""))
        ws.cell(row=row_idx, column=9, value=rec.get("status", ""))
        ws.cell(row=row_idx, column=10, value=rec.get("sk_dirjen_status", "BELUM"))
        ws.cell(row=row_idx, column=11, value=rec.get("nomor_sk") or "-")
        ws.cell(row=row_idx, column=12, value="Terverifikasi" if rec["record_type"] == "verified" else "Belum Diverifikasi")
        ws.cell(row=row_idx, column=13, value=rec.get("pengusul") or "-")

    for col in ws.columns:
        max_len = 0
        col_letter = col[0].column_letter
        for cell in col:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[col_letter].width = min(max_len + 2, 40)

    import io
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    filename = f"Pencarian_Global_{kabupaten or 'Semua'}_{status or 'Semua'}_SK_{sk_dirjen}.xlsx"
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.get("/api/stage/{stage_id}/export")
def export_excel(stage_id: int, batch_id: int = None):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT s.name as stage_name, p.name as province_name 
        FROM invers_stages s 
        LEFT JOIN provinces p ON s.province_id = p.id 
        WHERE s.id = ?
    """, (stage_id,))
    stage_row = cursor.fetchone()
    stage_name = stage_row['stage_name'] if stage_row else "Tahap"
    prov_name_clean = clean_province_for_export(stage_row['province_name'] if stage_row and stage_row['province_name'] else "SULAWESI SELATAN")
    
    if batch_id:
        cursor.execute("SELECT name FROM verified_batches WHERE id = ?", (batch_id,))
        batch_row = cursor.fetchone()
        batch_title = batch_row['name'] if batch_row else "Berita Acara"
        
        cursor.execute("""
            SELECT vr.*, re.nama_pengganti, re.no_ktp_pengganti, re.no_kk_pengganti, re.alamat_pengganti,
                   re.desa_kelurahan_pengganti, re.kecamatan_pengganti, re.kabupaten_pengganti,
                   re.jenis_kelamin_pengganti,
                   ro.id as override_id
            FROM verified_records vr
            LEFT JOIN replacement_events re ON re.disqualified_record_id = vr.id
            LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = ?
            WHERE vr.batch_id = ?
        """, (stage_id, batch_id))
    else:
        batch_title = "Semua Berita Acara"
        cursor.execute("""
            SELECT vr.*, re.nama_pengganti, re.no_ktp_pengganti, re.no_kk_pengganti, re.alamat_pengganti,
                   re.desa_kelurahan_pengganti, re.kecamatan_pengganti, re.kabupaten_pengganti,
                   re.jenis_kelamin_pengganti,
                   ro.id as override_id
            FROM verified_records vr
            JOIN verified_batches vb ON vr.batch_id = vb.id
            LEFT JOIN replacement_events re ON re.disqualified_record_id = vr.id
            LEFT JOIN reconciliation_overrides ro ON ro.original_no_ktp = vr.no_ktp AND ro.stage_id = vb.stage_id
            WHERE vb.stage_id = ?
        """, (stage_id,))
        
    records = [dict(row) for row in cursor.fetchall()]
    conn.close()
    
    # Filter: include non-duplicates OR reconciled duplicates
    def should_include_in_ba(r):
        if r['is_duplicate_in_previous'] == 0:
            return True
        return r.get('override_id') is not None
    
    lolos_records = [r for r in records if r['status'] == 'LOLOS' and should_include_in_ba(r)]
    tidak_lolos_records = [r for r in records if r['status'] == 'TIDAK LOLOS' and should_include_in_ba(r)]
    
    summary_by_kab = {}
    for r in records:
        if not should_include_in_ba(r):
            continue
        kab = (r['kabupaten_kota'] or "KAB. LUWU UTARA").upper().strip()
        if kab not in summary_by_kab:
            summary_by_kab[kab] = {"cpb": 0, "lolos": 0, "tidak_lolos": 0, "pengganti": 0}
            
        summary_by_kab[kab]["cpb"] += 1
        if r['status'] == 'LOLOS':
            summary_by_kab[kab]["lolos"] += 1
        else:
            summary_by_kab[kab]["tidak_lolos"] += 1
            if r['nama_pengganti']:
                summary_by_kab[kab]["pengganti"] += 1
                
    wb = openpyxl.Workbook()
    
    font_family = "Bookman Old Style"
    f_body = Font(name=font_family, size=12)
    f_body_bold = Font(name=font_family, size=12, bold=True)
    f_title = Font(name=font_family, size=14, bold=True)
    f_subtitle = Font(name=font_family, size=12, italic=True)
    
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    
    fill_header_blue = PatternFill(start_color="8EA9DB", end_color="8EA9DB", fill_type="solid")
    fill_header_red = PatternFill(start_color="C00000", end_color="C00000", fill_type="solid")
    fill_header_green = PatternFill(start_color="375623", end_color="375623", fill_type="solid")
    
    thin_side = Side(border_style="thin", color="000000")
    double_bottom_side = Side(border_style="double", color="000000")
    
    border_data = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
    border_total = Border(left=thin_side, right=thin_side, top=thin_side, bottom=double_bottom_side)
    
    # --- Sheet 1: Lamp.IA (Summary of Kabupaten) ---
    ws_ia = wb.active
    ws_ia.title = "Lamp.IA"
    ws_ia.views.sheetView[0].showGridLines = True
    ws_ia.page_setup.orientation = ws_ia.ORIENTATION_PORTRAIT
    ws_ia.page_margins.top = ws_ia.page_margins.bottom = ws_ia.page_margins.left = ws_ia.page_margins.right = 0.5
    ws_ia.sheet_properties.pageSetUpPr.fitToPage = True
    ws_ia.page_setup.fitToWidth = 1
    ws_ia.page_setup.fitToHeight = 0
    
    ws_ia.append([])
    ws_ia.merge_cells("A2:F2")
    ws_ia.merge_cells("A3:F3")
    ws_ia.merge_cells("A4:F4")
    ws_ia.cell(row=2, column=1, value="HASIL VERIFIKASI CALON PENERIMA BANTUAN").font = f_title
    ws_ia.cell(row=2, column=1).alignment = align_center
    ws_ia.cell(row=3, column=1, value=f"KEGIATAN BANTUAN STIMULAN PERUMAHAN SWADAYA {stage_name.upper()} TAHUN 2026").font = f_title
    ws_ia.cell(row=3, column=1).alignment = align_center
    ws_ia.cell(row=4, column=1, value=f"PROVINSI {prov_name_clean}").font = f_title
    ws_ia.cell(row=4, column=1).alignment = align_center
    
    headers_ia_row1 = ["No.", "KABUPATEN/KOTA", "CPB (unit)", "Hasil Verifikasi", "", ""]
    headers_ia_row2 = ["", "", "", "Lolos (unit)", "Tidak Lolos (Unit)", "CPB Pengganti (unit)"]
    
    ws_ia.append([])
    ws_ia.append(headers_ia_row1)
    ws_ia.append(headers_ia_row2)
    
    ws_ia.merge_cells("A6:A7")
    ws_ia.merge_cells("B6:B7")
    ws_ia.merge_cells("C6:C7")
    ws_ia.merge_cells("D6:F6")
    
    for r in [6, 7]:
        for c in range(1, 7):
            cell = ws_ia.cell(row=r, column=c)
            cell.font = f_body_bold
            cell.alignment = align_center
            cell.border = border_data
            cell.fill = fill_header_blue
            if r == 6 and c >= 4:
                cell.font = Font(name=font_family, size=12, bold=True, color="000000")
                cell.fill = fill_header_blue if c == 4 else fill_header_red
                
    row_num = 8
    tot_cpb = tot_lolos = tot_tl = tot_peng = 0
    for idx, (kab, val) in enumerate(summary_by_kab.items()):
        ws_ia.append([idx+1, kab, val["cpb"], val["lolos"], val["tidak_lolos"], val["pengganti"]])
        tot_cpb += val["cpb"]
        tot_lolos += val["lolos"]
        tot_tl += val["tidak_lolos"]
        tot_peng += val["pengganti"]
        
        for c in range(1, 7):
            cell = ws_ia.cell(row=row_num, column=c)
            cell.font = f_body
            cell.border = border_data
            cell.alignment = align_center if c != 2 else align_left
        row_num += 1
        
    ws_ia.append(["", "TOTAL", tot_cpb, tot_lolos, tot_tl, tot_peng])
    ws_ia.merge_cells(start_row=row_num, start_column=1, end_row=row_num, end_column=2)
    for c in range(1, 7):
        cell = ws_ia.cell(row=row_num, column=c)
        cell.font = f_body_bold
        cell.border = border_total
        cell.alignment = align_center if c != 2 else align_left
        
    ws_ia.cell(row=row_num+2, column=5, value="PARAF").font = f_body_bold
    ws_ia.cell(row=row_num+2, column=5).alignment = align_center
    
    # Border box di bawah PARAF untuk kolom tanda tangan
    paraf_box_start_row = row_num + 3
    paraf_box_end_row = row_num + 5
    for r_box in range(paraf_box_start_row, paraf_box_end_row + 1):
        for c_box in [5, 6]:
            cell = ws_ia.cell(row=r_box, column=c_box)
            cell.border = border_data
    
    ws_ia.column_dimensions['A'].width = 8
    ws_ia.column_dimensions['B'].width = 30
    ws_ia.column_dimensions['C'].width = 11.5
    ws_ia.column_dimensions['D'].width = 15
    ws_ia.column_dimensions['E'].width = 18
    ws_ia.column_dimensions['F'].width = 22

    # --- Sheet 2: Lamp.IIA (Lolos) ---
    ws_iia = wb.create_sheet(title="Lamp.IIA")
    ws_iia.views.sheetView[0].showGridLines = True
    ws_iia.page_setup.orientation = ws_iia.ORIENTATION_LANDSCAPE
    ws_iia.page_margins.top = ws_iia.page_margins.bottom = ws_iia.page_margins.left = ws_iia.page_margins.right = 0.5
    ws_iia.page_margins.header = ws_iia.page_margins.footer = 0.25
    ws_iia.sheet_properties.pageSetUpPr.fitToPage = True
    ws_iia.page_setup.fitToWidth = 1
    ws_iia.page_setup.fitToHeight = 0
    
    ws_iia.append([])
    ws_iia.merge_cells("A2:P2")
    ws_iia.merge_cells("A3:P3")
    ws_iia.merge_cells("A4:P4")
    ws_iia.cell(row=2, column=1, value="DAFTAR CALON PENERIMA BANTUAN KEGIATAN BANTUAN STIMULAN").font = f_title
    ws_iia.cell(row=2, column=1).alignment = align_center
    ws_iia.cell(row=3, column=1, value=f"KEGIATAN BANTUAN STIMULAN PERUMAHAN SWADAYA {stage_name.upper()} TAHUN 2026").font = f_title
    ws_iia.cell(row=3, column=1).alignment = align_center
    ws_iia.cell(row=4, column=1, value=f"PROVINSI {prov_name_clean}").font = f_title
    ws_iia.cell(row=4, column=1).alignment = align_center
    
    headers_iia = [
        "NO. URUT", "KODE DESA / KEL", "NAMA", "JENIS KELAMIN (L/P)", "NO.KTP", "NO.KK",
        "ALAMAT TEMPAT TINGGAL", "DESA / KELURAHAN", "KECAMATAN", "KABUPATEN / KOTA",
        "*) LOLOS / PENGGANTI", "LATITUDE", "LONGITUDE", "TAHAP", "TANGGAL", "KETERANGAN"
    ]
    ws_iia.append([])
    ws_iia.append(headers_iia)
    
    for c in range(1, len(headers_iia)+1):
        cell = ws_iia.cell(row=6, column=c)
        cell.font = f_body_bold
        cell.alignment = align_center
        cell.border = border_data
        cell.fill = fill_header_blue
        
    row_num = 7
    for idx, r in enumerate(lolos_records):
        ws_iia.append([
            idx+1, r.get('kode_desa') or "", r.get('nama') or "", r.get('jenis_kelamin') or "",
            f"'{r.get('no_ktp')}" if r.get('no_ktp') else "", f"'{r.get('no_kk')}" if r.get('no_kk') else "",
            r.get('alamat') or "", r.get('desa_kelurahan') or "", r.get('kecamatan') or "",
            r.get('kabupaten_kota') or "", "LOLOS", r.get('latitude') or "", r.get('longitude') or "",
            r.get('tahap') or "", r.get('tanggal') or "", r.get('keterangan') or ""
        ])
        for c in range(1, len(headers_iia)+1):
            cell = ws_iia.cell(row=row_num, column=c)
            cell.font = f_body
            cell.border = border_data
            if c in [1, 2, 4, 5, 6, 11, 12, 13, 14, 15]:
                cell.alignment = align_center
            else:
                cell.alignment = align_left
        ws_iia.row_dimensions[row_num].height = 28
        row_num += 1
        
    for col in ws_iia.columns:
        col_letter = get_column_letter(col[0].column)
        ws_iia.column_dimensions[col_letter].width = 18
    # Column widths sesuai permintaan user
    ws_iia.column_dimensions['A'].width = 10      # NO. URUT
    ws_iia.column_dimensions['B'].width = 18       # KODE DESA
    ws_iia.column_dimensions['C'].width = 36.5     # NAMA
    ws_iia.column_dimensions['D'].width = 18       # JENIS KELAMIN
    ws_iia.column_dimensions['E'].width = 27       # NO.KTP (NIK)
    ws_iia.column_dimensions['F'].width = 27       # NO.KK
    ws_iia.column_dimensions['G'].width = 30       # ALAMAT TEMPAT TINGGAL
    ws_iia.column_dimensions['H'].width = 27       # DESA / KELURAHAN
    ws_iia.column_dimensions['I'].width = 27       # KECAMATAN
    ws_iia.column_dimensions['J'].width = 42       # KABUPATEN / KOTA
    
    # --- Sheet 3: Lamp.IIIA (Tidak Lolos / Pengganti) ---
    ws_iiia = wb.create_sheet(title="Lamp.IIIA")
    ws_iiia.views.sheetView[0].showGridLines = True
    ws_iiia.page_setup.orientation = ws_iiia.ORIENTATION_LANDSCAPE
    ws_iiia.page_margins.top = ws_iiia.page_margins.bottom = ws_iiia.page_margins.left = ws_iiia.page_margins.right = 0.5
    ws_iiia.page_margins.header = ws_iiia.page_margins.footer = 0.25
    ws_iiia.sheet_properties.pageSetUpPr.fitToPage = True
    ws_iiia.page_setup.fitToWidth = 1
    ws_iiia.page_setup.fitToHeight = 0
    
    ws_iiia.append([])
    ws_iiia.merge_cells("A2:V2")
    ws_iiia.merge_cells("A3:V3")
    ws_iiia.merge_cells("A4:V4")
    ws_iiia.cell(row=2, column=1, value="DAFTAR CALON PENERIMA BANTUAN PENGGANTI KEGIATAN BANTUAN").font = f_title
    ws_iiia.cell(row=2, column=1).alignment = align_center
    ws_iiia.cell(row=3, column=1, value=f"STIMULAN PERUMAHAN SWADAYA {stage_name.upper()} TAHUN 2026").font = f_title
    ws_iiia.cell(row=3, column=1).alignment = align_center
    ws_iiia.cell(row=4, column=1, value=f"PROVINSI {prov_name_clean}").font = f_title
    ws_iiia.cell(row=4, column=1).alignment = align_center
    
    headers_iiia_top = ["NO.", "TIDAK LOLOS", "", "", "", "", "", "", "", "", "SESUDAH", "", "", "", "", "", "", "", "", "INSTRUKSI VERIFIKASI", "", "KETERANGAN"]
    headers_iiia_bottom = [
        "NO.", "NAMA", "JENIS KELAMIN (L/P)", "NO.KTP", "NO.KK", "ALAMAT TEMPAT TINGGAL",
        "DESA / KELURAHAN", "KECAMATAN", "KABUPATEN", "ALASAN TIDAK LOLOS *)",
        "BNBA", "NAMA (PENGGANTI)", "JENIS KELAMIN (L/P)", "NO.KTP (PENGGANTI)", "NO.KK (PENGGANTI)",
        "ALAMAT TEMPAT TINGGAL", "DESA / KELURAHAN", "KECAMATAN", "KABUPATEN",
        "TAHAP", "TANGGAL", "KETERANGAN"
    ]
    
    ws_iiia.append([])
    ws_iiia.append(headers_iiia_top)
    ws_iiia.append(headers_iiia_bottom)
    
    ws_iiia.merge_cells("A6:A7")
    ws_iiia.merge_cells("B6:J6")
    ws_iiia.merge_cells("K6:S6")
    ws_iiia.merge_cells("T6:U6")
    ws_iiia.merge_cells("V6:V7")
    
    for r in [6, 7]:
        for c in range(1, len(headers_iiia_bottom)+1):
            cell = ws_iiia.cell(row=r, column=c)
            cell.font = f_body_bold
            cell.border = border_data
            cell.alignment = align_center
            
            if c >= 2 and c <= 10:
                cell.fill = fill_header_red
                cell.font = Font(name=font_family, size=12, bold=True, color="FFFFFF")
            elif c >= 11 and c <= 19:
                cell.fill = fill_header_green
                cell.font = Font(name=font_family, size=12, bold=True, color="FFFFFF")
            else:
                cell.fill = fill_header_blue
                
    row_num = 8
    for idx, r in enumerate(tidak_lolos_records):
        ws_iiia.append([
            idx+1, r.get('nama') or "", r.get('jenis_kelamin') or "", f"'{r.get('no_ktp')}" if r.get('no_ktp') else "", f"'{r.get('no_kk')}" if r.get('no_kk') else "",
            r.get('alamat') or "", r.get('desa_kelurahan') or "", r.get('kecamatan') or "", r.get('kabupaten_kota') or "",
            r.get('alasan_tidak_lolos') or "", "", r.get('nama_pengganti') or "", r.get('jenis_kelamin_pengganti') or "",
            f"'{r.get('no_ktp_pengganti')}" if r.get('no_ktp_pengganti') else "", f"'{r.get('no_kk_pengganti')}" if r.get('no_kk_pengganti') else "",
            r.get('alamat_pengganti') or "", r.get('desa_kelurahan_pengganti') or "", r.get('kecamatan_pengganti') or "", r.get('kabupaten_pengganti') or "",
            r.get('tahap') or "", r.get('tanggal') or "", r.get('keterangan') or ""
        ])
        for c in range(1, len(headers_iiia_bottom)+1):
            cell = ws_iiia.cell(row=row_num, column=c)
            cell.font = f_body
            cell.border = border_data
            if c in [1, 3, 4, 5, 10, 11, 13, 14, 15, 20, 21]:
                cell.alignment = align_center
            else:
                cell.alignment = align_left
        ws_iiia.row_dimensions[row_num].height = 28
        row_num += 1

    # Catatan section di bawah tabel Lamp.IIIA
    catatan_start = row_num + 1
    ws_iiia.cell(row=catatan_start, column=1, value="Catatan:").font = Font(name=font_family, size=10, bold=True)
    catatan_items = [
        "*) Alasan Tidak Lolos, diisi dengan angka (1-8) sebagai berikut:",
        "1. Belum memiliki KK sendiri;",
        "2. Tanah bersengketa;",
        "3. Rumah dalam kondisi layak;",
        "4. Memiliki rumah lebih dari 1;",
        "5. Pernah memperoleh bantuan dari APBN/APBD/CSR/anggaran lainnya;",
        "6. Penghasilan lebih dari UMP;",
        "7. Memilih untuk dibantu dengan sumber anggaran lain;",
        "8. Menghuni kurang dari 3 tahun;",
        "9. Lainnya (diisi pada kolom keterangan);"
    ]
    for i, item in enumerate(catatan_items):
        ws_iiia.cell(row=catatan_start + 1 + i, column=1, value=item).font = Font(name=font_family, size=9)
        
    for col in ws_iiia.columns:
        col_letter = get_column_letter(col[0].column)
        ws_iiia.column_dimensions[col_letter].width = 18
    # Column widths sesuai permintaan user
    ws_iiia.column_dimensions['A'].width = 8        # NO.
    ws_iiia.column_dimensions['B'].width = 25       # NAMA (Tidak Lolos)
    ws_iiia.column_dimensions['D'].width = 27       # NO.KTP (Tidak Lolos)
    ws_iiia.column_dimensions['E'].width = 27       # NO.KK (Tidak Lolos)
    ws_iiia.column_dimensions['F'].width = 30       # ALAMAT (Tidak Lolos)
    ws_iiia.column_dimensions['G'].width = 27       # DESA / KELURAHAN (Tidak Lolos)
    ws_iiia.column_dimensions['H'].width = 27       # KECAMATAN (Tidak Lolos)
    ws_iiia.column_dimensions['I'].width = 50       # KABUPATEN (Tidak Lolos)
    ws_iiia.column_dimensions['J'].width = 60       # ALASAN TIDAK LOLOS
    ws_iiia.column_dimensions['L'].width = 25       # NAMA (Pengganti)
    ws_iiia.column_dimensions['N'].width = 27       # NO.KTP (Pengganti)
    ws_iiia.column_dimensions['O'].width = 27       # NO.KK (Pengganti)
    ws_iiia.column_dimensions['P'].width = 30       # ALAMAT (Pengganti)
    ws_iiia.column_dimensions['Q'].width = 27       # DESA / KELURAHAN (Pengganti)
    ws_iiia.column_dimensions['R'].width = 27       # KECAMATAN (Pengganti)
    ws_iiia.column_dimensions['S'].width = 50       # KABUPATEN (Pengganti)
    ws_iiia.column_dimensions['V'].width = 45       # KETERANGAN
    
    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    
    filename = f"{stage_name.upper()}_LAMPIRAN_BAHV_{batch_title}.xlsx"
    
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# ============================================================
# SK DIRJEN ENDPOINTS
# ============================================================

@app.post("/api/sk-dirjen/upload")
async def upload_sk_dirjen(
    stage_name: str = Form(...),
    province_id: int = Form(1),
    file: UploadFile = File(...)
):
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="File harus berformat .xlsx atau .xls")
    
    content = await file.read()
    wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
    ws = wb.active
    
    rows = list(ws.iter_rows(values_only=True))
    headers = None
    header_idx = None
    for i, row in enumerate(rows):
        cleaned = [str(c).strip().replace('\n', ' ').upper() if c else '' for c in row]
        if 'NAMA' in cleaned and any('KTP' in c or 'NIK' in c for c in cleaned):
            header_idx = i
            headers = [str(c).strip().replace('\n', ' ') if c else f'COL_{j}' for j, c in enumerate(row)]
            break
    
    if header_idx is None:
        raise HTTPException(status_code=400, detail="Header tidak ditemukan. Pastikan kolom NAMA dan NO. KTP ada.")
    
    data_rows = []
    for row in rows[header_idx+1:]:
        if any(c is not None for c in row):
            data_rows.append(row)
    
    if not data_rows:
        raise HTTPException(status_code=400, detail="Tidak ada data ditemukan di file")
    
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT id FROM sk_dirjen_batches WHERE stage_name = ? AND province_id = ?", (stage_name, province_id))
    existing = cursor.fetchone()
    if existing:
        cursor.execute("DELETE FROM sk_dirjen_matches WHERE sk_record_id IN (SELECT id FROM sk_dirjen_records WHERE batch_id = ?)", (existing['id'],))
        cursor.execute("DELETE FROM sk_dirjen_records WHERE batch_id = ?", (existing['id'],))
        cursor.execute("DELETE FROM sk_dirjen_batches WHERE id = ?", (existing['id'],))
        conn.commit()
    
    cursor.execute("INSERT INTO sk_dirjen_batches (stage_name, filename, province_id) VALUES (?, ?, ?)", (stage_name, file.filename, province_id))
    batch_id = cursor.lastrowid
    
    header_lower = [h.lower().replace(' ', '').replace('.', '') for h in headers]
    
    def find_col(keywords):
        for i, h in enumerate(header_lower):
            if any(k in h for k in keywords):
                return i
        return None
    
    col_nama = find_col(['nama'])
    col_nik = find_col(['ktp', 'nik'])
    col_kk = find_col(['kk', 'kartu'])
    col_jk = find_col(['jenis', 'kelamin', 'jk'])
    col_alamat = find_col(['alamat', 'tempat'])
    col_desa = find_col(['kelurahan'])
    if col_desa is None:
        col_desa = find_col(['desa'])
    col_kec = find_col(['kecamatan'])
    col_kab = find_col(['kabupaten', 'kota'])
    col_kode_desa = find_col(['kode'])
    col_no = find_col(['no', 'urut'])
    col_ket = find_col(['keterangan'])
    
    inserted = 0
    for row in data_rows:
        cells = list(row)
        def safe(idx):
            if idx is None or idx >= len(cells):
                return None
            v = cells[idx]
            if v is None:
                return None
            if isinstance(v, float):
                return str(int(v)) if v == int(v) else str(v)
            return str(v).strip()
        
        nama = safe(col_nama)
        nik = safe(col_nik)
        if not nama or not nik:
            continue
        
        cursor.execute("""
            INSERT INTO sk_dirjen_records 
            (batch_id, no_urut, kode_desa, nama, jenis_kelamin, no_ktp, no_kk, alamat, desa_kelurahan, kecamatan, kabupaten_kota, keterangan)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            batch_id, safe(col_no), safe(col_kode_desa), nama, safe(col_jk),
            nik, safe(col_kk), safe(col_alamat), safe(col_desa),
            safe(col_kec), safe(col_kab), safe(col_ket)
        ))
        sk_rec_id = cursor.lastrowid
        
        cursor.execute("""
            SELECT vr.id, vr.nama, vr.no_ktp, vr.no_kk 
            FROM verified_records vr
            JOIN verified_batches vb ON vb.id = vr.batch_id
            JOIN invers_stages ist ON ist.id = vb.stage_id
            WHERE vr.no_ktp = ? AND vr.status = 'LOLOS'
              AND (ist.province_id = ? OR (? = 1 AND (ist.province_id IS NULL OR ist.province_id = 0)))
            ORDER BY vb.id DESC NULLS LAST
            LIMIT 1
        """, (nik, province_id, province_id))
        matched = cursor.fetchone()
        
        if matched:
            nama_match = (nama.upper().strip() == (matched['nama'] or '').upper().strip())
            kk_match = (safe(col_kk) or '') == (matched['no_kk'] or '')
            
            if nama_match and kk_match:
                match_type = 'PERFECT'
            else:
                match_type = 'NEEDS_APPROVAL'
            
            cursor.execute("""
                SELECT vb.id as batch_id, vb.name as batch_name, vb.stage_id, ist.name as stage_name
                FROM verified_batches vb
                JOIN invers_stages ist ON vb.stage_id = ist.id
                WHERE vb.id = (SELECT batch_id FROM verified_records WHERE id = ?)
            """, (matched['id'],))
            batch_info = cursor.fetchone()
            
            if not batch_info:
                cursor.execute("""
                    SELECT vr.batch_id, vb.name as batch_name, vb.stage_id, ist.name as stage_name
                    FROM verified_records vr
                    LEFT JOIN verified_batches vb ON vb.id = vr.batch_id
                    LEFT JOIN invers_stages ist ON ist.id = vb.stage_id
                    WHERE vr.id = ?
                """, (matched['id'],))
                batch_info = cursor.fetchone()
            
            cursor.execute("""
                INSERT INTO sk_dirjen_matches (sk_record_id, verified_record_id, verified_batch_id, verified_stage_id, match_type)
                VALUES (?, ?, ?, ?, ?)
            """, (sk_rec_id, matched['id'], 
                  batch_info['batch_id'] if batch_info and batch_info['batch_id'] else None, 
                  batch_info['stage_id'] if batch_info and batch_info['stage_id'] else None, 
                  match_type))
        else:
            cursor.execute("INSERT INTO sk_dirjen_matches (sk_record_id, match_type) VALUES (?, 'NO_MATCH')", (sk_rec_id,))
        
        inserted += 1
    
    conn.commit()
    conn.close()
    
    return {
        "batch_id": batch_id,
        "stage_name": stage_name,
        "inserted_records": inserted,
        "total_rows": len(data_rows)
    }

@app.get("/api/sk-dirjen/batches")
def get_sk_dirjen_batches(province_id: int = 1):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    prov_sql = "WHERE (b.province_id = 1 OR b.province_id IS NULL OR b.province_id = 0)" if (not province_id or province_id == 1) else "WHERE b.province_id = ?"
    params = () if (not province_id or province_id == 1) else (province_id,)

    cursor.execute(f"""
        SELECT b.*, 
            (SELECT COUNT(*) FROM sk_dirjen_records WHERE batch_id = b.id) as total_records,
            (SELECT COUNT(*) FROM sk_dirjen_matches m JOIN sk_dirjen_records r ON m.sk_record_id = r.id WHERE r.batch_id = b.id AND m.match_type = 'PERFECT') as perfect_count,
            (SELECT COUNT(*) FROM sk_dirjen_matches m JOIN sk_dirjen_records r ON m.sk_record_id = r.id WHERE r.batch_id = b.id AND m.match_type = 'NEEDS_APPROVAL') as needs_approval_count,
            (SELECT COUNT(*) FROM sk_dirjen_matches m JOIN sk_dirjen_records r ON m.sk_record_id = r.id WHERE r.batch_id = b.id AND m.match_type = 'NO_MATCH') as no_match_count
        FROM sk_dirjen_batches b {prov_sql} ORDER BY b.id DESC
    """, params)
    batches = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return {"batches": batches}

@app.get("/api/sk-dirjen/search-verified")
def search_verified_for_pairing(q: str = "", desa: str = ""):
    if not q or len(q.strip()) < 2:
        return {"records": []}
    
    conn = get_db_connection()
    cursor = conn.cursor()
    
    term = f"%{q.upper()}%"
    if desa and desa.strip():
        desa_term = f"%{desa.upper().strip()}%"
        cursor.execute("""
            SELECT vr.id, vr.nama, vr.no_ktp, vr.no_kk, vr.desa_kelurahan, 
                   vr.kecamatan, vr.kabupaten_kota, vr.status, vb.name as batch_name, ist.name as tahap
            FROM verified_records vr
            JOIN verified_batches vb ON vb.id = vr.batch_id
            JOIN invers_stages ist ON ist.id = vb.stage_id
            WHERE UPPER(vr.nama) LIKE ? AND UPPER(vr.desa_kelurahan) LIKE ? AND vr.status = 'LOLOS'
            ORDER BY vr.nama
            LIMIT 20
        """, (term, desa_term))
    else:
        cursor.execute("""
            SELECT vr.id, vr.nama, vr.no_ktp, vr.no_kk, vr.desa_kelurahan, 
                   vr.kecamatan, vr.kabupaten_kota, vr.status, vb.name as batch_name, ist.name as tahap
            FROM verified_records vr
            JOIN verified_batches vb ON vb.id = vr.batch_id
            JOIN invers_stages ist ON ist.id = vb.stage_id
            WHERE (UPPER(vr.nama) LIKE ? OR vr.no_ktp LIKE ? OR vr.no_kk LIKE ?
                   OR UPPER(vr.desa_kelurahan) LIKE ?) AND vr.status = 'LOLOS'
            ORDER BY vr.nama
            LIMIT 20
        """, (term, term, term, term))
    
    records = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return {"records": records}

@app.post("/api/sk-dirjen/pair/{sk_record_id}")
def pair_sk_dirjen_record(sk_record_id: int, body: dict = Body(...)):
    verified_record_id = body.get("verified_record_id")
    if not verified_record_id:
        raise HTTPException(status_code=400, detail="verified_record_id harus diisi")
    
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT id FROM sk_dirjen_records WHERE id = ?", (sk_record_id,))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="Record SK Dirjen tidak ditemukan")
    
    cursor.execute("SELECT id, status FROM verified_records WHERE id = ?", (verified_record_id,))
    vr_check = cursor.fetchone()
    if not vr_check:
        conn.close()
        raise HTTPException(status_code=404, detail="Record Terverifikasi tidak ditemukan")
    if vr_check['status'] != 'LOLOS':
        conn.close()
        raise HTTPException(status_code=400, detail="Hanya bisa memasangkan dengan record yang LOLOS")
    
    cursor.execute("""
        SELECT vb.id as vb_id, vb.stage_id as vs_id
        FROM verified_records vr
        JOIN verified_batches vb ON vb.id = vr.batch_id
        WHERE vr.id = ?
    """, (verified_record_id,))
    vb_row = cursor.fetchone()
    
    cursor.execute("""
        INSERT INTO sk_dirjen_matches (sk_record_id, verified_record_id, verified_batch_id, verified_stage_id, match_type, override_status)
        VALUES (?, ?, ?, ?, 'MANUAL_PAIR', 'APPROVED')
        ON CONFLICT(sk_record_id) DO UPDATE SET
            verified_record_id = excluded.verified_record_id,
            verified_batch_id = excluded.verified_batch_id,
            verified_stage_id = excluded.verified_stage_id,
            match_type = 'MANUAL_PAIR',
            override_status = 'APPROVED'
    """, (sk_record_id, verified_record_id, vb_row['vb_id'] if vb_row else None, vb_row['vs_id'] if vb_row else None))
    
    conn.commit()
    conn.close()
    
    return {"success": True, "message": "Berhasil memasangkan data"}

@app.get("/api/sk-dirjen/{batch_id}/records")
def get_sk_dirjen_records(batch_id: int, kabupaten: str = None, kecamatan: str = None, desa: str = None, tahap: str = None, status: str = None, q: str = None, asal_batch: str = None):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    query = """
        SELECT r.*, m.match_type, m.override_status, m.verified_record_id, m.verified_batch_id, m.verified_stage_id,
            COALESCE(vb.name, vb2.name) as verified_batch_name, COALESCE(ist.name, ist2.name) as verified_stage_name,
            vr.nama as verified_nama, vr.no_ktp as verified_no_ktp, vr.no_kk as verified_no_kk,
            vr.alamat as verified_alamat, vr.desa_kelurahan as verified_desa_kelurahan,
            vr.kecamatan as verified_kecamatan, vr.kabupaten_kota as verified_kabupaten_kota
        FROM sk_dirjen_records r
        LEFT JOIN sk_dirjen_matches m ON m.sk_record_id = r.id
        LEFT JOIN verified_batches vb ON vb.id = m.verified_batch_id
        LEFT JOIN invers_stages ist ON ist.id = m.verified_stage_id
        LEFT JOIN verified_records vr ON vr.id = m.verified_record_id
        LEFT JOIN verified_batches vb2 ON vb2.id = vr.batch_id
        LEFT JOIN invers_stages ist2 ON ist2.id = vb2.stage_id
        WHERE r.batch_id = ?
    """
    params = [batch_id]
    
    if q and q.strip():
        search = f"%{q.strip().upper()}%"
        query += " AND (UPPER(r.nama) LIKE ? OR r.no_ktp LIKE ? OR r.no_kk LIKE ? OR UPPER(r.desa_kelurahan) LIKE ? OR UPPER(r.kecamatan) LIKE ? OR UPPER(r.kabupaten_kota) LIKE ?)"
        params.extend([search, search, search, search, search, search])
    
    if kabupaten:
        query += " AND UPPER(r.kabupaten_kota) = ?"
        params.append(kabupaten.upper())
    if kecamatan:
        query += " AND UPPER(r.kecamatan) = ?"
        params.append(kecamatan.upper())
    if desa:
        query += " AND UPPER(r.desa_kelurahan) = ?"
        params.append(desa.upper())
    if tahap:
        query += " AND UPPER(ist.name) = ?"
        params.append(tahap.upper())
    if asal_batch:
        query += " AND UPPER(COALESCE(vb.name, vb2.name)) = ?"
        params.append(asal_batch.upper())
    if status:
        if status == 'PERFECT':
            query += " AND m.match_type = 'PERFECT'"
        elif status == 'NEEDS_APPROVAL':
            query += " AND m.match_type = 'NEEDS_APPROVAL'"
        elif status == 'NO_MATCH':
            query += " AND m.match_type = 'NO_MATCH'"
        elif status == 'APPROVED':
            query += " AND m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED'"
        elif status == 'MANUAL_PAIR':
            query += " AND m.match_type = 'MANUAL_PAIR'"
    
    query += " ORDER BY r.no_urut, r.id"
    cursor.execute(query, params)
    records = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return {"records": records}

@app.get("/api/sk-dirjen/all-records")
def get_sk_dirjen_all_records(q: str = None, kabupaten: str = None, kecamatan: str = None, desa: str = None, status: str = None, tahap: str = None, asal_batch: str = None, province_id: int = 1):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    query = """
        SELECT r.*, m.match_type, m.override_status, m.verified_record_id, m.verified_batch_id, m.verified_stage_id,
            COALESCE(vb.name, vb2.name) as verified_batch_name, COALESCE(ist.name, ist2.name) as verified_stage_name,
            vr.nama as verified_nama, vr.no_ktp as verified_no_ktp, vr.no_kk as verified_no_kk,
            vr.alamat as verified_alamat, vr.desa_kelurahan as verified_desa_kelurahan,
            vr.kecamatan as verified_kecamatan, vr.kabupaten_kota as verified_kabupaten_kota,
            sb.stage_name as batch_stage_name
        FROM sk_dirjen_records r
        JOIN sk_dirjen_batches sb ON sb.id = r.batch_id
        LEFT JOIN sk_dirjen_matches m ON m.sk_record_id = r.id
        LEFT JOIN verified_batches vb ON vb.id = m.verified_batch_id
        LEFT JOIN invers_stages ist ON ist.id = m.verified_stage_id
        LEFT JOIN verified_records vr ON vr.id = m.verified_record_id
        LEFT JOIN verified_batches vb2 ON vb2.id = vr.batch_id
        LEFT JOIN invers_stages ist2 ON ist2.id = vb2.stage_id
        WHERE 1=1
    """
    params = []
    
    if not province_id or province_id == 1:
        query += " AND (sb.province_id = 1 OR sb.province_id IS NULL OR sb.province_id = 0)"
    else:
        query += " AND sb.province_id = ?"
        params.append(province_id)

    if q and q.strip():
        search = f"%{q.strip().upper()}%"
        query += " AND (UPPER(r.nama) LIKE ? OR r.no_ktp LIKE ? OR r.no_kk LIKE ? OR UPPER(r.desa_kelurahan) LIKE ? OR UPPER(r.kecamatan) LIKE ? OR UPPER(r.kabupaten_kota) LIKE ?)"
        params.extend([search, search, search, search, search, search])
    
    if kabupaten:
        query += " AND UPPER(r.kabupaten_kota) = ?"
        params.append(kabupaten.upper())
    if kecamatan:
        query += " AND UPPER(r.kecamatan) = ?"
        params.append(kecamatan.upper())
    if desa:
        query += " AND UPPER(r.desa_kelurahan) = ?"
        params.append(desa.upper())
    if status:
        if status == 'PERFECT':
            query += " AND m.match_type = 'PERFECT'"
        elif status == 'NEEDS_APPROVAL':
            query += " AND m.match_type = 'NEEDS_APPROVAL'"
        elif status == 'NO_MATCH':
            query += " AND m.match_type = 'NO_MATCH'"
        elif status == 'APPROVED':
            query += " AND m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED'"
        elif status == 'MANUAL_PAIR':
            query += " AND m.match_type = 'MANUAL_PAIR'"
    
    if tahap:
        query += " AND UPPER(COALESCE(ist.name, ist2.name)) = ?"
        params.append(tahap.upper())
    if asal_batch:
        query += " AND UPPER(COALESCE(vb.name, vb2.name)) = ?"
        params.append(asal_batch.upper())
    
    query += " ORDER BY sb.id DESC, r.no_urut, r.id"
    cursor.execute(query, params)
    records = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return {"records": records}

@app.post("/api/sk-dirjen/record/{record_id}/approve")
def approve_sk_dirjen_record(record_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("UPDATE sk_dirjen_matches SET override_status = 'APPROVED' WHERE sk_record_id = ?", (record_id,))
    conn.commit()
    conn.close()
    return {"success": True, "message": "Persetujuan berhasil disimpan"}

@app.post("/api/sk-dirjen/record/{record_id}/reject")
def reject_sk_dirjen_record(record_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("UPDATE sk_dirjen_matches SET override_status = 'REJECTED' WHERE sk_record_id = ?", (record_id,))
    conn.commit()
    conn.close()
    return {"success": True, "message": "Penolakan berhasil disimpan"}

@app.delete("/api/sk-dirjen/batch/{batch_id}")
def delete_sk_dirjen_batch(batch_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT id, stage_name FROM sk_dirjen_batches WHERE id = ?", (batch_id,))
    batch = cursor.fetchone()
    if not batch:
        conn.close()
        raise HTTPException(status_code=404, detail="Batch SK Dirjen tidak ditemukan")
    
    cursor.execute("DELETE FROM sk_dirjen_matches WHERE sk_record_id IN (SELECT id FROM sk_dirjen_records WHERE batch_id = ?)", (batch_id,))
    cursor.execute("DELETE FROM sk_dirjen_records WHERE batch_id = ?", (batch_id,))
    cursor.execute("DELETE FROM sk_dirjen_batches WHERE id = ?", (batch_id,))
    conn.commit()
    conn.close()
    
    return {"success": True, "message": f"Batch '{batch['stage_name']}' dan semua data terkait berhasil dihapus"}

@app.get("/api/sk-dirjen/rekap-per-tahap")
def get_sk_dirjen_rekap_per_tahap(province_id: int = 1):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    prov_sql = "WHERE (province_id = 1 OR province_id IS NULL OR province_id = 0)" if (not province_id or province_id == 1) else "WHERE province_id = ?"
    params = () if (not province_id or province_id == 1) else (province_id,)

    cursor.execute(f"SELECT id, stage_name FROM sk_dirjen_batches {prov_sql} ORDER BY id", params)
    sk_batches = [dict(row) for row in cursor.fetchall()]
    
    cursor.execute(f"SELECT id, name FROM invers_stages {prov_sql} ORDER BY id", params)
    invers_stages_raw = [dict(row) for row in cursor.fetchall()]
    
    def sort_stages(stages):
        import re
        murni = []
        pengganti = []
        for s in stages:
            if 'pengganti' in s['name'].lower():
                pengganti.append(s)
            else:
                murni.append(s)
        def stage_num(s):
            nums = re.findall(r'\d+', s['name'])
            return int(nums[0]) if nums else 999
        murni.sort(key=stage_num)
        pengganti.sort(key=stage_num)
        return murni + pengganti
    
    invers_stages = sort_stages(invers_stages_raw)
    
    rekap = []
    for sk_batch in sk_batches:
        row_data = {"batch_id": sk_batch['id'], "stage_name": sk_batch['stage_name'], "total": 0, "tahap": {}}
        
        for inv_stage in invers_stages:
            cursor.execute("""
                SELECT COUNT(*) as cnt FROM sk_dirjen_matches m
                JOIN sk_dirjen_records r ON m.sk_record_id = r.id
                LEFT JOIN verified_records vr ON vr.id = m.verified_record_id
                LEFT JOIN verified_batches vb ON vb.id = COALESCE(vr.batch_id, m.verified_batch_id)
                WHERE r.batch_id = ? AND vb.stage_id = ?
                AND (m.match_type = 'PERFECT' OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED') OR m.match_type = 'MANUAL_PAIR')
            """, (sk_batch['id'], inv_stage['id']))
            cnt = cursor.fetchone()['cnt']
            if cnt > 0:
                row_data['tahap'][inv_stage['name']] = cnt
                row_data['total'] += cnt
        
        rekap.append(row_data)
    
    conn.close()
    return {"rekap": rekap, "invers_stages": [s['name'] for s in invers_stages]}

@app.get("/api/sk-dirjen/rekap-per-kabupaten/all")
def get_sk_dirjen_rekap_per_kabupaten_all(province_id: int = 1):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    prov_b_sql = "WHERE (b.province_id = 1 OR b.province_id IS NULL OR b.province_id = 0)" if (not province_id or province_id == 1) else "WHERE b.province_id = ?"
    prov_s_sql = "WHERE (s.province_id = 1 OR s.province_id IS NULL OR s.province_id = 0)" if (not province_id or province_id == 1) else "WHERE s.province_id = ?"
    params_b = () if (not province_id or province_id == 1) else (province_id,)
    params_s = () if (not province_id or province_id == 1) else (province_id,)

    cursor.execute(f"SELECT DISTINCT UPPER(r.kabupaten_kota) as kab FROM sk_dirjen_records r JOIN sk_dirjen_batches b ON r.batch_id = b.id {prov_b_sql} ORDER BY kab", params_b)
    kabupatens = [row['kab'] for row in cursor.fetchall()]
    
    cursor.execute(f"""
        SELECT DISTINCT ist.name as inv_stage_name
        FROM sk_dirjen_matches m
        JOIN sk_dirjen_records r ON m.sk_record_id = r.id
        JOIN sk_dirjen_batches b ON r.batch_id = b.id
        LEFT JOIN verified_records vr ON vr.id = m.verified_record_id
        LEFT JOIN verified_batches vb ON vb.id = COALESCE(vr.batch_id, m.verified_batch_id)
        JOIN invers_stages ist ON ist.id = vb.stage_id
        {prov_b_sql} AND (m.match_type = 'PERFECT' OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED') OR m.match_type = 'MANUAL_PAIR')
        ORDER BY ist.name
    """, params_b)
    inv_stages_raw = [row['inv_stage_name'] for row in cursor.fetchall()]
    
    import re
    def sort_stage_names(names):
        murni = []
        pengganti = []
        for n in names:
            if 'pengganti' in n.lower():
                pengganti.append(n)
            else:
                murni.append(n)
        def stage_num(n):
            nums = re.findall(r'\d+', n)
            return int(nums[0]) if nums else 999
        murni.sort(key=stage_num)
        pengganti.sort(key=stage_num)
        return murni + pengganti
    
    inv_stages = sort_stage_names(inv_stages_raw)
    
    cursor.execute(f"""
        SELECT UPPER(TRIM(COALESCE(vr.kabupaten_kota, ''))) as kab, COUNT(*) as cnt
        FROM verified_records vr
        JOIN verified_batches vb ON vb.id = vr.batch_id
        JOIN invers_stages s ON s.id = vb.stage_id
        {prov_s_sql} AND vr.status = 'LOLOS' AND vb.is_published = 1
        GROUP BY kab
    """, params_s)
    lolos_by_kab = {row['kab']: row['cnt'] for row in cursor.fetchall()}
    
    table_data = []
    for kab in kabupatens:
        kab_row = {"kabupaten": kab}
        total = 0
        
        for inv_stage in inv_stages:
            cursor.execute("""
                SELECT COUNT(*) as cnt FROM sk_dirjen_matches m
                JOIN sk_dirjen_records r ON m.sk_record_id = r.id
                LEFT JOIN verified_records vr ON vr.id = m.verified_record_id
                LEFT JOIN verified_batches vb ON vb.id = COALESCE(vr.batch_id, m.verified_batch_id)
                JOIN invers_stages ist ON ist.id = vb.stage_id
                WHERE UPPER(r.kabupaten_kota) = ? AND ist.name = ?
                AND (m.match_type = 'PERFECT' OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED') OR m.match_type = 'MANUAL_PAIR')
            """, (kab, inv_stage))
            cnt = cursor.fetchone()['cnt']
            kab_row[inv_stage] = cnt
            total += cnt
        
        cpb_lolos = lolos_by_kab.get(kab, 0)
        kab_row['cpb_lolos'] = cpb_lolos
        kab_row['total'] = total
        kab_row['selisih_sk'] = cpb_lolos - total
        
        sumber_list = []
        if cpb_lolos - total > 0:
            cursor.execute("""
                SELECT vb.name as batch_name, ist.name as stage_name, COUNT(*) as cnt
                FROM verified_records vr
                JOIN verified_batches vb ON vr.batch_id = vb.id
                JOIN invers_stages ist ON vb.stage_id = ist.id
                WHERE vr.status = 'LOLOS' 
                AND vb.is_published = 1
                AND UPPER(TRIM(vr.kabupaten_kota)) = ?
                AND vr.id NOT IN (
                    SELECT m.verified_record_id FROM sk_dirjen_matches m
                    WHERE m.verified_record_id IS NOT NULL
                    AND (m.match_type = 'PERFECT' 
                        OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED')
                        OR m.match_type = 'MANUAL_PAIR')
                )
                GROUP BY vb.name, ist.name
                ORDER BY ist.name, vb.name
            """, (kab,))
            sumber_list = [dict(row) for row in cursor.fetchall()]
        kab_row['sumber_selisih'] = sumber_list
        
        table_data.append(kab_row)
    
    conn.close()
    return {
        "batch_id": None,
        "stage_name": "SEMUA BATCH",
        "kabupatens": table_data,
        "invers_stages": inv_stages
    }

@app.get("/api/sk-dirjen/sumber-selisih-detail")
def get_sumber_selisih_detail(kabupaten: str, batch_name: str, stage_name: str):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT vr.nama, vr.no_ktp, vr.no_kk, vr.desa_kelurahan, 
               vr.kecamatan, vr.kabupaten_kota
        FROM verified_records vr
        JOIN verified_batches vb ON vr.batch_id = vb.id
        JOIN invers_stages ist ON vb.stage_id = ist.id
        WHERE vr.status = 'LOLOS' 
        AND vb.is_published = 1
        AND UPPER(TRIM(vr.kabupaten_kota)) = UPPER(TRIM(?))
        AND vb.name = ?
        AND ist.name = ?
        AND vr.id NOT IN (
            SELECT m.verified_record_id FROM sk_dirjen_matches m
            WHERE m.verified_record_id IS NOT NULL
            AND (m.match_type = 'PERFECT' 
                OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED')
                OR m.match_type = 'MANUAL_PAIR')
        )
        ORDER BY vr.nama
    """, (kabupaten, batch_name, stage_name))
    
    records = [dict(row) for row in cursor.fetchall()]
    conn.close()
    
    return {"records": records}

@app.get("/api/sk-dirjen/rekap-per-kabupaten/{batch_id}")
def get_sk_dirjen_rekap_per_kabupaten(batch_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT id, stage_name FROM sk_dirjen_batches ORDER BY id")
    all_sk_batches = [dict(row) for row in cursor.fetchall()]
    
    cursor.execute("SELECT DISTINCT UPPER(kabupaten_kota) as kab FROM sk_dirjen_records WHERE batch_id = ?", (batch_id,))
    kabupatens = sorted([row['kab'] for row in cursor.fetchall()])
    
    current_batch = None
    for b in all_sk_batches:
        if b['id'] == batch_id:
            current_batch = b
            break
    
    cursor.execute("""
        SELECT DISTINCT ist.name as inv_stage_name
        FROM sk_dirjen_matches m
        JOIN sk_dirjen_records r ON m.sk_record_id = r.id
        LEFT JOIN verified_records vr ON vr.id = m.verified_record_id
        LEFT JOIN verified_batches vb ON vb.id = COALESCE(vr.batch_id, m.verified_batch_id)
        JOIN invers_stages ist ON ist.id = vb.stage_id
        WHERE r.batch_id = ? AND (m.match_type = 'PERFECT' OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED') OR m.match_type = 'MANUAL_PAIR')
        ORDER BY ist.name
    """, (batch_id,))
    inv_stages_in_match_raw = [row['inv_stage_name'] for row in cursor.fetchall()]
    
    import re
    def sort_stage_names(names):
        murni = []
        pengganti = []
        for n in names:
            if 'pengganti' in n.lower():
                pengganti.append(n)
            else:
                murni.append(n)
        def stage_num(n):
            nums = re.findall(r'\d+', n)
            return int(nums[0]) if nums else 999
        murni.sort(key=stage_num)
        pengganti.sort(key=stage_num)
        return murni + pengganti
    
    inv_stages_in_match = sort_stage_names(inv_stages_in_match_raw)
    
    table_data = []
    for kab in kabupatens:
        kab_row = {"kabupaten": kab}
        total = 0
        
        for inv_stage in inv_stages_in_match:
            cursor.execute("""
                SELECT COUNT(*) as cnt FROM sk_dirjen_matches m
                JOIN sk_dirjen_records r ON m.sk_record_id = r.id
                LEFT JOIN verified_records vr ON vr.id = m.verified_record_id
                LEFT JOIN verified_batches vb ON vb.id = COALESCE(vr.batch_id, m.verified_batch_id)
                JOIN invers_stages ist ON ist.id = vb.stage_id
                WHERE r.batch_id = ? AND UPPER(r.kabupaten_kota) = ? AND ist.name = ?
                AND (m.match_type = 'PERFECT' OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED') OR m.match_type = 'MANUAL_PAIR')
            """, (batch_id, kab, inv_stage))
            cnt = cursor.fetchone()['cnt']
            kab_row[inv_stage] = cnt
            total += cnt
        
        kab_row['total'] = total
        table_data.append(kab_row)
    
    conn.close()
    return {
        "batch_id": batch_id,
        "stage_name": current_batch['stage_name'] if current_batch else "",
        "kabupatens": table_data,
        "invers_stages": inv_stages_in_match
    }

@app.get("/api/sk-dirjen/export/{batch_id}")
def export_sk_dirjen_daftar_pb(batch_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT stage_name FROM sk_dirjen_batches WHERE id = ?", (batch_id,))
    batch = cursor.fetchone()
    if not batch:
        conn.close()
        raise HTTPException(status_code=404, detail="Batch SK Dirjen tidak ditemukan")
    stage_name = batch['stage_name']
    
    cursor.execute("""
        SELECT r.*, m.match_type, m.override_status, vb.name as verified_batch_name, ist.name as verified_stage_name
        FROM sk_dirjen_records r
        LEFT JOIN sk_dirjen_matches m ON m.sk_record_id = r.id
        LEFT JOIN verified_batches vb ON vb.id = m.verified_batch_id
        LEFT JOIN invers_stages ist ON ist.id = m.verified_stage_id
        WHERE r.batch_id = ?
        ORDER BY r.no_urut, r.id
    """, (batch_id,))
    records = [dict(row) for row in cursor.fetchall()]
    conn.close()
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "DAFTAR PB SK DIRJEN"
    
    headers_excel = ["NO", "KODE DESA/KEL", "NAMA", "JENIS KELAMIN", "NO. KTP", "NO. KK", "ALAMAT TEMPAT TINGGAL", "DESA/KELURAHAN", "KECAMATAN", "KABUPATEN/KOTA", "KETERANGAN", "STATUS", "ASAL VERIFIKASI"]
    
    f_header = Font(name='Arial', size=11, bold=True)
    fill_header = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
    a_header = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    
    for col_idx, h in enumerate(headers_excel, 1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.font = f_header
        cell.fill = fill_header
        cell.alignment = a_header
        cell.border = thin_border
    
    fill_perfect = PatternFill(start_color='DCfce7', end_color='DCfce7', fill_type='solid')
    fill_needs = PatternFill(start_color='fef9c3', end_color='fef9c3', fill_type='solid')
    fill_none = PatternFill(start_color='fee2e2', end_color='fee2e2', fill_type='solid')
    f_body = Font(name='Arial', size=11)
    
    for i, r in enumerate(records, 2):
        status_label = {'PERFECT': 'Cocok', 'NEEDS_APPROVAL': 'Perlu Persetujuan', 'NO_MATCH': 'Tidak Ditemukan'}.get(r['match_type'], '')
        asal = ''
        if r['match_type'] in ('PERFECT', 'NEEDS_APPROVAL') and r['verified_batch_name']:
            asal = f"BA {r['verified_batch_name']} Tahap {r['verified_stage_name']}"
        
        vals = [r['no_urut'], r['kode_desa'], r['nama'], r['jenis_kelamin'], r['no_ktp'], r['no_kk'], r['alamat'], r['desa_kelurahan'], r['kecamatan'], r['kabupaten_kota'], r['keterangan'], status_label, asal]
        
        row_fill = fill_perfect if r['match_type'] == 'PERFECT' else (fill_needs if r['match_type'] == 'NEEDS_APPROVAL' else fill_none)
        
        for col_idx, v in enumerate(vals, 1):
            cell = ws.cell(row=i, column=col_idx, value=v)
            cell.font = f_body
            cell.border = thin_border
            if col_idx == 12:
                cell.fill = row_fill
    
    col_widths = [6, 18, 30, 16, 20, 20, 30, 22, 22, 22, 25, 18, 35]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    
    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    
    safe_stage = stage_name.replace(' ', '_').upper()
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={safe_stage}_DAFTAR_PB_SK_DIRJEN.xlsx"}
    )

@app.get("/api/sk-dirjen/rekap-per-tahap/export")
def export_sk_dirjen_rekap_per_tahap():
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT id, stage_name FROM sk_dirjen_batches ORDER BY id")
    sk_batches = [dict(row) for row in cursor.fetchall()]
    
    cursor.execute("SELECT id, name FROM invers_stages ORDER BY id")
    invers_stages = [dict(row) for row in cursor.fetchall()]
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "REKAP PB PER TAHAP"
    
    f_header = Font(name='Arial', size=11, bold=True)
    fill_header = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
    a_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    f_total = Font(name='Arial', size=11, bold=True)
    
    headers_excel = ["TAHAP SK DIRJEN"] + [s['name'] for s in invers_stages] + ["TOTAL"]
    for col_idx, h in enumerate(headers_excel, 1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.font = f_header
        cell.fill = fill_header
        cell.alignment = a_center
        cell.border = thin_border
    
    for row_idx, sk_batch in enumerate(sk_batches, 2):
        ws.cell(row=row_idx, column=1, value=sk_batch['stage_name']).font = f_header
        ws.cell(row=row_idx, column=1).border = thin_border
        total = 0
        for col_idx, inv_stage in enumerate(invers_stages, 2):
            cursor.execute("""
                SELECT COUNT(*) as cnt FROM sk_dirjen_matches m
                JOIN sk_dirjen_records r ON m.sk_record_id = r.id
                LEFT JOIN verified_records vr ON vr.id = m.verified_record_id
                LEFT JOIN verified_batches vb ON vb.id = COALESCE(vr.batch_id, m.verified_batch_id)
                WHERE r.batch_id = ? AND vb.stage_id = ?
                AND (m.match_type = 'PERFECT' OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED') OR m.match_type = 'MANUAL_PAIR')
            """, (sk_batch['id'], inv_stage['id']))
            cnt = cursor.fetchone()['cnt']
            ws.cell(row=row_idx, column=col_idx, value=cnt if cnt > 0 else None).font = f_body
            ws.cell(row=row_idx, column=col_idx).border = thin_border
            ws.cell(row=row_idx, column=col_idx).alignment = a_center
            total += cnt
        ws.cell(row=row_idx, column=len(invers_stages)+2, value=total).font = f_total
        ws.cell(row=row_idx, column=len(invers_stages)+2).border = thin_border
        ws.cell(row=row_idx, column=len(invers_stages)+2).alignment = a_center
    
    total_row = len(sk_batches) + 2
    ws.cell(row=total_row, column=1, value="TOTAL").font = f_total
    ws.cell(row=total_row, column=1).border = thin_border
    for col_idx, inv_stage in enumerate(invers_stages, 2):
        cursor.execute("""
            SELECT COUNT(*) as cnt FROM sk_dirjen_matches m
            JOIN sk_dirjen_records r ON m.sk_record_id = r.id
            LEFT JOIN verified_records vr ON vr.id = m.verified_record_id
            LEFT JOIN verified_batches vb ON vb.id = COALESCE(vr.batch_id, m.verified_batch_id)
            WHERE vb.stage_id = ?
            AND (m.match_type = 'PERFECT' OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED') OR m.match_type = 'MANUAL_PAIR')
        """, (inv_stage['id'],))
        cnt = cursor.fetchone()['cnt']
        ws.cell(row=total_row, column=col_idx, value=cnt if cnt > 0 else None).font = f_total
        ws.cell(row=total_row, column=col_idx).border = thin_border
        ws.cell(row=total_row, column=col_idx).alignment = a_center
    grand_total = sum(ws.cell(row=r, column=len(invers_stages)+2).value or 0 for r in range(2, total_row))
    ws.cell(row=total_row, column=len(invers_stages)+2, value=grand_total).font = f_total
    ws.cell(row=total_row, column=len(invers_stages)+2).border = thin_border
    
    ws.column_dimensions['A'].width = 25
    for i in range(2, len(invers_stages)+3):
        ws.column_dimensions[get_column_letter(i)].width = 18
    
    conn.close()
    
    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=REKAP_PB_SK_DIRJEN_PER_TAHAP.xlsx"}
    )

@app.get("/api/sk-dirjen/rekap-per-kabupaten/{batch_id}/export")
def export_sk_dirjen_rekap_per_kabupaten(batch_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT stage_name FROM sk_dirjen_batches WHERE id = ?", (batch_id,))
    batch = cursor.fetchone()
    if not batch:
        conn.close()
        raise HTTPException(status_code=404, detail="Batch SK Dirjen tidak ditemukan")
    stage_name = batch['stage_name']
    
    cursor.execute("SELECT DISTINCT UPPER(kabupaten_kota) as kab FROM sk_dirjen_records WHERE batch_id = ?", (batch_id,))
    kabupatens = sorted([row['kab'] for row in cursor.fetchall()])
    
    cursor.execute("""
        SELECT DISTINCT ist.name as inv_stage_name
        FROM sk_dirjen_matches m
        JOIN sk_dirjen_records r ON m.sk_record_id = r.id
        LEFT JOIN verified_records vr ON vr.id = m.verified_record_id
        LEFT JOIN verified_batches vb ON vb.id = COALESCE(vr.batch_id, m.verified_batch_id)
        JOIN invers_stages ist ON ist.id = vb.stage_id
        WHERE r.batch_id = ? AND (m.match_type = 'PERFECT' OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED') OR m.match_type = 'MANUAL_PAIR')
        ORDER BY ist.name
    """, (batch_id,))
    inv_stages_raw = [row['inv_stage_name'] for row in cursor.fetchall()]
    
    import re
    def sort_stage_names(names):
        murni = []
        pengganti = []
        for n in names:
            if 'pengganti' in n.lower():
                pengganti.append(n)
            else:
                murni.append(n)
        def stage_num(n):
            nums = re.findall(r'\d+', n)
            return int(nums[0]) if nums else 999
        murni.sort(key=stage_num)
        pengganti.sort(key=stage_num)
        return murni + pengganti
    
    inv_stages = sort_stage_names(inv_stages_raw)
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "REKAP PB PER KABUPATEN"
    
    f_header = Font(name='Arial', size=11, bold=True)
    fill_header = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
    a_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    f_total = Font(name='Arial', size=11, bold=True)
    f_body = Font(name='Arial', size=11)
    
    ws.merge_cells(start_row=1, start_column=1, end_row=2, end_column=1)
    ws.merge_cells(start_row=1, start_column=2, end_row=2, end_column=2)
    ws.merge_cells(start_row=1, start_column=len(inv_stages)+3, end_row=2, end_column=len(inv_stages)+3)
    
    cell_no = ws.cell(row=1, column=1, value="No.")
    cell_no.font = f_header
    cell_no.fill = fill_header
    cell_no.alignment = a_center
    cell_no.border = thin_border
    
    cell_kab = ws.cell(row=1, column=2, value="KABUPATEN/KOTA")
    cell_kab.font = f_header
    cell_kab.fill = fill_header
    cell_kab.alignment = a_center
    cell_kab.border = thin_border
    
    cell_total_header = ws.cell(row=1, column=len(inv_stages)+3, value="TOTAL")
    cell_total_header.font = f_header
    cell_total_header.fill = fill_header
    cell_total_header.alignment = a_center
    cell_total_header.border = thin_border
    
    ws.cell(row=2, column=1).border = thin_border
    ws.cell(row=2, column=2).border = thin_border
    ws.cell(row=2, column=len(inv_stages)+3).border = thin_border
    
    for col_idx, inv_name in enumerate(inv_stages, 3):
        cell = ws.cell(row=1, column=col_idx, value=f"SK DIRJEN {stage_name.upper()}")
        cell.font = f_header
        cell.fill = fill_header
        cell.alignment = a_center
        cell.border = thin_border
        
        cell2 = ws.cell(row=2, column=col_idx, value=inv_name)
        cell2.font = f_header
        cell2.fill = fill_header
        cell2.alignment = a_center
        cell2.border = thin_border
    
    for row_idx, kab in enumerate(kabupatens, 3):
        ws.cell(row=row_idx, column=1, value=row_idx-2).font = f_body
        ws.cell(row=row_idx, column=1).border = thin_border
        ws.cell(row=row_idx, column=1).alignment = a_center
        
        ws.cell(row=row_idx, column=2, value=kab).font = f_body
        ws.cell(row=row_idx, column=2).border = thin_border
        
        total = 0
        for col_idx, inv_name in enumerate(inv_stages, 3):
            cursor.execute("""
                SELECT COUNT(*) as cnt FROM sk_dirjen_matches m
                JOIN sk_dirjen_records r ON m.sk_record_id = r.id
                LEFT JOIN verified_records vr ON vr.id = m.verified_record_id
                LEFT JOIN verified_batches vb ON vb.id = COALESCE(vr.batch_id, m.verified_batch_id)
                JOIN invers_stages ist ON ist.id = vb.stage_id
                WHERE r.batch_id = ? AND UPPER(r.kabupaten_kota) = ? AND ist.name = ?
                AND (m.match_type = 'PERFECT' OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED') OR m.match_type = 'MANUAL_PAIR')
            """, (batch_id, kab, inv_name))
            cnt = cursor.fetchone()['cnt']
            ws.cell(row=row_idx, column=col_idx, value=cnt if cnt > 0 else None).font = f_body
            ws.cell(row=row_idx, column=col_idx).border = thin_border
            ws.cell(row=row_idx, column=col_idx).alignment = a_center
            total += cnt
        ws.cell(row=row_idx, column=len(inv_stages)+3, value=total).font = f_total
        ws.cell(row=row_idx, column=len(inv_stages)+3).border = thin_border
        ws.cell(row=row_idx, column=len(inv_stages)+3).alignment = a_center
    
    total_row = len(kabupatens) + 3
    ws.cell(row=total_row, column=1, value="").border = thin_border
    ws.cell(row=total_row, column=2, value="TOTAL").font = f_total
    ws.cell(row=total_row, column=2).border = thin_border
    for col_idx, inv_name in enumerate(inv_stages, 3):
        cursor.execute("""
            SELECT COUNT(*) as cnt FROM sk_dirjen_matches m
            JOIN sk_dirjen_records r ON m.sk_record_id = r.id
            LEFT JOIN verified_records vr ON vr.id = m.verified_record_id
            LEFT JOIN verified_batches vb ON vb.id = COALESCE(vr.batch_id, m.verified_batch_id)
            JOIN invers_stages ist ON ist.id = vb.stage_id
            WHERE r.batch_id = ? AND ist.name = ?
            AND (m.match_type = 'PERFECT' OR (m.match_type = 'NEEDS_APPROVAL' AND m.override_status = 'APPROVED') OR m.match_type = 'MANUAL_PAIR')
        """, (batch_id, inv_name))
        cnt = cursor.fetchone()['cnt']
        ws.cell(row=total_row, column=col_idx, value=cnt if cnt > 0 else None).font = f_total
        ws.cell(row=total_row, column=col_idx).border = thin_border
        ws.cell(row=total_row, column=col_idx).alignment = a_center
    grand_total = sum(ws.cell(row=r, column=len(inv_stages)+3).value or 0 for r in range(3, total_row))
    ws.cell(row=total_row, column=len(inv_stages)+3, value=grand_total).font = f_total
    ws.cell(row=total_row, column=len(inv_stages)+3).border = thin_border
    
    ws.column_dimensions['A'].width = 6
    ws.column_dimensions['B'].width = 30
    for i in range(3, len(inv_stages)+4):
        ws.column_dimensions[get_column_letter(i)].width = 22
    
    conn.close()
    
    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    
    safe_stage = stage_name.replace(' ', '_').upper()
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={safe_stage}_REKAP_PB_PER_KABUPATEN.xlsx"}
    )

if __name__ == "__main__":
    import os
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run("app:app", host="0.0.0.0", port=port)
